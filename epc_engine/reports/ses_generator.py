"""
Generátor příloh ke Smlouvě o energetických službách (SES / EPC).

Výstup: ZIP archiv s 9 Word dokumenty:
  Příloha_1_Výchozí_stav.docx
  Příloha_2_Popis_opatření.docx
  Příloha_3_Cena_a_úhrada.docx
  Příloha_4_Harmonogram.docx
  Příloha_5_Garantovaná_úspora.docx
  Příloha_6_Vyhodnocování_úspor_MV.docx
  Příloha_7_Energetický_management.docx
  Příloha_8_Oprávněné_osoby.docx
  Příloha_9_Poddodavatelé.docx
"""
from __future__ import annotations

import zipfile
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from epc_engine.models import BuildingInfo, ProjectResult
from epc_engine.op_descriptions import OP_INFO

from . import _styles as S
from ._tables import _header_row, _data_row

_TEMPLATE = __import__("pathlib").Path(__file__).parent / "templates" / "ep_template.docx"

# ── Sdílené pomocné funkce ────────────────────────────────────────────────────

def _doc() -> Document:
    return Document(_TEMPLATE)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 1")
    for run in p.runs:
        run.font.name = S.FONT_BODY
        run.font.color.rgb = S.COLOR_H1


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 2")
    for run in p.runs:
        run.font.name = S.FONT_BODY
        run.font.color.rgb = S.COLOR_H2


def _p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.style.font.name = S.FONT_BODY
    if bold:
        for run in p.runs:
            run.bold = True


_DZ = "<Doplní zhotovitel>"   # standardní placeholder pro ESCO


def _ph(doc: Document, text: str = "") -> None:
    """Placeholder – zvýrazněný text pro místo k doplnění zhotovitelem."""
    label = text if text else _DZ
    p = doc.add_paragraph()
    run = p.add_run(f"<Doplní zhotovitel: {label}>")
    run.font.name = S.FONT_BODY
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)   # oranžová – dobře viditelné


def _tbl(doc: Document, headers: list[str], rows: list[list[str]],
         widths: list[float] | None = None,
         aligns: list | None = None) -> None:
    n = len(headers)
    if widths is None:
        widths = [16.0 / n] * n
    table = doc.add_table(rows=1, cols=n)
    table.style = "Table Grid"
    for col, w in zip(table.columns, widths):
        col.width = Cm(w)
    _header_row(table, headers)
    al = aligns or [WD_ALIGN_PARAGRAPH.LEFT] * n
    for row in rows:
        _data_row(table, [str(c) for c in row], aligns=al)
    doc.add_paragraph()


def _fmt(v: float, decimals: int = 0) -> str:
    fmt = f"{{:,.{decimals}f}}"
    return fmt.format(v).replace(",", "\u00a0")


def _titulni_strana(doc: Document, cislo: int, nazev: str, ss: dict,
                    budova: BuildingInfo) -> None:
    """Titulní strana každé přílohy."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Příloha č. {cislo}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = S.FONT_BODY
    run.font.color.rgb = S.COLOR_H1

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(nazev)
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = S.FONT_BODY
    run2.font.color.rgb = S.COLOR_H2

    doc.add_paragraph()
    for label, value in [
        ("Objekt:", f"{budova.objekt_nazev or ''}  {budova.objekt_adresa or ''}".strip()),
        ("Zadavatel:", ss.get("zadavatel_nazev", "")),
        ("ESCO / zhotovitel:", _DZ),
        ("Datum:", ss.get("datum", "")),
    ]:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p3.add_run(f"{label}  ")
        r1.bold = True
        r1.font.name = S.FONT_BODY
        r1.font.size = Pt(11)
        r2 = p3.add_run(value)
        r2.font.name = S.FONT_BODY
        r2.font.size = Pt(11)
    doc.add_page_break()


def _save(doc: Document) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Příloha č. 1 ─────────────────────────────────────────────────────────────

def _priloha_1(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 1,
                    "Popis výchozího stavu vč. referenční spotřeby a referenčních nákladů",
                    ss, budova)

    _h1(doc, "A. Obecné informace o objektu")
    _celk_plocha = budova.celkova_plocha_m2
    _budova0 = budova.budovy[0] if budova.budovy else None
    _tbl(doc,
         ["Parametr", "Hodnota"],
         [
             ["Název objektu", budova.objekt_nazev or ""],
             ["Adresa", budova.objekt_adresa or ""],
             ["Zadavatel / vlastník", ss.get("zadavatel_nazev", "")],
             ["Druh činnosti / účel budovy", budova.druh_cinnosti or ""],
             ["Celková podlahová plocha", f"{_celk_plocha:,.0f} m²".replace(",", "\u00a0") if _celk_plocha else ""],
             ["Objem budovy", f"{sum(b.objem_m3 for b in budova.budovy):,.0f} m³".replace(",", "\u00a0") if budova.budovy else ""],
             ["Rok výstavby / rekonstrukce", ss.get("rok_vystavby", "")],
             ["Počet uživatelů / zaměstnanců", budova.pocet_zamestnancu or ""],
             ["Provozní doba", budova.provozni_rezim or ss.get("provozni_rezim", "")],
         ],
         widths=[7.0, 9.0])
    doc.add_paragraph()

    _h1(doc, "B. Stávající technický stav")

    _h2(doc, "B.1  Tepelné hospodářství – zdroj tepla a otopná soustava")
    _ph(doc, "Popis stávajícího zdroje tepla: typ, výkon, rok instalace, palivo")
    _ph(doc, "Popis otopné soustavy: teplotní spád, počet těles, stav regulace")

    _h2(doc, "B.2  Příprava teplé vody")
    _ph(doc, "Popis způsobu přípravy TV: centrální/decentrální, zásobník, cirkulace")

    _h2(doc, "B.3  Osvětlení")
    _ph(doc, "Popis stávajícího osvětlení: typ zdrojů, příkon, způsob řízení")

    _h2(doc, "B.4  Vzduchotechnika a větrání")
    _ph(doc, "Popis VZT: nucené/přirozené větrání, výměna vzduchu, ZZT")

    _h2(doc, "B.5  Řízení energetických spotřeb")
    _ph(doc, "Popis řídicích systémů (MaR, BMS, IRC) a energetického managementu")

    doc.add_paragraph()
    _h1(doc, "C. Referenční spotřeby a referenční náklady")
    _p(doc, "Referenční spotřeby jsou stanoveny jako průměr spotřeb za referenční období "
            "(zpravidla 3 roky před realizací opatření), normalizované na klimatické podmínky "
            "referenčního roku. Jsou základem pro výpočet dosažených úspor.")

    en = result.energie
    cena_zp  = ss.get("cena_zp",    1800.0)
    cena_t   = ss.get("cena_teplo", 2200.0)
    cena_ee  = ss.get("cena_ee",    4500.0)
    cena_v   = ss.get("cena_voda", 100.0)
    spotr_v  = 0.0  # doplní zhotovitel / zadavatel

    ref_rows = [
        ["Zemní plyn – vytápění",   _fmt(en.zp_ut, 1),      "MWh/rok", _fmt(cena_zp, 0),  _fmt(en.zp_ut * cena_zp / 1000, 0)],
        ["Zemní plyn – TUV",        _fmt(en.zp_tuv, 1),     "MWh/rok", _fmt(cena_zp, 0),  _fmt(en.zp_tuv * cena_zp / 1000, 0)],
        ["Teplo (CZT)",             _fmt(en.teplo_total, 1), "MWh/rok", _fmt(cena_t, 0),   _fmt(en.teplo_total * cena_t / 1000, 0)],
        ["Elektrická energie",      _fmt(en.ee, 1),          "MWh/rok", _fmt(cena_ee, 0),  _fmt(en.ee * cena_ee / 1000, 0)],
        ["Voda",                    _fmt(spotr_v, 0),        "m³/rok",  _fmt(cena_v, 0),   _fmt(spotr_v * cena_v / 1000, 0)],
        ["CELKEM",                  "",                      "",        "",                 _fmt(en.celkove_naklady, 0)],
    ]
    _tbl(doc,
         ["Energonosič", "Spotřeba", "Jednotka", "Cena [Kč/MWh resp. m³]", "Roční náklady [Kč]"],
         ref_rows,
         widths=[4.5, 2.5, 1.8, 3.2, 4.0],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                 WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT])

    _h2(doc, "C.1  Referenční období a klimatické podmínky")
    lokalita = ss.get("lokalita_projekt", "Praha (Karlov)")
    theta_i  = ss.get("theta_i", 21.0)
    theta_e  = ss.get("theta_e", -13.0)
    _tbl(doc,
         ["Parametr", "Hodnota", "Jednotka"],
         [
             ["Referenční lokalita", lokalita, ""],
             ["Vnitřní výpočtová teplota", f"{theta_i:.1f}", "°C"],
             ["Venkovní výpočtová teplota", f"{theta_e:.1f}", "°C"],
             ["Referenční rok", ss.get("datum", "")[:4] or "dle smlouvy", ""],
             ["Normalizace klimatu", "Denostupňová metoda dle ČSN EN ISO 15316-4-1", ""],
         ],
         widths=[6.0, 6.0, 4.0])

    _h2(doc, "C.2  Referenční ceny energií")
    _p(doc, "Referenční ceny energií jsou platné pro celou dobu trvání smlouvy "
            "a jsou základem pro výpočet dosažených úspor v Kč.")
    _tbl(doc,
         ["Energonosič", "Referenční cena", "Jednotka"],
         [
             ["Zemní plyn",           _fmt(cena_zp, 0),  "Kč/MWh"],
             ["Teplo (CZT)",          _fmt(cena_t, 0),   "Kč/MWh"],
             ["Elektrická energie",   _fmt(cena_ee, 0),  "Kč/MWh"],
             ["Voda",                 _fmt(cena_v, 0),   "Kč/m³"],
         ],
         widths=[6.0, 4.0, 6.0],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT])

    return _save(doc)


# ── Příloha č. 2 ─────────────────────────────────────────────────────────────

def _priloha_2(budova: BuildingInfo, result: ProjectResult,
               aktivni_op: list[str], ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 2, "Popis základních opatření", ss, budova)

    _h1(doc, "Technický popis opatření")
    _p(doc, "Níže jsou popsána všechna základní energeticky úsporná opatření, která budou "
            "v rámci tohoto EPC projektu realizována. Zhotovitel je povinen zajistit dosažení "
            "minimálních technických parametrů uvedených v zadávací dokumentaci.")
    doc.add_paragraph()

    for i, op_id in enumerate(aktivni_op, start=1):
        info = OP_INFO.get(op_id, {})
        mr = next((m for m in result.vysledky if m.id == op_id), None)

        _h2(doc, f"Opatření č. {i}: {info.get('title', op_id)}")

        # Popis z OP_INFO
        popis = info.get("popis", "")
        if popis:
            _p(doc, popis)

        # Ekonomické parametry opatření
        if mr:
            uspora_kc = mr.uspora_kc
            inv = mr.investice
            doc.add_paragraph()
            _tbl(doc,
                 ["Parametr", "Hodnota", "Jednotka"],
                 [
                     ["Investiční náklady",            _fmt(inv, 0),      "Kč bez DPH"],
                     ["Roční úspora nákladů na energii", _fmt(uspora_kc, 0), "Kč/rok"],
                     ["Prostá návratnost",              _fmt(inv / uspora_kc, 1) if uspora_kc > 0 else "—", "let"],
                 ],
                 widths=[7.0, 4.0, 5.0],
                 aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT])

        _ph(doc, f"Technické parametry navrhovaného řešení pro {op_id} – upřesnit dle projektové dokumentace")
        doc.add_paragraph()

    return _save(doc)


# ── Příloha č. 3 ─────────────────────────────────────────────────────────────

def _priloha_3(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 3, "Cena a její úhrada", ss, budova)

    inv = result.celkova_investice
    dph = inv * 0.21
    inv_s_dph = inv + dph
    horizont   = int(ss.get("horizont", 10))

    _h1(doc, "Cena za realizaci úsporných opatření")
    _p(doc, f"Celková cena za úsporná opatření činí "
            f"{_fmt(inv, 0)} Kč (bez DPH), tj. {_fmt(inv_s_dph, 0)} Kč včetně DPH 21 %.")
    _p(doc, "Platba Klienta za úsporná opatření proběhne jednorázově, a to do 30 dnů "
            "od dokončení realizace úsporných opatření a podpisu předávacího protokolu.")

    doc.add_paragraph()
    _tbl(doc,
         ["Opatření", "Investice bez DPH [Kč]", "DPH 21 % [Kč]", "Investice vč. DPH [Kč]"],
         [
             [f"{mr.id} – {OP_INFO.get(mr.id, {}).get('title', mr.id)}",
              _fmt(mr.investice, 0),
              _fmt(mr.investice * 0.21, 0),
              _fmt(mr.investice * 1.21, 0)]
             for mr in result.vysledky if mr.aktivni
         ] + [
             ["CELKEM",
              _fmt(inv, 0),
              _fmt(dph, 0),
              _fmt(inv_s_dph, 0)],
         ],
         widths=[7.5, 3.5, 2.5, 3.0],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                 WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT])

    _h1(doc, "Cena za energetický management")
    _p(doc, f"Roční cena za energetický management bude součástí nabídky zhotovitele. "
            f"Cena bude účtována ročně vždy po skončení příslušného zúčtovacího období "
            f"po dobu {horizont} let.")
    _ph(doc, "Roční cena za energetický management [Kč/rok bez DPH]")
    _ph(doc, f"Kumulovaná cena za EnM za {horizont} let [Kč bez DPH]")

    doc.add_paragraph()
    _h1(doc, "Celková cena projektu")
    _tbl(doc,
         ["Položka", "Cena bez DPH [Kč]"],
         [
             ["Realizace úsporných opatření", _fmt(inv, 0)],
             [f"Energetický management ({horizont} let)", _DZ],
             ["CELKEM", _DZ],
             ["DPH 21 %", _DZ],
             ["CELKEM vč. DPH", _DZ],
         ],
         widths=[9.0, 7.0],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT])

    doc.add_paragraph()
    _h1(doc, "Fakturační a korespondenční údaje")
    _ph(doc, "Fakturační adresa Klienta")
    _ph(doc, "Korespondenční adresa Klienta")
    _ph(doc, "Elektronická adresa pro zasílání faktur (datová schránka / e-mail)")

    return _save(doc)


# ── Příloha č. 4 ─────────────────────────────────────────────────────────────

def _priloha_4(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 4, "Harmonogram realizace projektu", ss, budova)

    horizont = int(ss.get("horizont", 10))

    _h1(doc, "Harmonogram základních fází projektu")
    _p(doc, "Níže je uveden hrubý harmonogram realizace EPC projektu. Podrobný harmonogram "
            "bude součástí projektové dokumentace.")

    doc.add_paragraph()
    _tbl(doc,
         ["Fáze", "Popis", "Předpokládaný termín"],
         [
             ["I.",   "Předběžné činnosti – ověření stávajícího stavu, upřesnění projektové dokumentace",
              "doplnit"],
             ["II.",  "Projektová dokumentace – zpracování PD pro stavební povolení a realizaci",
              "doplnit"],
             ["III.", "Realizace základních úsporných opatření – stavební a technologické práce, "
                      "uvedení do provozu, komplexní zkoušky",
              "doplnit"],
             ["IV.",  f"Poskytování garance – M&V, průběžné zprávy, energetický management "
                      f"po dobu {horizont} let",
              "doplnit"],
         ],
         widths=[1.0, 10.0, 5.0])

    doc.add_paragraph()
    _h2(doc, "Podmínky harmonogramu")
    _p(doc, "Harmonogram je závazný za předpokladu, že Klient zajistí potřebnou součinnost "
            "(přístup do objektu, vydání stavebního povolení, předání podkladů) v dohodnutých termínech. "
            "Případné zpoždění na straně Klienta opravňuje ESCO k odpovídajícímu posunutí termínů.")
    _p(doc, "Zúčtovací období garance začíná běžet dnem podpisu předávacího protokolu "
            "základních úsporných opatření.")

    return _save(doc)


# ── Příloha č. 5 ─────────────────────────────────────────────────────────────

def _priloha_5(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 5, "Výše garantované úspory", ss, budova)

    horizont  = int(ss.get("horizont", 10))

    uspora_kc  = result.celkova_uspora_kc
    uspora_zp  = result.celkova_uspora_zp
    uspora_t   = result.celkova_uspora_teplo
    uspora_ee  = result.celkova_uspora_ee
    uspora_mwh = uspora_zp + uspora_t + uspora_ee
    ref_naklady = result.energie.celkove_naklady
    uspora_pct  = uspora_kc / ref_naklady * 100 if ref_naklady > 0 else 0.0

    _h1(doc, "Výše garantované úspory")
    _p(doc, "Výše úspory je smluvně garantovaná a bude součástí nabídky zhotovitele. "
            "Níže jsou uvedeny vypočtené (kalkulované) úspory jako podklad pro nabídku. "
            "Zhotovitel nabídne výši garantované úspory v procentech kalkulované úspory.")
    doc.add_paragraph()

    _tbl(doc,
         ["Ukazatel", "Kalkulovaná hodnota", "Jednotka"],
         [
             ["Roční úspora nákladů – kalkulovaná",  _fmt(uspora_kc, 0),   "Kč/rok"],
             ["Roční úspora energie – celkem",        _fmt(uspora_mwh, 1),  "MWh/rok"],
             ["  z toho zemní plyn",                  _fmt(uspora_zp, 1),   "MWh/rok"],
             ["  z toho teplo (CZT)",                 _fmt(uspora_t, 1),    "MWh/rok"],
             ["  z toho elektrická energie",          _fmt(uspora_ee, 1),   "MWh/rok"],
             ["Relativní úspora nákladů",             f"{uspora_pct:.1f}",  "%"],
             [f"Kumulovaná úspora za {horizont} let", _fmt(uspora_kc * horizont, 0), "Kč"],
         ],
         widths=[8.0, 4.5, 4.0],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()
    _p(doc, "Nabídka zhotovitele musí obsahovat:", bold=True)
    _ph(doc, "Výše garantované roční úspory [Kč/rok]")
    _ph(doc, "Výše garantované roční úspory [% z kalkulované hodnoty]")
    _ph(doc, f"Kumulovaná garantovaná úspora za {horizont} let [Kč]")

    doc.add_paragraph()
    _h1(doc, "Sankce za nedosažení garantované úspory")
    _p(doc, "V případě, že z důvodů výlučně na straně ESCO nebude v zúčtovacím období "
            "dosaženo garantované výše úspory, ESCO uhradí Klientovi záporný rozdíl mezi "
            "garantovanou úsporou a skutečně dosaženou úsporou v plné výši.")
    _p(doc, "Sankce = garantovaná úspora − skutečně dosažená úspora  [Kč/zúčtovací období]")

    doc.add_paragraph()
    _h1(doc, "Prémie za překročení garantované úspory")
    _p(doc, "Bude-li v zúčtovacím období dosaženo vyšší než garantované úspory, ESCO má nárok "
            "na prémii za překročení ve výši:")
    _p(doc, "Prémie = 0,05 × (dosažená úspora − garantovaná úspora)  [Kč/zúčtovací období]",
       bold=True)

    return _save(doc)


# ── Příloha č. 6 ─────────────────────────────────────────────────────────────

def _priloha_6(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 6, "Vyhodnocování dosažených úspor (M&V)", ss, budova)

    horizont = int(ss.get("horizont", 10))
    lokalita = ss.get("lokalita_projekt", "Praha (Karlov)")
    theta_i  = ss.get("theta_i", 21.0)

    _h1(doc, "Druh úspor a metodika vyhodnocení")
    _p(doc, "Vyhodnocované úspory jsou definovány jako nerealizovaná spotřeba energie "
            "(nerealizované náklady) – tj. rozdíl mezi tím, jaká by byla spotřeba bez realizace "
            "opatření (referenční stav upravený na aktuální podmínky) a skutečnou spotřebou "
            "v zúčtovacím období.")

    _h2(doc, "Vybraná varianta IPMVP a hranice systému")
    _p(doc, "ESCO volí variantu C dle IPMVP – měření na úrovni celého objektu. "
            "Využívají se data z fakturačních měřidel nebo z podružných měřidel. "
            "Veškeré úspory uvažované ve smlouvě jsou měřitelné.")

    doc.add_paragraph()
    _h1(doc, "Zúčtovací období")
    _ph(doc, "Datum zahájení garance (podpis předávacího protokolu)")
    _p(doc, f"Délka zúčtovacího období: 1 rok (12 měsíců).")
    _p(doc, f"Celková doba garance: {horizont} let.")

    doc.add_paragraph()
    _h1(doc, "Výpočet úspory – metodika")
    _h2(doc, "Zemní plyn a teplo (CZT) – denostupňová normalizace")
    _p(doc, "Referenční spotřeba paliv na vytápění je normalizována na klimatické podmínky "
            "zúčtovacího období pomocí denostupňové metody:")
    _p(doc, "Q_ref_upravená = Q_ref_nezávislá + Q_ref_závislá × (DS_skutečné / DS_referenční)",
       bold=True)
    _p(doc, f"Kde denostupně jsou počítány jako: DS = Σ (θ_i − θ_e,průměrná) × počet topných dnů, "
            f"referenční vnitřní teplota θ_i = {theta_i:.0f} °C, lokalita: {lokalita}.")

    _h2(doc, "Elektrická energie – výpočtová metoda")
    _p(doc, "Úspora elektrické energie je stanovena jako rozdíl příkonu stávajících "
            "a nových zařízení vynásobený dobou provozu dle projektu.")

    _h2(doc, "Voda")
    _p(doc, "Úspora vody je stanovena jako rozdíl spotřeby v referenčním a zúčtovacím období "
            "dle podružných vodoměrů.")

    doc.add_paragraph()
    _h1(doc, "Ceny energie pro výpočet úspory nákladů")
    _p(doc, "Úspory nákladů se stanovují použitím referenčních cen energií "
            "uvedených v Příloze č. 1 dle vzorce:")
    _p(doc, "Úspory nákladů = C_baseline − C_reporting  [Kč/zúčtovací období]", bold=True)

    doc.add_paragraph()
    _h1(doc, "Přesnost dat a zdroje dat")
    _tbl(doc,
         ["Energonosič / médium", "Zdroj dat", "Přesnost"],
         [
             ["Zemní plyn",         "Fakturační plynoměr (dodavatel energie)",        "±0,5 % (dle zákona o metrologii)"],
             ["Teplo (CZT)",        "Fakturační kalorimetr (dodavatel tepla)",         "±1 % (dle ČSN EN 1434)"],
             ["Elektrická energie", "Fakturační elektroměr / podružné elektroměry",   "±1 % (dle zákona o metrologii)"],
             ["Voda",               "Vodoměr / podružné vodoměry",                    "±2 %"],
             ["Klimatická data",    f"ČHMÚ, lokalita {lokalita}",                     "Průměrné měsíční teploty a topné dny"],
         ],
         widths=[4.0, 6.5, 6.0])

    doc.add_paragraph()
    _h1(doc, "Odpovědnosti a postup vyhodnocení")
    _tbl(doc,
         ["Odpovědnost", "ESCO", "Klient"],
         [
             ["Sběr dat z měřidel",                 "✓", "součinnost"],
             ["Měsíční průběžná zpráva",             "✓", "—"],
             ["Roční zpráva o dosažených úsporách",  "✓ (do 60 dnů po ukončení ZO)", "schválení"],
             ["Kalibrace a údržba měřidel",          "✓", "—"],
             ["Informování o změnách v objektu",     "—", "✓"],
             ["Nápravná opatření při nedosažení",    "✓ (návrh do 30 dnů)", "schválení"],
         ],
         widths=[7.0, 4.5, 5.0])

    doc.add_paragraph()
    _h1(doc, "Průběžná zpráva o vyhodnocení úspor")
    _p(doc, "ESCO bude ročně předkládat průběžnou zprávu hodnotící uplynulé zúčtovací "
            "období (do 60 dnů po ukončení ZO). Zpráva bude obsahovat minimálně:")
    for bod in [
        "Popis provozu energetického systému v zúčtovacím období",
        "Surová měřená data za zúčtovací období",
        "Výpočet dosažené úspory v MWh a Kč za jednotlivé energonosiče",
        "Porovnání dosažené a garantované úspory",
        "Závěr o splnění / nesplnění garance, příp. výpočet sankce nebo prémie",
        "Jméno zpracovatele zprávy, datum, podpis oprávněné osoby ESCO",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    return _save(doc)


# ── Příloha č. 7 ─────────────────────────────────────────────────────────────

def _priloha_7(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 7, "Energetický management", ss, budova)

    horizont = int(ss.get("horizont", 10))

    _h1(doc, "Úvod")
    _p(doc, "Energetický management (EnM) je nedílnou součástí služeb poskytovaných ESCO v rámci "
            "této smlouvy. Je nezbytný pro dosažení garantované úspory, pro její průkazné měření "
            "a vyhodnocení. EnM není možné vykonávat bez náležité součinnosti Klienta.")

    doc.add_paragraph()
    _h1(doc, "Povinnosti ESCO")
    for bod in [
        "Pravidelné měsíční vyhodnocování spotřeb energií a výpočet úspor",
        "Předkládání průběžných měsíčních zpráv se spotřebami a výpočtem úspory",
        "Zpracování roční zprávy o dosažených úsporách (do 60 dnů po ukončení ZO)",
        "Sledování a vizualizace provozu klíčových technologií prostřednictvím dispečinku",
        "Návrh nápravných opatření v případě nedosažení plánovaných úspor",
        "Údržba, kalibrace a servis instalovaných měřicích zařízení",
        "Aktivní hledání dalších potenciálu úspor a jejich předkládání Klientovi",
        "Školení obsluhy a správy objektu v oblasti energeticky úsporného provozu",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    doc.add_paragraph()
    _h1(doc, "Povinnosti Klienta")
    for bod in [
        "Informovat ESCO o plánovaných změnách v provozu objektu, jež mohou ovlivnit spotřebu energie",
        "Zajistit přístup ESCO k měřidlům a technologiím pro účely M&V a servisu",
        "Předávat fakturační data za energie do 15 dnů po jejich obdržení",
        "Schvalovat průběžné a roční zprávy ESCO do 30 dnů od jejich předložení",
        "Zachovat realizovaná opatření v původním technickém stavu; "
            "případné změny konzultovat s ESCO",
        "Provozovat objekt způsobem, který nebrání dosažení garantovaných úspor",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    doc.add_paragraph()
    _h1(doc, "Nástroje a systémy EnM")
    _ph(doc, "Popis instalovaných měřicích zařízení a řídicích systémů (MaR, BMS, dispečink)")
    _ph(doc, "Popis softwaru pro sběr, archivaci a vizualizaci dat o spotřebách")

    doc.add_paragraph()
    _h1(doc, "Cena za energetický management")
    _ph(doc, f"Roční cena za energetický management [Kč/rok bez DPH]")
    _ph(doc, f"Kumulovaná cena za EnM za {horizont} let [Kč bez DPH]")

    return _save(doc)


# ── Příloha č. 8 ─────────────────────────────────────────────────────────────

def _priloha_8(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 8, "Oprávněné osoby", ss, budova)

    _h1(doc, "Oprávněné osoby ESCO (zhotovitele)")
    _p(doc, "Kontaktní osoby oprávněné jednat ve věcech technických a smluvních za ESCO:")
    _tbl(doc,
         ["Funkce", "Jméno a příjmení", "E-mail", "Telefon"],
         [
             ["Projektový manažer / kontakt pro VŘ", _DZ, _DZ, _DZ],
             ["Vedoucí realizace",                   _DZ, _DZ, _DZ],
             ["Garant / EnM manažer",                _DZ, _DZ, _DZ],
             ["Technický dozor",                     _DZ, _DZ, _DZ],
         ],
         widths=[4.5, 4.0, 4.0, 4.0])

    doc.add_paragraph()
    _h1(doc, "Oprávněné osoby Klienta (objednatele)")
    _p(doc, "Kontaktní osoby oprávněné jednat ve věcech technických a smluvních za Klienta:")
    _tbl(doc,
         ["Funkce", "Jméno a příjmení", "E-mail", "Telefon"],
         [
             ["Kontaktní osoba pro VŘ",  ss.get("zadavatel_kontakt", ""),
              ss.get("zadavatel_email", ""), ss.get("zadavatel_tel", "")],
             ["Osoba odpovědná za objekt", "", "", ""],
             ["Energetický manažer",       "", "", ""],
         ],
         widths=[4.5, 4.0, 4.0, 4.0])

    doc.add_paragraph()
    _p(doc, "Oprávněné osoby mohou být změněny písemným oznámením druhé smluvní straně "
            "bez nutnosti uzavření dodatku ke smlouvě.")

    return _save(doc)


# ── Příloha č. 9 ─────────────────────────────────────────────────────────────

def _priloha_9(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 9, "Seznam poddodavatelů", ss, budova)

    _h1(doc, "Poddodavatelé s podílem 10 % a vyšším na celkové hodnotě zakázky")
    _p(doc, "V souladu s podmínkami zadávacího řízení jsou níže uvedeni poddodavatelé, "
            "jejichž podíl na celkové hodnotě zakázky činí 10 % nebo více.")
    doc.add_paragraph()
    _tbl(doc,
         ["Obchodní firma", "IČO", "Předmět plnění", "Podíl na zakázce [%]"],
         [
             ["", "", "", ""],
             ["", "", "", ""],
             ["", "", "", ""],
         ],
         widths=[5.0, 2.5, 6.5, 2.5],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                 WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT])

    doc.add_paragraph()
    _p(doc, "Poznámka: Pokud ESCO nema poddodavatele s podilem >= 10 %, tuto skutecnost "
            "explicitne uvede a tabulku oznaci jako 'Bez poddodavatelu'.")

    return _save(doc)


# ── Terezín variant – P04 ─────────────────────────────────────────────────────

def _t_priloha_4(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 4, "Harmonogram plnění smlouvy", ss, budova)

    _h1(doc, "Podrobný harmonogram")
    _p(doc, "Klient vyžaduje, aby ESCO předložilo podrobný harmonogram realizace projektu "
            "do 30 dnů od podpisu smlouvy. Harmonogram musí obsahovat:")
    for bod in [
        "Měsíční časový a finanční harmonogram",
        "Výrobní strukturu (technologické etapy se začátky, konci a dobami trvání)",
        "Označení kritické cesty",
        "Rezervy pro činnosti mimo kritickou cestu",
        "Harmonogram funkčních zkoušek a přejímacích řízení",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    doc.add_paragraph()
    _h1(doc, "Závazné milníky realizace")
    _p(doc, "ESCO je povinno dodržet následující závazné termíny:")
    _tbl(doc,
         ["Milník", "Závazný termín"],
         [
             ["Předložení dokumentace ke kontrole (střecha)", _DZ],
             ["Podání žádosti o stavební povolení",           _DZ],
             ["Realizace rekonstrukce střechy",               _DZ],
             ["Realizace výměny zdroje tepla",                _DZ],
             ["Realizace všech základních opatření",          _DZ],
             ["Energetický management – zahájení",            _DZ],
         ],
         widths=[10.0, 6.0])

    doc.add_paragraph()
    _h1(doc, "Podmínky harmonogramu")
    _p(doc, "Harmonogram je závazný za předpokladu, že Klient zajistí potřebnou součinnost "
            "(přístup do objektu, vydání stavebního povolení, předání podkladů) v dohodnutých termínech. "
            "Případné zpoždění na straně Klienta opravňuje ESCO k odpovídajícímu posunutí termínů.")
    _p(doc, "Zúčtovací období garance začíná běžet dnem podpisu předávacího protokolu "
            "základních úsporných opatření.")

    return _save(doc)


# ── Terezín variant – P05 ─────────────────────────────────────────────────────

def _t_priloha_5(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 5, "Výše garantované úspory", ss, budova)

    en = result.energie
    cena_zp  = ss.get("cena_zp",    1800.0)
    cena_t   = ss.get("cena_teplo", 2200.0)
    cena_ee  = ss.get("cena_ee",    4500.0)

    uspora_kc  = result.celkova_uspora_kc
    uspora_zp  = result.celkova_uspora_zp
    uspora_t   = result.celkova_uspora_teplo
    uspora_ee  = result.celkova_uspora_ee
    uspora_mwh = uspora_zp + uspora_t + uspora_ee
    horizont   = int(ss.get("horizont", 10))

    _h1(doc, "Vyčíslení úspory")
    _p(doc, "Tabulka uvádí kalkulované úspory po jednotlivých účelech spotřeby energie. "
            "Hodnoty jsou podkladem pro nabídku zhotovitele, který doplní výši garantované úspory. "
            "Pro hodnocení nabídek je rozhodná celková úspora primární energie z neobnovitelných "
            "zdrojů – nabídky s hodnotou nižší než 40 % referenční spotřeby budou vyloučeny.")
    doc.add_paragraph()

    _tbl(doc,
         ["Účel spotřeby energie",
          "Úspora EE\n[MWh]",
          "Úspora ZP\n[MWh]",
          "Úspora dodané energie\n[MWh]",
          "Úspora\nvč. DPH [Kč]",
          "Úspora primární energie\n(neobnovit.) [MWh]"],
         [
             ["Vytápění",             _DZ, _fmt(uspora_zp + uspora_t, 1), _DZ, _DZ, _DZ],
             ["Větrání",              _DZ, _DZ, _DZ, _DZ, _DZ],
             ["Příprava teplé vody",  _DZ, _DZ, _DZ, _DZ, _DZ],
             ["Osvětlení",            _fmt(uspora_ee, 1), "0,0", _fmt(uspora_ee, 1), _DZ, _DZ],
             ["Využití EE z FVE",     _DZ, _DZ, _DZ, _DZ, _DZ],
             ["CELKEM",               _fmt(uspora_ee, 1), _fmt(uspora_zp + uspora_t, 1),
              _fmt(uspora_mwh, 1), _fmt(uspora_kc, 0), _DZ],
         ],
         widths=[3.8, 1.9, 1.9, 2.8, 2.6, 3.5],
         aligns=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 5)

    doc.add_paragraph()
    _p(doc, "Nabídka zhotovitele musí obsahovat:", bold=True)
    _ph(doc, "Výše garantované roční usp. energie z neobnovitelnych zdrojů [MWh/rok]")
    _ph(doc, "Výše garantované roční úspory nákladů [Kč/rok vč. DPH]")
    _ph(doc, f"Kumulovaná garantovaná úspora za {horizont} let [Kč vč. DPH]")

    doc.add_paragraph()
    _h1(doc, "Sankce za nedosažení garantované úspory")
    _p(doc, "V případě, že z důvodů výlučně na straně ESCO nebude v zúčtovacím období "
            "dosaženo garantované výše úspory, ESCO uhradí Klientovi záporný rozdíl mezi "
            "garantovanou úsporou a skutečně dosaženou úsporou v plné výši.")

    return _save(doc)


# ── Terezín variant – P06 ─────────────────────────────────────────────────────

def _t_priloha_6(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 6, "Metodika vyhodnocování úspor", ss, budova)

    en = result.energie
    has_zp    = en.zp_ut + en.zp_tuv > 0
    has_teplo = en.teplo_total > 0
    has_ee    = en.ee > 0
    ano_ne    = lambda b: "ano" if b else "ne"
    lokalita  = ss.get("lokalita_projekt", "Praha (Karlov)")
    theta_i   = ss.get("theta_i", 21.0)
    horizont  = int(ss.get("horizont", 10))

    _h1(doc, "Druh vyhodnocovaných úspor")
    _p(doc, "V projektu jsou vyhodnocovány následující druhy úspor. "
            "Tabulka uvádí, pro které energonosiče jsou úspory v daném objektu vyhodnocovány.")
    doc.add_paragraph()

    _tbl(doc,
         ["Č.", "Název objektu", "Adresa",
          "plyn\nÚSP_P", "teplo\nÚSP_T", "elektřina\nÚSP_E", "voda\nÚSP_V", "ostatní\nÚSP_O"],
         [
             ["1",
              budova.objekt_nazev or "",
              budova.objekt_adresa or "",
              ano_ne(has_zp),
              ano_ne(has_teplo),
              ano_ne(has_ee),
              "ano",
              "ne"],
         ],
         widths=[0.8, 3.5, 4.0, 1.3, 1.3, 1.7, 1.3, 1.6],
         aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
                + [WD_ALIGN_PARAGRAPH.CENTER] * 5)

    doc.add_paragraph()
    _h1(doc, "Definice druhů vyhodnocovaných úspor")
    _p(doc, "Vyhodnocované úspory jsou definovány v souladu s IPMVP sv. I (EVO 10000-1:2009):")
    doc.add_paragraph()

    _tbl(doc,
         ["Č.", "Název objektu", "Adresa",
          "plyn", "teplo", "elektřina", "voda", "ostatní"],
         [
             ["1",
              budova.objekt_nazev or "",
              budova.objekt_adresa or "",
              "nerealizovaná spotřeba / normalizované náklady" if has_zp else "—",
              "nerealizovaná spotřeba / normalizované náklady" if has_teplo else "—",
              "nerealizovaná spotřeba / normalizované náklady" if has_ee else "—",
              "nerealizovaná spotřeba",
              "—"],
         ],
         widths=[0.8, 3.5, 3.5, 1.5, 1.5, 1.9, 1.4, 1.4],
         aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
                + [WD_ALIGN_PARAGRAPH.CENTER] * 5)

    doc.add_paragraph()
    _h1(doc, "Metodika výpočtu – normalizace klimatu")
    _p(doc, "Referenční spotřeba paliv na vytápění a TUV je normalizována denostupňovou metodou:")
    _p(doc, "Q_ref_upravená = Q_ref_nezávislá + Q_ref_závislá × (DS_skutečné / DS_referenční)",
       bold=True)
    _p(doc, f"Vnitřní výpočtová teplota θ_i = {theta_i:.0f} °C, lokalita: {lokalita}. "
            f"Klimatická data: ČHMÚ.")

    doc.add_paragraph()
    _h1(doc, "Zúčtovací období")
    _ph(doc, "Datum zahájení garance (podpis předávacího protokolu)")
    _p(doc, f"Délka zúčtovacího období: 1 rok. Celková doba garance: {horizont} let.")

    return _save(doc)


# ── Terezín variant – P07 ─────────────────────────────────────────────────────

def _t_priloha_7(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 7, "Energetický management", ss, budova)

    horizont = int(ss.get("horizont", 10))

    _h1(doc, "Energetický management – činnosti a povinnosti ESCO")
    for bod in [
        "Měsíční sledování a archivace spotřeby zemního plynu na vytápění a TUV dle objektů/odběrných míst",
        "Výpočet spotřeby plynu v měsíčním intervalu",
        "Porovnání s referenčním obdobím nebo předchozím cyklem",
        "Výpočet dosažených úspor s teplotní korekcí a zohledněním provozních změn",
        "Prognózování a hodnocení očekávané vs. skutečné spotřeby",
        "Analýza příčin zvýšené spotřeby",
        "Návrh řešení při nadměrné spotřebě",
        "Odborné poradenství a spolupráce s Klientem",
        "Dálkové poradenství a intervence pro optimalizaci MaR systémů (dispečink)",
        "Sledování správné funkce realizovaných opatření",
        "Optimalizace provozu, zejména při nadměrné spotřebě",
        "Identifikace a doporučení dalších úsporných opatření",
        "Zpracování roční zprávy o dosažených úsporách (do 60 dnů po ukončení ZO)",
        "Předkládání průběžných měsíčních zpráv se spotřebami a výpočtem úspory",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    doc.add_paragraph()
    _h1(doc, "Energetický management – činnosti a povinnosti klienta")
    _p(doc, "Klient je povinen zajistit potřebnou součinnost, zejména:")
    for bod in [
        "Sledovat stavy měřidel a předávat informace ESCO v dohodnutých termínech",
        "Předávat měsíční faktury za spotřebu plynu (do 5 pracovních dnů od obdržení)",
        "Předávat měsíční stavy plynoměrů (do 7. dne v měsíci)",
        "Předávat faktury za plyn, elektřinu a vodu (do 5 pracovních dnů od obdržení)",
        "Informovat ESCO o plánovaných změnách: významné změny (30 dnů předem e-mailem + poštou), "
            "drobné změny (7 dnů předem e-mailem), mimořádné skutečnosti (bezodkladně e-mailem)",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    doc.add_paragraph()
    _h1(doc, "Standardní provozní podmínky")
    _p(doc, "Nastavení HVAC systémů dle vyhlášky č. 194/2007 Sb. Provozní a útlumové režimy "
            "jsou nastaveny dohodou poskytovatele a Klienta a průběžně aktualizovány "
            "podle aktuálního využití objektu.")
    _ph(doc, "Teplotní nastavení a útlumové časy – doplní zhotovitel dle projektu")

    doc.add_paragraph()
    _h1(doc, "Nástroje a systémy EnM")
    _ph(doc, "Popis instalovaných měřicích zařízení a řídicích systémů (MaR, BMS, dispečink)")
    _ph(doc, "Popis softwaru pro sběr, archivaci a vizualizaci dat o spotřebách")

    doc.add_paragraph()
    _h1(doc, "Cena za energetický management")
    _ph(doc, f"Roční cena za energetický management [Kč/rok bez DPH]")
    _ph(doc, f"Kumulovaná cena za EnM za {horizont} let [Kč bez DPH]")

    return _save(doc)


# ── Terezín variant – P08 ─────────────────────────────────────────────────────

def _t_priloha_8(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 8, "Oprávněné osoby", ss, budova)

    _h1(doc, "Oprávněné osoby ESCO")
    for oblast in [
        "pro obchodní a smluvní záležitosti",
        "pro technické a provozní záležitosti",
        "pro fakturační záležitosti",
    ]:
        _p(doc, f"Oprávněná osoba {oblast}:")
        _ph(doc, "doplní účastník")
        doc.add_paragraph()

    _h1(doc, "Oprávněné osoby Klienta")
    for oblast in [
        "pro obchodní a smluvní záležitosti",
        "pro technické a provozní záležitosti",
        "pro fakturační záležitosti",
    ]:
        _p(doc, f"Oprávněná osoba {oblast}:")
        _ph(doc, "doplní zadavatel")
        doc.add_paragraph()

    _p(doc, "Oprávněné osoby mohou být změněny písemným oznámením druhé smluvní straně "
            "bez nutnosti uzavření dodatku ke smlouvě.")

    return _save(doc)


# ── Terezín variant – P09 ─────────────────────────────────────────────────────

def _t_priloha_9(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 9, "Seznam poddodavatelů", ss, budova)
    _ph(doc, "doplní zhotovitel")
    return _save(doc)


# ── Terezín variant – P10 (Inflační doložka) ──────────────────────────────────

def _t_priloha_10(budova: BuildingInfo, result: ProjectResult, ss: dict) -> BytesIO:
    doc = _doc()
    _titulni_strana(doc, 10, "Inflační doložka", ss, budova)

    _h1(doc, "Účel a rozsah")
    _p(doc, "Cena základních opatření bude čtvrtletně upravována v závislosti na změnách "
            "materiálových, mzdových a dalších vstupních nákladů. Doložka se vztahuje na veškeré "
            "položky a práce provedené ESCO v průběhu realizace základních opatření, "
            "a to od podání závazné cenové nabídky v zadávacím řízení až po skutečnou realizaci "
            "každé položky / práce.")

    doc.add_paragraph()
    _h1(doc, "Podklad pro výpočet")
    _p(doc, "Používá se Index cen stavebních konstrukcí a prací podle klasifikace TSKPstat "
            "(Český statistický úřad). Využívá se specifický index dle kódu TSKPstat "
            "nejbližšímu fakturované základní práci. Uplatňují se čtvrtletní hodnoty "
            "\"předchozí období = 100\".")

    doc.add_paragraph()
    _h1(doc, "Vzorec pro výpočet cenové úpravy")
    _p(doc, "Výše cenové úpravy za kalendářní čtvrtletí n:", bold=True)
    _p(doc, "    UCn  =  Fnz × (Pnz − 1)", bold=True)

    doc.add_paragraph()
    _tbl(doc,
         ["Symbol", "Význam"],
         [
             ["UCn",  "Částka k připočtení / odečtení za kalendářní čtvrtletí n"],
             ["Fnz",  "Součet smluvních cen položek/prací realizovaných ve čtvrtletí n"],
             ["Pnz",  "Korekční násobitel pro čtvrtletí n (viz níže)"],
         ],
         widths=[2.5, 13.5])

    doc.add_paragraph()
    _p(doc, "Korekční násobitel:", bold=True)
    _p(doc, "    Pnz  =  Π  (Li / 100)   pro i = o … n", bold=True)

    _tbl(doc,
         ["Symbol", "Význam"],
         [
             ["Li",  "Index cen pro i-té čtvrtletí (\"předchozí období = 100\")"],
             ["o",   "Čtvrtletí, v němž bylo podáno závazné cenové nabídkové plnění"],
             ["n",   "Čtvrtletí skutečné realizace položky/práce (dle stavebního deníku)"],
         ],
         widths=[2.5, 13.5])

    doc.add_paragraph()
    _h1(doc, "Podmínky uplatnění")
    for bod in [
        "ESCO předloží zvláštní vyúčtování cenové úpravy jako přílohu faktury, "
            "rozepsané po čtvrtletích kalendářního roku.",
        "Splatnost: 30 dnů od doručení faktury.",
        "Klient může do 14 dnů požádat o přepracování, pokud je vyúčtování věcně nesprávné.",
        "Úprava se neuplatní, pokud korekční násobitel Pnz (zaokrouhlený na 4 desetinná místa) "
            "leží v intervalu <0,99; 1,01>.",
        "Datum rozhodné pro výpočet: datum realizace dle stavebního deníku, nikoli datum vystavení faktury.",
    ]:
        p = doc.add_paragraph(bod, style="List Bullet")
        p.style.font.name = S.FONT_BODY

    doc.add_paragraph()
    _h1(doc, "Tabulka procentního rozdělení podle TSKPstat")
    _p(doc, "Předpokládané procentní rozdělení položek dle TSKPstat kódů bude upřesněno "
            "v průběhu ověřování energetického potenciálu a bude přílohou tohoto dokumentu.")
    _ph(doc, "Tabulka TSKPstat kódů a procentního rozdělení – doplní zhotovitel")

    return _save(doc)


# ── Veřejné API ───────────────────────────────────────────────────────────────

_PRILOHY_APES = [
    ("P01_Výchozí_stav.docx",          _priloha_1),
    ("P02_Popis_opatření.docx",         _priloha_2),
    ("P03_Cena_a_úhrada.docx",          _priloha_3),
    ("P04_Harmonogram.docx",            _priloha_4),
    ("P05_Garantovaná_úspora.docx",     _priloha_5),
    ("P06_Vyhodnocování_úspor_MV.docx", _priloha_6),
    ("P07_Energetický_management.docx", _priloha_7),
    ("P08_Oprávněné_osoby.docx",        _priloha_8),
    ("P09_Poddodavatelé.docx",          _priloha_9),
]

_PRILOHY_DPU = [
    ("P01_Výchozí_stav.docx",          _priloha_1),
    ("P02_Popis_opatření.docx",         _priloha_2),
    ("P03_Cena_a_úhrada.docx",          _priloha_3),
    ("P04_Harmonogram.docx",            _t_priloha_4),
    ("P05_Garantovaná_úspora.docx",     _t_priloha_5),
    ("P06_Vyhodnocování_úspor_MV.docx", _t_priloha_6),
    ("P07_Energetický_management.docx", _t_priloha_7),
    ("P08_Oprávněné_osoby.docx",        _t_priloha_8),
    ("P09_Poddodavatelé.docx",          _t_priloha_9),
    ("P10_Inflační_doložka.docx",        _t_priloha_10),
]


def generuj_ses_prilohy(
    budova: BuildingInfo,
    result: ProjectResult,
    aktivni_op: list[str],
    ss: dict,
    variant: str = "kusk",
) -> BytesIO:
    """
    Sestaví přílohy ke Smlouvě o energetických službách (SES) a vrátí ZIP BytesIO.

    Parametry:
        budova     – BuildingInfo ze session state
        result     – ProjectResult z build_project().vypocitej()
        aktivni_op – seznam aktivních OP ID
        ss         – st.session_state jako dict
        variant    – "apes" (9 příloh, APES standard) nebo "dpu" (10 příloh, DPU standard)
    """
    prilohy = _PRILOHY_DPU if variant == "dpu" else _PRILOHY_APES
    buf_zip = BytesIO()
    with zipfile.ZipFile(buf_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, gen_fn in prilohy:
            if gen_fn is _priloha_2:
                buf = gen_fn(budova, result, aktivni_op, ss)
            else:
                buf = gen_fn(budova, result, ss)
            zf.writestr(filename, buf.getvalue())
    buf_zip.seek(0)
    return buf_zip
