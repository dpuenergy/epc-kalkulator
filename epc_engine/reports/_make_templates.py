"""
Jednorázový skript pro vytvoření .docx šablon se styly.

Spuštění:
    python -m epc_engine.reports._make_templates

Výstup: epc_engine/reports/templates/ep_template.docx
        epc_engine/reports/templates/ea_template.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_OUT = Path(__file__).parent / "templates"


def _rgb(hex6: str) -> RGBColor:
    h = hex6.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_style_font(style, name: str, size_pt: int,
                    bold: bool = False, color: str | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color:
        style.font.color.rgb = _rgb(color)


def _add_footer(doc: Document, text: str) -> None:
    """Přidá zápatí s textem vlevo a číslem stránky vpravo."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.clear()

    # Levá část – text
    run_left = para.add_run(text + "   ")
    run_left.font.name = "Calibri"
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = _rgb("888888")

    # Pravá část – číslo stránky (XML element)
    run_page = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page.font.name = "Calibri"
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = _rgb("888888")


def _build_template(typ: str) -> Document:
    """Sestaví prázdný dokument se Named Styles pro daný typ (EP / EA)."""
    doc = Document()

    # Okraje
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # Normal
    normal = doc.styles["Normal"]
    _set_style_font(normal, "Calibri", 10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(14)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    _set_style_font(h1, "Calibri", 16, bold=True, color="1B4F72")
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    _set_style_font(h2, "Calibri", 13, bold=True, color="2E86AB")
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    _set_style_font(h3, "Calibri", 11, bold=True, color="2C3E50")
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(2)

    # List Bullet
    try:
        lb = doc.styles["List Bullet"]
        _set_style_font(lb, "Calibri", 10)
    except KeyError:
        pass

    # Zápatí
    footer_text = f"DPU ENERGY s.r.o.  |  {typ}  |  Důvěrné"
    _add_footer(doc, footer_text)

    # Smazat prázdný odstavec, který Word přidá automaticky
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    return doc


def make_templates() -> None:
    _OUT.mkdir(parents=True, exist_ok=True)

    ep = _build_template("Energetický posudek")
    ep.save(_OUT / "ep_template.docx")
    print(f"Vytvořeno: {_OUT / 'ep_template.docx'}")

    ea = _build_template("Energetický audit")
    ea.save(_OUT / "ea_template.docx")
    print(f"Vytvořeno: {_OUT / 'ea_template.docx'}")


if __name__ == "__main__":
    make_templates()
