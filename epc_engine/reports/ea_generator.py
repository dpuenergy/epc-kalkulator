"""
Generátor Energetického auditu (EA) ve formátu .docx.

Vyhláška č. 140/2021 Sb. – §5 až §10

Struktura dokumentu:
  Titulní strana
  Obsah
  1. Souhrn energetického auditu (§6)
     1.1 Identifikační údaje zadavatele
     1.2 Identifikační údaje energetického specialisty
     1.3 Základní údaje o EA
     1.4 Souhrn příležitostí – evidenční list (příl. 1B)
     1.5 Program realizace příležitostí
  2. Vymezení předmětu energetického auditu (§7)
     2.1 Hranice energetického hospodářství
     2.2 Specifikace předmětu EA
     2.3 Vstupní podklady
  3. Popis stávajícího stavu (§8 odst. 2–3)
     3.1 Bilance energetických vstupů (příl. 3)
     3.2 Analýza užití energie (příl. 4)
     3.3 Ukazatele energetické náročnosti EnPI (příl. 5)
     3.4 Historie spotřeby energie
     3.5 Tepelně technické vlastnosti obálky budovy
     3.6 Technické systémy budovy
     3.7 Systém managementu hospodaření s energií (EnMS)
     3.8 Energetická bilance stávajícího stavu
  4. Vyhodnocení stávajícího stavu
     4.1 Vyhodnocení účinnosti užití energie ve zdrojích energie
     4.2 Principy výpočtu tepelně technického hodnocení
     4.3 Klasifikace obálky budovy
     4.4 Vyhodnocení plnění legislativních požadavků
     4.5 Emisní bilance
  5. Příležitosti ke snížení energetické náročnosti (§9)
     5.1 Přehled navržených příležitostí
     5.2 Popis příležitostí
     5.3 Ekonomická analýza (příl. 7)
     5.4 Ekologické hodnocení (příl. 8)
     5.5 Multikriteriální hodnocení (příl. 9)
     5.6 Možnosti využití dotačních programů
     5.7 Primární energie z neobnovitelných zdrojů
  6. Závěr a doporučení energetického specialisty
  7. Přílohy (§10)

Použití::
    from epc_engine.reports import generuj_ea
    buf = generuj_ea(budova, result)
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from epc_engine.models import BuildingInfo, ProjectResult, EAData, MKHKriterium, Fotografie, PlanEA
from epc_engine.op_descriptions import OP_INFO
from epc_engine.tepelna_technika import TYP_POPISY, vypocitej_uem_z_konstrukci

from . import _styles as S
from ._grafy import (
    graf_spotreba_rocni,
    graf_spotreba_mesicni,
    graf_podily_spotreby,
    graf_cf_kumulativni,
    graf_investice_uspory,
)
from ._tables import (
    tabulka_energeticka_bilance,
    tabulka_opatreni,
    tabulka_ekonomika,
    tabulka_emise,
    tabulka_klasifikace,
    tabulka_budovy,
    tabulka_podklady,
    tabulka_konstrukce,
    tabulka_bilance_pouziti,
    tabulka_penb,
    tabulka_mesicni_spotreba,
    tabulka_u_porovnani,
    tabulka_enms,
    tabulka_bilance_vstupu_ea,
    tabulka_enpi,
    tabulka_mkh,
    tabulka_evidencni_list_1b,
    tabulka_bilance_rocni,
    tabulka_cash_flow,
    tabulka_soucet_variant,
    tabulka_ucinnost_zdroju,
)

_TEMPLATE = Path(__file__).parent / "templates" / "ea_template.docx"

# Výchozí kritéria MKH (použijí se pokud EAData není zadán)
_DEFAULT_MKH_KRITERIA: list[MKHKriterium] = [
    MKHKriterium("Čistá současná hodnota (NPV)", "tis. Kč", "max", 30.0, "npv"),
    MKHKriterium("Prostá doba návratnosti", "roky", "min", 25.0, "td"),
    MKHKriterium("Úspora energie", "MWh/rok", "max", 25.0, "mwh"),
    MKHKriterium("Úspora nákladů na energii", "tis. Kč/rok", "max", 20.0, "kc"),
]

# Legislativa a normy pro EA
_LEGISLATIVA_EA = [
    "Zákon č. 406/2000 Sb., o hospodaření energií, ve znění pozdějších předpisů",
    "Vyhláška č. 140/2021 Sb., o energetickém auditu",
    "Vyhláška č. 264/2020 Sb., o energetické náročnosti budov (PENB)",
    "Vyhláška č. 193/2007 Sb., o účinnosti užití energie při rozvodu tepelné energie",
    "ČSN EN ISO 6946:2017 – Tepelný odpor a součinitel prostupu tepla stavebních prvků",
    "ČSN 73 0540-2:2011 – Tepelná ochrana budov – Požadavky",
    "ČSN EN ISO 13789:2009 – Tepelné chování budov – Součinitel tepelné ztráty",
    "ČSN EN ISO 50001:2018 – Systémy managementu hospodaření s energií",
]

# Obsah dokumentu (pro statický přehled)
_OBSAH_EA = [
    (0, "1. SOUHRN ENERGETICKÉHO AUDITU"),
    (1, "1.1 Identifikační údaje zadavatele"),
    (1, "1.2 Identifikační údaje energetického specialisty"),
    (1, "1.3 Základní údaje o EA"),
    (1, "1.4 Souhrn příležitostí ke snížení energetické náročnosti (příl. 1B)"),
    (1, "1.5 Program realizace příležitostí"),
    (0, "2. VYMEZENÍ PŘEDMĚTU ENERGETICKÉHO AUDITU"),
    (1, "2.1 Hranice energetického hospodářství"),
    (1, "2.2 Specifikace předmětu EA"),
    (1, "2.3 Vstupní podklady"),
    (0, "3. POPIS STÁVAJÍCÍHO STAVU"),
    (1, "3.1 Bilance energetických vstupů (příl. 3)"),
    (1, "3.2 Analýza užití energie (příl. 4)"),
    (1, "3.3 Ukazatele energetické náročnosti EnPI (příl. 5)"),
    (1, "3.4 Historie spotřeby energie"),
    (1, "3.5 Tepelně technické vlastnosti obálky budovy"),
    (1, "3.6 Technické systémy budovy"),
    (1, "3.7 Systém managementu hospodaření s energií (EnMS)"),
    (1, "3.8 Energetická bilance stávajícího stavu"),
    (0, "4. VYHODNOCENÍ STÁVAJÍCÍHO STAVU"),
    (1, "4.1 Vyhodnocení účinnosti užití energie ve zdrojích energie"),
    (1, "4.2 Principy výpočtu tepelně technického hodnocení"),
    (1, "4.3 Klasifikace obálky budovy"),
    (1, "4.4 Vyhodnocení plnění legislativních požadavků"),
    (1, "4.5 Emisní bilance"),
    (0, "5. PŘÍLEŽITOSTI KE SNÍŽENÍ ENERGETICKÉ NÁROČNOSTI"),
    (1, "5.1 Přehled navržených příležitostí"),
    (1, "5.2 Popis příležitostí"),
    (1, "5.3 Ekonomická analýza (příl. 7)"),
    (1, "5.4 Ekologické hodnocení (příl. 8)"),
    (1, "5.5 Multikriteriální hodnocení (příl. 9)"),
    (1, "5.6 Možnosti využití dotačních programů"),
    (1, "5.7 Primární energie z neobnovitelných zdrojů"),
    (0, "6. ZÁVĚR A DOPORUČENÍ ENERGETICKÉHO SPECIALISTY"),
    (0, "7. PŘÍLOHY"),
]


# ── Veřejné API ───────────────────────────────────────────────────────────────

def generuj_ea(budova: BuildingInfo, result: ProjectResult) -> BytesIO:
    """Sestaví EA dokument dle vyhl. 140/2021 Sb. a vrátí BytesIO ke stažení."""
    doc = Document(_TEMPLATE)

    _titulni_strana(doc, budova)
    _obsah(doc)

    _sekce_1_souhrn(doc, budova, result)
    _sekce_2_vymezeni(doc, budova)
    _sekce_3_popis(doc, budova, result)
    _sekce_4_vyhodnoceni(doc, budova, result)
    _sekce_5_prilezitosti(doc, budova, result)
    _sekce_6_zaver(doc, budova)
    _sekce_7_prilohy(doc, budova)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Titulní strana ────────────────────────────────────────────────────────────

def _titulni_strana(doc: Document, budova: BuildingInfo) -> None:
    h = doc.add_heading("ENERGETICKÝ AUDIT", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if budova.objekt_nazev:
        h2 = doc.add_heading(budova.objekt_nazev, level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    ea = budova.ea_data
    ev_cislo = (ea.evidencni_cislo_ea if ea and ea.evidencni_cislo_ea
                else budova.evidencni_cislo or "–")

    rows = [
        ("Zakázka", budova.nazev_zakazky or "–"),
        ("Zadavatel", budova.zadavatel_nazev or "–"),
        ("Adresa objektu", budova.objekt_adresa or "–"),
        ("Evidenční číslo EA", ev_cislo),
        ("Energetický specialista",
         f"DPU ENERGY s.r.o.\n{budova.zpracovatel_zastupce}"
         if budova.zpracovatel_zastupce else "DPU ENERGY s.r.o."),
        ("Číslo oprávnění", budova.cislo_opravneni or "–"),
        ("Datum vypracování", budova.datum or "–"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        S.cell_text(c0, label, bold=True, font_size=10)
        S.cell_text(c1, value, font_size=10)
        S.set_col_width(c0, Cm(5.0))
        S.set_col_width(c1, Cm(11.0))

    doc.add_page_break()


# ── Obsah ─────────────────────────────────────────────────────────────────────

def _obsah(doc: Document) -> None:
    doc.add_heading("OBSAH", level=2)

    para = doc.add_paragraph()
    run = para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fc_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fc_sep)
    fallback = OxmlElement("w:t")
    fallback.text = "[Obsah se aktualizuje po otevření v MS Word: Ctrl+A → F9]"
    run._r.append(fallback)
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run._r.append(fc_end)

    doc.add_paragraph()
    doc.add_paragraph("Přehled sekcí dokumentu:")
    for uroven, text in _OBSAH_EA:
        indent = "\u00a0\u00a0\u00a0\u00a0" * uroven
        p = doc.add_paragraph(f"{indent}{text}")
        if uroven == 0:
            for run in p.runs:
                run.bold = True
        for run in p.runs:
            run.font.size = Pt(9 if uroven > 0 else 10)

    doc.add_page_break()


# ── 1. Souhrn energetického auditu (§6) ───────────────────────────────────────

def _sekce_1_souhrn(doc: Document, budova: BuildingInfo,
                    result: ProjectResult) -> None:
    doc.add_heading("1. SOUHRN ENERGETICKÉHO AUDITU", level=1)

    ea = budova.ea_data

    # 1.1 Identifikační údaje zadavatele
    doc.add_heading("1.1 Identifikační údaje zadavatele", level=2)
    _kv_tabulka(doc, [
        ("Název / jméno zadavatele", budova.zadavatel_nazev or "–"),
        ("Sídlo / adresa pro doručování", budova.zadavatel_adresa or "–"),
        ("Identifikační číslo (IČ)", budova.zadavatel_ico or "–"),
        ("Kontaktní zástupce", budova.zadavatel_kontakt or "–"),
        ("Telefon", budova.zadavatel_telefon or "–"),
        ("E-mail", budova.zadavatel_email or "–"),
    ])

    # 1.2 Identifikační údaje energetického specialisty
    doc.add_heading("1.2 Identifikační údaje energetického specialisty", level=2)
    _kv_tabulka(doc, [
        ("Zpracovatel", "DPU ENERGY s.r.o."),
        ("Jméno a příjmení specialisty", budova.zpracovatel_zastupce or "–"),
        ("Číslo oprávnění", budova.cislo_opravneni or "–"),
    ])

    # 1.3 Základní údaje o EA
    doc.add_heading("1.3 Základní údaje o energetickém auditu", level=2)
    ev_cislo = (ea.evidencni_cislo_ea if ea and ea.evidencni_cislo_ea
                else budova.evidencni_cislo or "–")
    cil = ea.cil if ea and ea.cil else "–"
    datum_z = ea.datum_zahajeni if ea and ea.datum_zahajeni else "–"
    datum_u = ea.datum_ukonceni if ea and ea.datum_ukonceni else "–"
    _kv_tabulka(doc, [
        ("Evidenční číslo EA", ev_cislo),
        ("Cíl energetického auditu", cil),
        ("Datum zahájení EA", datum_z),
        ("Datum ukončení EA", datum_u),
        ("Předmět auditu", budova.objekt_nazev or "–"),
        ("Adresa předmětu auditu", budova.objekt_adresa or "–"),
    ])

    # 1.4 Evidenční list – souhrn příležitostí (příl. 1B)
    doc.add_heading(
        "1.4 Souhrn příležitostí ke snížení energetické náročnosti", level=2)
    doc.add_paragraph(
        "Evidenční list navržených příležitostí ke snížení energetické náročnosti "
        "seřazený dle priority (příloha č.\u00a01 vyhlášky č.\u00a0140/2021\u00a0Sb., část B)."
    )
    aktivni = result.aktivni
    if aktivni:
        kriteria = ea.mkh_kriteria if ea and ea.mkh_kriteria else _DEFAULT_MKH_KRITERIA
        tabulka_evidencni_list_1b(doc, aktivni, kriteria)
    else:
        doc.add_paragraph("Žádné příležitosti ke snížení energetické náročnosti nebyly identifikovány.")

    # 1.5 Program realizace
    doc.add_heading("1.5 Program realizace příležitostí", level=2)
    program = ea.program_realizace if ea and ea.program_realizace else ""
    if program:
        doc.add_paragraph(program)
    else:
        doc.add_paragraph(
            "Přínosy realizovaných opatření budou průběžně vyhodnocovány na základě "
            "fakturovaných spotřeb energií a porovnávány s výchozím stavem upraveným "
            "o klimatické korekce. Vyhodnocení bude prováděno ročně v rámci "
            "pravidelného reportingu."
        )


# ── 2. Vymezení předmětu energetického auditu (§7) ────────────────────────────

def _sekce_2_vymezeni(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("2. VYMEZENÍ PŘEDMĚTU ENERGETICKÉHO AUDITU", level=1)

    # 2.1 Hranice energetického hospodářství
    doc.add_heading("2.1 Hranice energetického hospodářství", level=2)
    doc.add_paragraph(
        "Hodnocené energetické hospodářství zahrnuje budovy, technické systémy "
        "a energetické vstupy předmětu energetického auditu. "
        "Energetické hospodářství je vymezeno územně dle níže uvedené specifikace."
    )
    doc.add_paragraph(
        "Do hodnoceného energetického hospodářství nejsou zahrnuty spotřeby energie "
        "třetích osob realizované prostřednictvím jejich odběrných míst a přímých "
        "smluvních vztahů s dodavatelem energie."
    )

    # 2.2 Specifikace předmětu EA
    doc.add_heading("2.2 Specifikace předmětu energetického auditu", level=2)
    for line in _nonempty([
        budova.objekt_nazev,
        budova.objekt_adresa,
        f"Katastrální území: {budova.objekt_ku}" if budova.objekt_ku else None,
        f"Parcelní číslo: {budova.objekt_parcelni_cislo}" if budova.objekt_parcelni_cislo else None,
        f"Druh činnosti: {budova.druh_cinnosti}" if budova.druh_cinnosti else None,
        f"Počet zaměstnanců / uživatelů: {budova.pocet_zamestnancu}" if budova.pocet_zamestnancu else None,
        f"Provozní režim: {budova.provozni_rezim}" if budova.provozni_rezim else None,
        f"Účel EA/EP: {budova.ucel_ep}" if budova.ucel_ep else None,
    ]):
        doc.add_paragraph(line)

    budovy_s_daty = [b for b in budova.budovy if b.objem_m3 > 0 or b.nazev]
    if budovy_s_daty:
        doc.add_heading("Soupis budov a geometrické parametry", level=3)
        tabulka_budovy(doc, budovy_s_daty)

    if budova.prostory:
        doc.add_heading("Využití prostor", level=3)
        for p in budova.prostory:
            text = p.nazev
            if p.ucel:
                text += f" – {p.ucel}"
            if p.provoz:
                text += f" ({p.provoz})"
            doc.add_paragraph(text, style="List Bullet")

    # 2.3 Vstupní podklady
    doc.add_heading("2.3 Vstupní podklady ke zpracování energetického auditu", level=2)

    if budova.podklady:
        tabulka_podklady(doc, budova.podklady)
    if budova.poznamka_podklady:
        doc.add_paragraph(budova.poznamka_podklady)

    doc.add_heading("Použitá legislativa a technické normy", level=3)
    for predpis in _LEGISLATIVA_EA:
        doc.add_paragraph(predpis, style="List Bullet")

    ea = budova.ea_data
    plan = ea.plan_text if ea and ea.plan_text else ""
    if plan:
        doc.add_heading("Plán energetického auditu – shrnutí (§4)", level=3)
        doc.add_paragraph(plan)


# ── 3. Popis stávajícího stavu (§8) ──────────────────────────────────────────

def _sekce_3_popis(doc: Document, budova: BuildingInfo,
                   result: ProjectResult) -> None:
    doc.add_heading("3. POPIS STÁVAJÍCÍHO STAVU", level=1)

    ea = budova.ea_data

    # 3.1 Bilance energetických vstupů (příl. 3)
    if ea and ea.energonositele:
        doc.add_heading(
            "3.1 Bilance energetických vstupů (příl. 3 vyhl. č.\u00a0140/2021\u00a0Sb.)",
            level=2)
        doc.add_paragraph(
            "Přehled energetických vstupů energetického hospodářství s klasifikací "
            "energonositelů dle kategorie (NOZE / OZE / Druhotné) a oblasti užití "
            "(Budovy / Výrobní procesy / Doprava)."
        )
        tabulka_bilance_vstupu_ea(doc, ea.energonositele, result.energie)

    # 3.2 Analýza užití energie (příl. 4)
    doc.add_heading(
        "3.2 Analýza užití energie (příl. 4 vyhl. č.\u00a0140/2021\u00a0Sb.)", level=2)
    if budova.bilance_pouziti:
        doc.add_paragraph(
            "Rozpad celkové spotřeby energie dle způsobu užití (vytápění, chlazení, "
            "příprava TUV, větrání, osvětlení, technologie, PHM)."
        )
        tabulka_bilance_pouziti(doc, budova.bilance_pouziti)
    else:
        doc.add_paragraph(
            "Analýza užití energie dle přílohy č.\u00a04 nebyla zadána. "
            "Viz záložka Obálka & systémy, sekce Bilance dle účelu."
        )

    # 3.3 Ukazatele EnPI (příl. 5)
    doc.add_heading(
        "3.3 Ukazatele energetické náročnosti EnPI (příl. 5 vyhl. č.\u00a0140/2021\u00a0Sb.)",
        level=2)
    if ea and ea.enpi:
        tabulka_enpi(doc, ea.enpi)
    else:
        doc.add_paragraph(
            "Ukazatele energetické náročnosti (EnPI) nebyly zadány. "
            "Viz záložka EA – rozšíření."
        )

    # 3.4 Historie spotřeby energie
    doc.add_heading("3.4 Historie spotřeby energie", level=2)
    kd = budova.klimaticka_data
    doc.add_heading("Klimatická data a klimatická korekce", level=3)
    if kd:
        for line in _nonempty([
            f"Lokalita: {kd.lokalita}" if kd.lokalita else None,
            f"Normované denostupně D\u00a0=\u00a0{kd.stupnodni_normovane:,.0f}\u00a0°C·dny/rok".replace(",", "\u00a0")
            if kd.stupnodni_normovane else None,
            f"Vnitřní výpočtová teplota ti\u00a0=\u00a0{kd.teplota_vnitrni:.0f}\u00a0°C",
            f"Venkovní výpočtová teplota te\u00a0=\u00a0{kd.teplota_exterieru:.0f}\u00a0°C",
        ]):
            doc.add_paragraph(line)
        if kd.stupnodni_normovane > 0:
            doc.add_paragraph(
                "Spotřeba tepla na vytápění je klimaticky korigována na normované denostupně. "
                f"Korigovaná spotřeba\u00a0=\u00a0skutečná spotřeba\u00a0×\u00a0D_norm\u00a0/\u00a0D_skut, "
                f"D_norm\u00a0=\u00a0{kd.stupnodni_normovane:,.0f}\u00a0°C·dny/rok "
                f"(lokalita: {kd.lokalita or '–'}).".replace(",", "\u00a0")
            )
    else:
        doc.add_paragraph("Klimatická data lokality nebyla zadána.")

    if budova.historie_spotreby:
        doc.add_heading("Přehled měsíčních spotřeb", level=3)
        nosice: dict[str, list] = {}
        for h in budova.historie_spotreby:
            nosice.setdefault(h.energonosic or "Neurčeno", []).append(h)
        for nosic, roky in nosice.items():
            doc.add_paragraph(f"{nosic}:")
            tabulka_mesicni_spotreba(doc, roky)

        doc.add_heading("Roční souhrn spotřeb a nákladů", level=3)
        d_norm = kd.stupnodni_normovane if kd and kd.stupnodni_normovane else 0.0
        has_korekce = d_norm > 0
        headers = ["Rok", "Energonosič", "Spotřeba [MWh/rok]",
                   "Náklady [tis. Kč/rok]", "Skut. D [°C·d]"]
        if has_korekce:
            headers += ["D_norm/D_skut", "Korig. spotřeba [MWh]"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hr = tbl.rows[0]
        for i, h in enumerate(headers):
            S.cell_text(hr.cells[i], h, bold=True, font_size=9)
            S.set_cell_bg(hr.cells[i], S.COLOR_TBL_HEAD_HEX)
        for h in budova.historie_spotreby:
            row = tbl.add_row()
            d_skut = h.stupnodni or 0.0
            if has_korekce and d_skut > 0:
                pomer = d_norm / d_skut
                korig = h.spotreba_mwh * pomer if h.spotreba_mwh else 0.0
                pomer_str = f"{pomer:.3f}"
                korig_str = f"{korig:,.1f}".replace(",", "\u00a0")
            else:
                pomer_str = "–"
                korig_str = "–"
            vals = [
                str(h.rok) if h.rok else "–",
                h.energonosic or "–",
                f"{h.spotreba_mwh:,.1f}".replace(",", "\u00a0") if h.spotreba_mwh else "–",
                f"{h.naklady_kc / 1000:.1f}" if h.naklady_kc else "–",
                f"{d_skut:,.0f}".replace(",", "\u00a0") if d_skut else "–",
            ]
            if has_korekce:
                vals += [pomer_str, korig_str]
            for i, val in enumerate(vals):
                S.cell_text(row.cells[i], val, font_size=9)
        if has_korekce:
            col_widths = [Cm(1.2), Cm(3.5), Cm(2.8), Cm(2.8), Cm(2.2), Cm(2.2), Cm(2.8)]
        else:
            col_widths = [Cm(1.5), Cm(4.0), Cm(3.5), Cm(3.5), Cm(3.0)]
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    S.set_col_width(cell, col_widths[i])

        # Grafy spotřeby
        doc.add_heading("Grafické vyhodnocení spotřeby energie", level=3)
        buf_rocni = graf_spotreba_rocni(budova)
        _vloz_obrazek(doc, buf_rocni,
                      "Roční spotřeba energie dle energonosiče [MWh]", cislo=1)
        # Měsíční grafy per energonosič
        graf_cislo = 2
        for en in sorted({h.energonosic for h in budova.historie_spotreby}):
            buf_mes = graf_spotreba_mesicni(budova, en)
            if buf_mes:
                _vloz_obrazek(doc, buf_mes,
                              f"Vývoj měsíčních spotřeb – {en} [MWh]",
                              cislo=graf_cislo)
                graf_cislo += 1
    else:
        doc.add_paragraph(
            "Data historické spotřeby za poslední roky nebyla zadána. "
            "Dle §8 odst. 3 vyhl. č.\u00a0140/2021\u00a0Sb. je potřeba uvést historii "
            "spotřeby za minimálně 2 předchozí ucelené kalendářní roky."
        )

    # 3.5 Tepelně technické vlastnosti
    doc.add_heading("3.5 Tepelně technické vlastnosti obálky budovy", level=2)
    doc.add_paragraph(
        "Systémová hranice budovy je definována dle ČSN EN ISO 13789:2009 jako obálka "
        "ohraničující vytápěný prostor od venkovního prostředí nebo od nevytápěných prostor."
    )
    if budova.konstrukce:
        SKUPINY = [
            ("stena",   "Fasády a obvodové stěny"),
            ("podlaha", "Podlahy"),
            ("strecha", "Střechy"),
            ("okno",    "Okna a prosklené výplně otvorů"),
            ("dvere",   "Vstupní dveře"),
        ]
        _visited_oken = False
        for typ, nadpis in SKUPINY:
            skupina = [k for k in budova.konstrukce if k.typ == typ]
            if not skupina:
                continue
            if typ in ("okno", "dvere"):
                if not _visited_oken:
                    doc.add_heading("Okna, dveře a prosklené výplně otvorů", level=3)
                    _visited_oken = True
            else:
                doc.add_heading(nadpis, level=3)
            tabulka_konstrukce(doc, skupina)
    else:
        doc.add_paragraph("Tepelně technické vlastnosti stavebních konstrukcí nebyly zadány.")

    # Fotografie budovy (obálka, exteriér)
    foto_budova = [f for f in budova.fotografie if f.sekce == "budova"]
    if foto_budova:
        doc.add_heading("Fotodokumentace objektu", level=3)
        for i, foto in enumerate(foto_budova, 1):
            _vloz_obrazek(doc, foto.data, foto.popisek,
                          sirka_cm=foto.sirka_cm, cislo=0)

    # 3.6 Technické systémy budovy
    doc.add_heading("3.6 Technické systémy budovy", level=2)
    ts = budova.technicke_systemy
    doc.add_heading("Vytápění a příprava teplé vody", level=3)
    for nazev, sys in [("Vytápění", ts.vytapeni), ("Příprava teplé vody (TUV)", ts.tuv)]:
        doc.add_heading(nazev, level=4)
        _popis_systemu(doc, sys)
    doc.add_heading("Větrání a klimatizace", level=3)
    _popis_systemu(doc, ts.vzt)
    doc.add_heading("Osvětlení", level=3)
    _popis_systemu(doc, ts.osvetleni)
    if ts.mereni_ridici:
        doc.add_heading("Měření, regulace a monitoring (MaR)", level=3)
        doc.add_paragraph(ts.mereni_ridici)

    # Fotografie technických zařízení (kotle, rozvaděče, schémata)
    foto_technika = [f for f in budova.fotografie if f.sekce == "technika"]
    if foto_technika:
        doc.add_heading("Fotodokumentace technických systémů", level=3)
        for foto in foto_technika:
            _vloz_obrazek(doc, foto.data, foto.popisek,
                          sirka_cm=foto.sirka_cm, cislo=0)

    # 3.7 EnMS
    doc.add_heading(
        "3.7 Systém managementu hospodaření s energií (EnMS)", level=2)
    if budova.enms:
        doc.add_paragraph(
            "Hodnocení úrovně systému managementu hospodaření s energií dle "
            "ČSN EN ISO 50001:2018, škála: 1 – nesplnění, 2 – částečné splnění, "
            "3 – plné splnění."
        )
        tabulka_enms(doc, budova.enms)
        hodnoty = [o.hodnoceni for o in budova.enms.oblasti if o.hodnoceni > 0]
        if hodnoty:
            prumer = sum(hodnoty) / len(hodnoty)
            doc.add_paragraph(
                f"Průměrné hodnocení EnMS: {prumer:.1f}\u00a0/\u00a03,0. "
                + ("Organizace má zaveden a certifikován EnMS dle ISO 50001."
                   if budova.enms.certifikovan
                   else "Organizace nemá certifikovaný EnMS dle ISO 50001.")
            )
        if budova.enms.komentar:
            doc.add_paragraph(f"Komentář: {budova.enms.komentar}")
    else:
        doc.add_paragraph(
            "Hodnocení EnMS nebylo provedeno. Doporučujeme zhodnotit zavedení systému "
            "managementu hospodaření s energií dle ČSN EN ISO 50001:2018."
        )

    # 3.8 Energetická bilance stávajícího stavu
    doc.add_heading("3.8 Energetická bilance stávajícího stavu", level=2)
    doc.add_paragraph(
        "Výchozí roční energetická bilance energetického hospodářství je sestavena "
        "pro referenční období na základě normalizovaných (klimaticky korigovaných) spotřeb. "
        "Tato bilance je použita jako výchozí stav pro výpočet úspor navrhovaných příležitostí."
    )
    tabulka_bilance_rocni(doc, result.energie)
    # Výsečový graf podílů spotřeby
    buf_podily = graf_podily_spotreby(budova)
    _vloz_obrazek(doc, buf_podily,
                  "Procentní podíly spotřeby a nákladů dle energonosiče", cislo=0)
    doc.add_paragraph(
        "Podrobná bilance spotřeb a nákladů v členění stávající stav / po realizaci příležitostí:"
    )
    tabulka_energeticka_bilance(doc, result.energie)


# ── 4. Vyhodnocení stávajícího stavu ─────────────────────────────────────────

def _sekce_4_vyhodnoceni(doc: Document, budova: BuildingInfo,
                          result: ProjectResult) -> None:
    doc.add_heading("4. VYHODNOCENÍ STÁVAJÍCÍHO STAVU", level=1)

    # 4.1 Vyhodnocení účinnosti užití energie (§ 6 vyhl. 406/2000 Sb.)
    doc.add_heading(
        "4.1 Vyhodnocení účinnosti užití energie ve zdrojích energie",
        level=2)
    doc.add_paragraph(
        "Dle § 6 zákona č.\u00a0406/2000\u00a0Sb. je vlastník zdroje tepla s jmenovitým "
        "výkonem nad 20\u00a0kW povinen zajistit pravidelnou kontrolu kotlů a rozvodů. "
        "Splnění minimální účinnosti je posuzováno dle vyhl. č.\u00a0441/2012\u00a0Sb."
    )
    tabulka_ucinnost_zdroju(doc, budova.technicke_systemy)

    # 4.2 Principy výpočtu
    doc.add_heading("4.2 Principy výpočtu tepelně technického hodnocení", level=2)
    doc.add_paragraph(
        "Součinitele prostupu tepla U jsou vypočteny dle ČSN EN ISO 6946:2017 "
        "z tepelných odporů vrstev (R\u00a0=\u00a0d/λ) a povrchových odporů (Rsi, Rse). "
        "Pro okna a dveře jsou použity hodnoty U zadané přímo (dle výrobce)."
    )
    doc.add_paragraph(
        "Požadované (UN), doporučené (Urec) a doporučené pro pasivní budovy (Upas) hodnoty "
        "jsou převzaty z ČSN 73 0540-2:2011, tabulka 3, "
        "pro budovy s převažující vnitřní teplotou 18–22\u00a0°C."
    )

    # 4.3 Klasifikace obálky budovy
    doc.add_heading("4.3 Klasifikace obálky budovy", level=2)
    if budova.konstrukce:
        doc.add_paragraph(
            "Hodnocení dle ČSN 73 0540-2:2011. "
            "UN = požadovaná, Urec = doporučená, Upas = pro pasivní budovy."
        )
        tabulka_u_porovnani(doc, budova.konstrukce)
        uem = vypocitej_uem_z_konstrukci(budova.konstrukce)
        if uem is not None:
            doc.add_paragraph(
                f"Průměrný součinitel prostupu tepla obálky budovy: "
                f"Uem\u00a0=\u00a0{uem:.3f}\u00a0W/(m²K)."
            )
    else:
        doc.add_paragraph("Tepelně technické vlastnosti stavebních konstrukcí nebyly zadány.")

    klas = result.klasifikace_pred
    if klas:
        doc.add_heading("Energetický štítek obálky budovy", level=3)
        tabulka_klasifikace(doc, klas)
        doc.add_paragraph(
            f"Budova je zařazena do energetické třídy {klas.trida} "
            f"(Uem\u00a0=\u00a0{klas.uem:.3f}\u00a0W/m²K, "
            f"Uem,N\u00a0=\u00a0{klas.uem_n:.3f}\u00a0W/m²K)."
        )

    # 4.4 Vyhodnocení plnění legislativních požadavků (§9 vyhl. 140/2021)
    doc.add_heading(
        "4.4 Vyhodnocení plnění legislativních požadavků (§9 vyhl. č.\u00a0140/2021\u00a0Sb.)",
        level=2)
    if klas is None:
        doc.add_paragraph(
            "Vstupní data pro posouzení obálky budovy (Uem, faktor tvaru A/V) "
            "nebyla zadána – legislativní hodnocení nelze provést."
        )
    else:
        splneni = klas.pomer <= 1.0
        if splneni:
            doc.add_paragraph(
                f"Budova splňuje požadavky ČSN\u00a073\u00a00540-2 na průměrný "
                f"součinitel prostupu tepla obálky. "
                f"Uem\u00a0=\u00a0{klas.uem:.3f}\u00a0W/m²K "
                f"(požadavek Uem,N\u00a0=\u00a0{klas.uem_n:.3f}\u00a0W/m²K, "
                f"poměr Uem/Uem,N\u00a0=\u00a0{klas.pomer:.3f}). "
                f"Klasifikační třída obálky: {klas.trida}."
            )
        else:
            doc.add_paragraph(
                f"Budova NESPLŇUJE požadavky ČSN\u00a073\u00a00540-2 na průměrný "
                f"součinitel prostupu tepla obálky. "
                f"Uem\u00a0=\u00a0{klas.uem:.3f}\u00a0W/m²K překračuje referenční hodnotu "
                f"Uem,N\u00a0=\u00a0{klas.uem_n:.3f}\u00a0W/m²K "
                f"(poměr Uem/Uem,N\u00a0=\u00a0{klas.pomer:.3f}). "
                f"Klasifikační třída obálky: {klas.trida}. "
                f"Navrhovaná opatření přispívají ke splnění tohoto požadavku."
            )
    doc.add_paragraph(
        "Dle §9 vyhlášky č.\u00a0140/2021\u00a0Sb. musí energetický audit obsahovat "
        "posouzení, zda předmět auditu splňuje požadavky na energetickou náročnost "
        "stanovené prováděcím právním předpisem."
    )

    # 4.5 Emisní bilance
    doc.add_heading("4.5 Emisní bilance", level=2)
    doc.add_paragraph(
        "Emisní bilance pro stávající stav a stav po realizaci navrhovaných příležitostí. "
        "Emise jsou vypočteny z měrných emisních faktorů dle přílohy č.\u00a08 "
        "vyhlášky č.\u00a0140/2021\u00a0Sb."
    )
    tabulka_emise(doc, result.emise_pred, result.emise_po)
    pred = result.emise_pred
    po = result.emise_po
    if pred and po:
        delta_co2 = pred.co2_kg - po.co2_kg
        if delta_co2 > 0:
            doc.add_paragraph(
                f"Realizací navržených příležitostí dojde ke snížení emisí CO\u2082 "
                f"o {delta_co2:,.0f}\u00a0kg/rok.".replace(",", "\u00a0")
            )


# ── 5. Příležitosti ke snížení energetické náročnosti (§9) ───────────────────

def _sekce_5_prilezitosti(doc: Document, budova: BuildingInfo,
                           result: ProjectResult) -> None:
    doc.add_heading(
        "5. PŘÍLEŽITOSTI KE SNÍŽENÍ ENERGETICKÉ NÁROČNOSTI", level=1)

    aktivni = result.aktivni
    ea = budova.ea_data
    kriteria = ea.mkh_kriteria if ea and ea.mkh_kriteria else _DEFAULT_MKH_KRITERIA

    # 5.1 Přehled navržených příležitostí – souhrnná tabulka variant
    doc.add_heading("5.1 Přehled navržených příležitostí ke snížení energetické náročnosti",
                    level=2)
    doc.add_paragraph(
        "Navrhované příležitosti ke snížení energetické náročnosti jsou definovány "
        "v souladu s §9 vyhl. č.\u00a0140/2021\u00a0Sb. Každá příležitost je "
        "posouzena z hlediska technického, ekonomického a ekologického."
    )
    if aktivni:
        doc.add_paragraph(
            "Druhy opatření dle výše investice: Bezná. – organizační opatření bez nákladů; "
            "Nízkoná. – malé investice s rychlou návratností; Středoná. – systémová opatření; "
            "Vysokoná. – celkové rekonstrukce (zateplení, výměna zdroje)."
        )
        doc.add_heading("Souhrnná tabulka navrhovaných příležitostí", level=3)
        tabulka_soucet_variant(doc, aktivni, result.ekonomika_parametry)
        buf_inv = graf_investice_uspory(aktivni)
        _vloz_obrazek(doc, buf_inv,
                      "Poměr investičních nákladů a ročních úspor jednotlivých příležitostí",
                      cislo=0)
    else:
        doc.add_paragraph("Žádné příležitosti ke snížení energetické náročnosti nebyly identifikovány.")
        return

    # 5.2 Popis příležitostí – per opatření s CF tabulkou
    doc.add_heading("5.2 Popis příležitostí ke snížení energetické náročnosti", level=2)
    par = result.ekonomika_parametry
    for r in aktivni:
        doc.add_heading(f"{r.id} – {r.nazev}", level=3)
        info = OP_INFO.get(r.id, {})
        popis = info.get("popis", "")
        if popis:
            doc.add_paragraph(popis)

        # Klíčové parametry
        mwh = r.uspora_teplo + r.uspora_zp + r.uspora_ee
        doc.add_paragraph(
            f"Odhadované investiční náklady: {S.fmt_kc(r.investice)}.  "
            f"Roční úspora energie: {S.fmt_mwh(mwh)}.  "
            f"Roční úspora nákladů na energie: {S.fmt_kc(r.uspora_kc)}.  "
            f"Prostá doba návratnosti: {S.fmt_nav(r.prosta_navratnost)}."
        )
        if r.ekonomika:
            ek = r.ekonomika
            doc.add_paragraph(
                f"Čistá současná hodnota (NPV): {S.fmt_kc(ek.npv) if ek.npv is not None else '–'}.  "
                f"Vnitřní výnosové procento (IRR): {S.fmt_irr(ek.irr) if ek.irr is not None else '–'}.  "
                f"Reálná doba návratnosti (Tsd): {S.fmt_tsd(ek.tsd, 20)}."
            )

        dotace = info.get("dotace", "")
        if dotace:
            doc.add_paragraph(f"Možnosti finanční podpory: {dotace}")

        # Cash flow tabulka + graf rok po roku
        if r.investice > 0 and r.uspora_kc > 0:
            doc.add_heading(f"Ekonomické hodnocení příležitosti {r.id}", level=4)
            horizont = par.horizont if par else 20
            tabulka_cash_flow(doc, r, par, horizont)
            buf_cf = graf_cf_kumulativni(r, par, horizont)
            _vloz_obrazek(doc, buf_cf,
                          f"Kumulativní cash flow – {r.id} (nediskontovaný a diskontovaný)",
                          cislo=0)

    # 5.3 Ekonomická analýza (příl. 7)
    doc.add_heading(
        "5.3 Ekonomická analýza souhrnu příležitostí (příl. 7 vyhl. č.\u00a0140/2021\u00a0Sb.)",
        level=2)
    if par:
        doc.add_paragraph(
            f"Hodnotící horizont: {par.horizont}\u00a0let, "
            f"diskontní sazba: {par.diskontni_sazba * 100:.1f}\u00a0%, "
            f"inflace cen energií: {par.inflace_energie * 100:.1f}\u00a0%/rok."
        )
    tabulka_ekonomika(doc, aktivni, result.ekonomika_projekt, par)

    celk_inv = result.celkova_investice
    celk_kc = result.celkova_uspora_kc
    ek = result.ekonomika_projekt
    doc.add_paragraph(
        f"Celková investice všech příležitostí: {S.fmt_kc(celk_inv)}. "
        f"Celková roční úspora: {S.fmt_kc(celk_kc)}. "
        f"Prostá návratnost: {S.fmt_nav(result.prosta_navratnost_celkem)}."
        + (f" NPV souhrnu: {S.fmt_kc(ek.npv)}." if ek else "")
    )

    # 5.4 Ekologické hodnocení (příl. 8)
    doc.add_heading(
        "5.4 Ekologické hodnocení (příl. 8 vyhl. č.\u00a0140/2021\u00a0Sb.)", level=2)
    doc.add_paragraph(
        "Ekologické hodnocení je zpracováno dle emisních faktorů z přílohy č.\u00a08 "
        "vyhlášky č.\u00a0140/2021\u00a0Sb. "
        "Ekvivalentní prašná substance (EPS)\u00a0=\u00a01,0·TZL\u00a0+"
        "\u00a00,88·NOₓ\u00a0+\u00a00,54·SO₂\u00a0+\u00a00,64·NH₃."
    )
    tabulka_emise(doc, result.emise_pred, result.emise_po)
    pred = result.emise_pred
    po = result.emise_po
    if pred and po:
        delta_co2 = pred.co2_kg - po.co2_kg
        delta_pct = delta_co2 / pred.co2_kg * 100 if pred.co2_kg > 0 else 0
        if delta_co2 > 0:
            doc.add_paragraph(
                f"Realizací navržených příležitostí dojde ke snížení emisí CO₂ "
                f"o {delta_co2:,.0f}\u00a0kg/rok ({delta_pct:.1f}\u00a0%).".replace(",", "\u00a0")
            )

    # 5.5 Multikriteriální hodnocení (příl. 9)
    doc.add_heading(
        "5.5 Multikriteriální hodnocení příležitostí (příl. 9 vyhl. č.\u00a0140/2021\u00a0Sb.)",
        level=2)
    total_vaha = sum(k.vaha for k in kriteria if k.vaha > 0)
    doc.add_paragraph(
        "Příležitosti jsou seřazeny metodou váženého součtu (min-max normalizace) dle příl. č.\u00a09 "
        f"vyhlášky č.\u00a0140/2021\u00a0Sb. Součet vah: {total_vaha:.0f}\u00a0%. "
        "Pořadí 1 = nejvýhodnější příležitost."
    )
    for k in kriteria:
        if k.vaha > 0:
            doc.add_paragraph(
                f"• {k.nazev} [{k.jednotka}] – {k.typ}imalizace, váha: {k.vaha:.0f}\u00a0%",
                style="List Bullet",
            )
    tabulka_mkh(doc, aktivni, kriteria)

    # 5.6 Dotační příležitosti
    doc.add_heading("5.6 Možnosti využití dotačních programů", level=2)
    doc.add_paragraph(
        "V závislosti na charakteru vlastníka (obec, kraj, právnická osoba, fyzická osoba – podnikatel) "
        "jsou k dispozici tyto hlavní dotační programy:"
    )
    for program in [
        ("Nová zelená úsporám (NZÚ) – SFŽP",
         "Program MŽP / SFŽP pro snižování energetické náročnosti budov. Podporuje zateplení obálky, "
         "výměnu oken, výměnu zdroje tepla, instalaci OZE. Výše podpory závisí na dosažené třídě ENB po "
         "rekonstrukci. Podmínkou je zpracování PENB před a po opatření."),
        ("Operační program Životní prostředí (OPŽP) – výzvy pro veřejné budovy",
         "Dotace EU pro snižování energetické náročnosti veřejných budov. Vhodné pro municipality a "
         "příspěvkové organizace. Podmínkou je zpracování energetického posudku nebo auditu a dosažení "
         "min. 30 % úspory primární energie."),
        ("MPO EFEKT – program Ministerstva průmyslu a obchodu",
         "Program pro zvyšování energetické účinnosti a využívání OZE. Přímé dotace pro malé a střední "
         "podniky, školy a obce. Výzvy jsou vyhlašovány zpravidla 1× ročně."),
        ("Modernizační fond – Enerfin / Čistá energie",
         "Fond pro velké projekty OZE a úspor energie. Vhodný pro projekty s investicí nad 5 mil. Kč."),
    ]:
        doc.add_heading(program[0], level=3)
        doc.add_paragraph(program[1])

    # 5.7 Primární energie
    doc.add_heading(
        "5.7 Primární energie z neobnovitelných zdrojů", level=2)
    penb = budova.penb
    if penb:
        tabulka_penb(doc, penb)
        if penb.poznamka:
            doc.add_paragraph(f"Poznámka k PENB: {penb.poznamka}")
    else:
        doc.add_paragraph(
            "Data průkazu energetické náročnosti budovy (PENB) nebyla zadána. "
            "PENB je zpracováván v certifikovaném softwaru a je přikládán jako příloha."
        )


# ── 6. Závěr a doporučení ─────────────────────────────────────────────────────

def _sekce_6_zaver(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("6. ZÁVĚR A DOPORUČENÍ ENERGETICKÉHO SPECIALISTY", level=1)

    doc.add_heading("6.1 Závěr", level=2)
    if budova.okrajove_podminky:
        doc.add_paragraph(budova.okrajove_podminky)
    doc.add_paragraph(
        "Výpočty jsou provedeny v souladu s ČSN EN ISO 6946:2017, ČSN 73 0540-2:2011, "
        "ČSN EN ISO 13789:2009 a metodikou Státní energetické inspekce."
    )
    doc.add_paragraph(
        "Na základě provedeného energetického auditu jsou identifikovány příležitosti "
        "ke snížení energetické náročnosti uvedené v části 5. "
        "Doporučujeme zahájit přípravné práce pro realizaci navrhovaných opatření "
        "a ověřit aktuální podmínky dotačních programů."
    )

    doc.add_heading("6.2 Prohlášení a podpis energetického specialisty", level=2)
    doc.add_paragraph(
        "Energetický audit byl zpracován v souladu s požadavky zákona č.\u00a0406/2000\u00a0Sb. "
        "a vyhlášky č.\u00a0140/2021\u00a0Sb. osobou s oprávněním energetického specialisty. "
        "Energetický specialista prohlašuje, že údaje v auditu jsou pravdivé a úplné."
    )
    doc.add_paragraph()
    sig_rows = [
        ("Energetický specialista:", budova.zpracovatel_zastupce or ""),
        ("Číslo oprávnění:", budova.cislo_opravneni or ""),
        ("Datum:", budova.datum or ""),
        ("Podpis:", ""),
    ]
    tbl = doc.add_table(rows=len(sig_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(sig_rows):
        S.cell_text(tbl.rows[i].cells[0], label, bold=True, font_size=10)
        S.cell_text(tbl.rows[i].cells[1], value, font_size=10)
        S.set_col_width(tbl.rows[i].cells[0], Cm(5.0))
        S.set_col_width(tbl.rows[i].cells[1], Cm(11.0))


# ── 7. Přílohy (§10) ─────────────────────────────────────────────────────────

def _sekce_7_prilohy(doc: Document, budova: BuildingInfo) -> None:
    doc.add_page_break()
    doc.add_heading("7. PŘÍLOHY", level=1)

    doc.add_paragraph(
        "K tomuto energetickému auditu jsou přiloženy následující přílohy "
        "(§10 vyhl. č.\u00a0140/2021\u00a0Sb.):"
    )
    prilohy = [
        "7.1 Plán energetického auditu (§4 vyhl. č. 140/2021 Sb.)",
        "7.2 Seznam požadovaných a obdržených podkladů",
        "7.3 Průkaz energetické náročnosti budovy (PENB) dle vyhl. č. 264/2020 Sb.",
        "7.4 Ilustrativní fotodokumentace předmětu energetického hospodářství",
    ]
    for p in prilohy:
        doc.add_paragraph(p, style="List Bullet")

    # Příloha 7.1 – Plán energetického auditu (příloha č. 2 k vyhl. 140/2021)
    _priloha_plan_ea(doc, budova)

    # Příloha 7.2 – seznam podkladů
    doc.add_heading("7.2 Seznam požadovaných a obdržených podkladů", level=2)
    if budova.podklady:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hr = tbl.rows[0]
        for i, h in enumerate(["Podklad", "Požadován", "Obdržen"]):
            S.cell_text(hr.cells[i], h, bold=True, font_size=9)
            S.set_cell_bg(hr.cells[i], S.COLOR_TBL_HEAD_HEX)
        for p in budova.podklady:
            row = tbl.add_row()
            S.cell_text(row.cells[0], p.nazev, font_size=9)
            S.cell_text(row.cells[1], "Ano", font_size=9)
            S.cell_text(row.cells[2], "Ano" if p.k_dispozici else "Ne", font_size=9)
    else:
        doc.add_paragraph(
            "[Seznam podkladů je uveden v průvodním dokumentu k auditu.]"
        )

    # Příloha 7.4 – fotodokumentace označená sekce="priloha"
    foto_priloha = [f for f in budova.fotografie if f.sekce == "priloha"]
    if foto_priloha:
        doc.add_heading("7.4 Fotodokumentace", level=2)
        for foto in foto_priloha:
            _vloz_obrazek(doc, foto.data, foto.popisek,
                          sirka_cm=foto.sirka_cm, cislo=0)


def _priloha_plan_ea(doc: Document, budova: BuildingInfo) -> None:
    """Příloha č. 7.1 – Plán energetického auditu (příloha č. 2 k vyhl. 140/2021 Sb.)."""
    doc.add_heading("7.1 Plán energetického auditu", level=2)
    doc.add_paragraph(
        "Příloha č. 2 k vyhlášce č.\u00a0140/2021\u00a0Sb. "
        "Plán energetického auditu je dohodou energetického specialisty se "
        "zadavatelem (§\u00a04 odst.\u00a01 vyhl. č.\u00a0140/2021\u00a0Sb.). "
        "Podepsaný plán je povinnou přílohou zprávy o provedeném energetickém auditu "
        "(§\u00a010 písm.\u00a0a) téže vyhlášky)."
    )

    plan = budova.plan_ea

    # Záhlaví – identifikace stran
    doc.add_heading("Identifikace stran", level=3)
    tbl_id = doc.add_table(rows=0, cols=2)
    tbl_id.style = "Table Grid"
    id_rows = [
        ("Zadavatel (název):", budova.zadavatel_nazev),
        ("Adresa zadavatele:", budova.zadavatel_adresa),
        ("IČO:", budova.zadavatel_ico),
        ("Zástupce zadavatele:", plan.zadavatel_zastupce if plan else ""),
        ("Energetický specialista:", budova.zpracovatel_zastupce),
        ("Číslo oprávnění:", budova.cislo_opravneni),
        ("Předmět auditu:", budova.objekt_nazev),
        ("Datum plánu:", plan.datum_planu if plan else ""),
    ]
    for label, val in id_rows:
        row = tbl_id.add_row()
        S.cell_text(row.cells[0], label, bold=True, font_size=9)
        S.cell_text(row.cells[1], val or "–", font_size=9)
        S.set_col_width(row.cells[0], Cm(6))
        S.set_col_width(row.cells[1], Cm(10))

    if plan is None:
        doc.add_paragraph(
            "\n[Plán energetického auditu nebyl v aplikaci vyplněn. "
            "Doplňte jej v záložce 'Plán EA' a exportujte znovu.]"
        )
        return

    # Oddíl 1 – Míra detailu
    doc.add_heading("1. Požadavky na míru detailu provádění EA", level=3)
    doc.add_paragraph(plan.mira_detailu or "–")

    # Oddíl 2 – Předmět EA
    doc.add_heading("2. Předmět energetického auditu", level=3)
    doc.add_paragraph(
        f"Předmět: {plan.predmet_ea or budova.predmet_analyzy or '–'}"
    )
    if plan.lokalizace_predmetu:
        doc.add_paragraph(f"Lokalizace: {plan.lokalizace_predmetu}")

    # Oddíl 3 – Potřeby a očekávání
    doc.add_heading("3. Potřeby zadavatele a jeho očekávání", level=3)
    doc.add_paragraph(plan.potreby_a_cile or "–")

    # Oddíl 4 – Kritéria hodnocení
    doc.add_heading("4. Kritéria pro hodnocení příležitostí", level=3)
    tbl_krit = doc.add_table(rows=0, cols=2)
    tbl_krit.style = "Table Grid"
    krit_rows = [
        ("Ekonomické ukazatele:", plan.ekonomicke_ukazatele),
        ("Horizont hodnocení:", f"{plan.horizont_hodnoceni_let}\u00a0let"),
        ("Diskontní sazba (reálná):", f"{plan.diskontni_sazba_pct:.1f}\u00a0%"),
        ("Roční změna cen energií:", f"{plan.inflace_energie_pct:.1f}\u00a0%"),
        ("Zahrnutí finanční podpory:", "Ano" if plan.zahrnout_financni_podporu else "Ne"),
    ]
    for label, val in krit_rows:
        row = tbl_krit.add_row()
        S.cell_text(row.cells[0], label, bold=True, font_size=9)
        S.cell_text(row.cells[1], val, font_size=9)
        S.set_col_width(row.cells[0], Cm(6))
        S.set_col_width(row.cells[1], Cm(10))
    if plan.mkh_kriteria_popis:
        doc.add_paragraph(f"Vícekriteriální hodnocení: {plan.mkh_kriteria_popis}")

    # Oddíl 5 – Součinnost
    doc.add_heading("5. Požadavky na součinnost zadavatele", level=3)
    doc.add_paragraph(plan.soucinnost_pozadavky or "–")
    if plan.harmonogram:
        doc.add_paragraph(f"Harmonogram: {plan.harmonogram}")

    # Oddíl 6 – Strategické dokumenty
    doc.add_heading("6. Seznam strategických dokumentů a plánů zadavatele", level=3)
    doc.add_paragraph(plan.strategicke_dokumenty or "–")

    # Oddíl 7 – Formát zprávy
    doc.add_heading("7. Formát zprávy o provedeném energetickém auditu", level=3)
    doc.add_paragraph(plan.format_zpravy or "–")

    # Oddíl 8 – Projednání
    doc.add_heading("8. Způsob projednání dílčích výstupů", level=3)
    doc.add_paragraph(plan.projednani_vystupu or "–")

    # Dodatky (§ 4 odst. 3)
    if plan.dodatky:
        doc.add_heading("Dodatky k plánu energetického auditu (§\u00a04 odst.\u00a03)", level=3)
        for i, add in enumerate(plan.dodatky, 1):
            doc.add_paragraph(f"Dodatek č.\u00a0{i}:")
            doc.add_paragraph(add)

    # Podpisový blok
    doc.add_heading("Podpisy", level=3)
    doc.add_paragraph(
        "Tento plán energetického auditu byl odsouhlasen oběma stranami. "
        "Podepsané vyhotovení je uloženo u zadavatele a zpracovatele."
    )
    tbl_podpis = doc.add_table(rows=3, cols=2)
    tbl_podpis.style = "Table Grid"
    headers_p = ["Zadavatel / zástupce zadavatele", "Energetický specialista"]
    for i, h in enumerate(headers_p):
        S.cell_text(tbl_podpis.rows[0].cells[i], h, bold=True, font_size=9)
        S.set_cell_bg(tbl_podpis.rows[0].cells[i], S.COLOR_TBL_HEAD_HEX)
    vals_p = [
        [plan.zadavatel_zastupce or "–", budova.zpracovatel_zastupce],
        ["Datum, podpis:", f"Č. oprávnění: {budova.cislo_opravneni or '–'}  |  Datum, podpis:"],
    ]
    for ri, row_vals in enumerate(vals_p):
        for ci, val in enumerate(row_vals):
            S.cell_text(tbl_podpis.rows[ri + 1].cells[ci], val, font_size=9)
    for row in tbl_podpis.rows:
        for cell in row.cells:
            S.set_col_width(cell, Cm(8))


# ── Pomocné funkce ────────────────────────────────────────────────────────────

def _nonempty(items: list) -> list[str]:
    return [x for x in items if x]


def _vloz_obrazek(doc: Document, buf, popisek: str = "",
                  sirka_cm: float = 14.0, cislo: int = 0) -> None:
    """
    Vloží obrázek z BytesIO (PNG/JPEG) do dokumentu s popiskem.
    buf může být None (funkce pak nic neudělá).
    Odstavec obrázku má automatické řádkování (předchází překrývání v různých šablonách).
    """
    if buf is None:
        return
    from io import BytesIO as _BytesIO
    from docx.enum.text import WD_LINE_SPACING
    data = buf if isinstance(buf, _BytesIO) else _BytesIO(buf)
    data.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(4)
    p.add_run().add_picture(data, width=Cm(sirka_cm))
    if popisek:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(10)
        run_cap = cap.add_run(
            f"Obr.\u00a0{cislo}\u00a0–\u00a0{popisek}" if cislo else popisek
        )
        run_cap.italic = True
        run_cap.font.size = Pt(8)


def _kv_tabulka(doc: Document, radky: list[tuple[str, str]]) -> None:
    """Vypíše tabulku klíč–hodnota (2 sloupce)."""
    tbl = doc.add_table(rows=len(radky), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(radky):
        S.cell_text(tbl.rows[i].cells[0], label, bold=True, font_size=9)
        S.cell_text(tbl.rows[i].cells[1], value, font_size=9)
        S.set_col_width(tbl.rows[i].cells[0], Cm(6.0))
        S.set_col_width(tbl.rows[i].cells[1], Cm(10.0))


def _popis_systemu(doc: Document, sys) -> None:
    """Vypíše parametry jednoho technického systému."""
    if sys.typ or sys.vykon_kw or sys.ucinnost_pct or sys.rok_instalace or sys.popis:
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        if sys.typ:
            _add_kv_row(tbl, "Typ / zdroj", sys.typ)
        if sys.vykon_kw:
            _add_kv_row(tbl, "Instalovaný výkon",
                        f"{sys.vykon_kw:,.1f}\u00a0kW".replace(",", "\u00a0"))
        if sys.ucinnost_pct:
            _add_kv_row(tbl, "Účinnost", f"{sys.ucinnost_pct:.1f}\u00a0%")
        if sys.rok_instalace:
            _add_kv_row(tbl, "Rok instalace", str(sys.rok_instalace))
        for row in tbl.rows:
            S.set_col_width(row.cells[0], Cm(5.5))
            S.set_col_width(row.cells[1], Cm(10.5))
        if sys.popis:
            doc.add_paragraph(sys.popis)
    else:
        doc.add_paragraph("Data k systému nebyla zadána.")


def _add_kv_row(table, klic: str, hodnota: str) -> None:
    row = table.add_row()
    S.cell_text(row.cells[0], klic, bold=True, font_size=9)
    S.cell_text(row.cells[1], hodnota, font_size=9)
