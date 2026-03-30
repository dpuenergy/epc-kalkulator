"""
Sdílené stavební bloky tabulek pro Word výstupy EP a EA.

Každá funkce přijme Document a potřebná data, přidá tabulku do dokumentu
a vrátí ji. Volající (ep_generator / ea_generator) ji může dále upravit.
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from epc_engine.building_class import KlasifikaceObaly
from epc_engine.emissions import EmiseBilance, EMISNI_FAKTORY
from epc_engine.economics import EkonomickeBilance, EkonomickeParametry
from epc_engine.models import (
    Budova, EnergyInputs, MeasureResult, Podklad,
    Konstrukce, BilancePouzitiEnergie, PenbData,
    HistorieRok, EnMSHodnoceni,
    EnergonositelEA, MKHKriterium, EnPIUkazatel,
)
from epc_engine.tepelna_technika import (
    TYP_POPISY, hodnoceni_splneni, urec_doporucena, upas_pasivni,
)

from . import _styles as S


# ── Helpers ───────────────────────────────────────────────────────────────────

def _header_row(table, texty: list[str], font_size: int = 9) -> None:
    """Zapíše první řádek tabulky: tučné záhlaví s modrým pozadím."""
    row = table.rows[0]
    for i, text in enumerate(texty):
        cell = row.cells[i]
        cell.text = text
        para = cell.paragraphs[0]
        for run in para.runs:
            run.bold = True
            run.font.name = S.FONT_BODY
            run.font.size = Pt(font_size)
        S.set_cell_bg(cell, S.COLOR_TBL_HEAD_HEX)


def _data_row(table, values: list[str], bold: bool = False,
              bg_hex: str | None = None, font_size: int = 9,
              aligns: list | None = None) -> None:
    """Přidá datový řádek do tabulky."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        para = cell.paragraphs[0]
        if aligns and i < len(aligns):
            para.alignment = aligns[i]
        for run in para.runs:
            run.bold = bold
            run.font.name = S.FONT_BODY
            run.font.size = Pt(font_size)
        if bg_hex:
            S.set_cell_bg(cell, bg_hex)


_R = WD_ALIGN_PARAGRAPH.RIGHT
_C = WD_ALIGN_PARAGRAPH.CENTER
_L = WD_ALIGN_PARAGRAPH.LEFT


# ── Tabulka energetické bilance ───────────────────────────────────────────────

def tabulka_energeticka_bilance(doc: Document, energie: EnergyInputs):
    """
    Energetická bilance stávajícího stavu.
    Sloupce: Energonosič | Spotřeba ÚT | Spotřeba TUV | Celkem | Cena | Náklady/rok
    """
    headers = [
        "Energonosič", "Spotřeba ÚT\n[MWh/rok]", "Spotřeba TUV\n[MWh/rok]",
        "Celkem\n[MWh/rok]", "Cena\n[Kč/MWh]", "Náklady\n[Kč/rok]",
    ]
    col_widths = [Cm(4.5), Cm(2.6), Cm(2.6), Cm(2.6), Cm(2.6), Cm(2.8)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    def _row(nosic, ut, tuv, cena):
        celkem = ut + tuv
        naklady = celkem * cena if cena > 0 else 0.0
        _data_row(table, [
            nosic,
            S.fmt_mwh(ut) if ut > 0 else "–",
            S.fmt_mwh(tuv) if tuv > 0 else "–",
            S.fmt_mwh(celkem) if celkem > 0 else "–",
            S.fmt_kc(cena) if cena > 0 else "–",
            S.fmt_kc(naklady) if naklady > 0 else "–",
        ], aligns=[_L, _R, _R, _R, _R, _R])

    if energie.pouzit_zp:
        _row("Zemní plyn", energie.zp_ut, energie.zp_tuv, energie.cena_zp)
    if energie.pouzit_teplo:
        _row("Teplo (CZT)", energie.teplo_ut, energie.teplo_tuv, energie.cena_teplo)
    if energie.ee > 0:
        _data_row(table, [
            "Elektrická energie",
            "–", "–",
            S.fmt_mwh(energie.ee),
            S.fmt_kc(energie.cena_ee),
            S.fmt_kc(energie.ee * energie.cena_ee),
        ], aligns=[_L, _R, _R, _R, _R, _R])
    if energie.voda > 0:
        _data_row(table, [
            "Voda a stočné", "–", "–",
            S.fmt_m3(energie.voda),
            S.fmt_kc(energie.cena_voda) + "/m³",
            S.fmt_kc(energie.voda * energie.cena_voda),
        ], aligns=[_L, _R, _R, _R, _R, _R])

    # Řádek CELKEM
    _data_row(table, [
        "CELKEM", "–", "–", "–", "–",
        S.fmt_kc(energie.celkove_naklady),
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
       aligns=[_L, _R, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka opatření ──────────────────────────────────────────────────────────

def tabulka_opatreni(doc: Document, aktivni: list[MeasureResult]):
    """
    Souhrnná tabulka aktivních opatření.
    Sloupce: ID | Název | Investice | Úspora teplo+ZP | Úspora EE | Úspora Kč/rok | Návratnost
    """
    headers = [
        "ID", "Název opatření",
        "Investice\n[Kč]", "Úspora tepla/ZP\n[MWh/rok]",
        "Úspora EE\n[MWh/rok]", "Úspora\n[Kč/rok]", "Prostá\nnávratnost",
    ]
    col_widths = [S.COL_ID, S.COL_NAZEV, S.COL_INVESTICE,
                  S.COL_USPORA_KW, S.COL_USPORA_KW, S.COL_USPORA_KC, S.COL_NAV]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for r in aktivni:
        _data_row(table, [
            r.id,
            r.nazev,
            S.fmt_kc(r.investice),
            S.fmt_mwh(r.uspora_teplo + r.uspora_zp) if (r.uspora_teplo + r.uspora_zp) > 0 else "–",
            S.fmt_mwh(r.uspora_ee) if r.uspora_ee != 0 else "–",
            S.fmt_kc(r.uspora_kc),
            S.fmt_nav(r.prosta_navratnost),
        ], aligns=[_C, _L, _R, _R, _R, _R, _R])

    # Řádek CELKEM
    celk_inv = sum(r.investice for r in aktivni)
    celk_tep = sum(r.uspora_teplo + r.uspora_zp for r in aktivni)
    celk_ee = sum(r.uspora_ee for r in aktivni)
    celk_kc = sum(r.uspora_kc for r in aktivni)
    celk_serv = sum(r.servisni_naklady for r in aktivni)
    net = celk_kc - celk_serv
    nav_celk = celk_inv / net if net > 0 else None
    _data_row(table, [
        "", "CELKEM",
        S.fmt_kc(celk_inv),
        S.fmt_mwh(celk_tep) if celk_tep > 0 else "–",
        S.fmt_mwh(celk_ee) if celk_ee != 0 else "–",
        S.fmt_kc(celk_kc),
        S.fmt_nav(nav_celk),
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
       aligns=[_C, _L, _R, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka ekonomická analýza ────────────────────────────────────────────────

def tabulka_ekonomika(doc: Document, aktivni: list[MeasureResult],
                      projekt_ek: EkonomickeBilance | None,
                      par: EkonomickeParametry | None):
    """
    Ekonomická analýza per opatření: Investice | NPV | IRR | Tsd | Prostá návratnost
    """
    horizont = par.horizont if par else 20
    headers = [
        "ID", "Název opatření",
        "Investice\n[Kč]", "Úspora\n[Kč/rok]",
        "NPV\n[Kč]", "IRR\n[%]",
        f"Tsd\n[roky]", "Prostá\nnávratnost",
    ]
    col_widths = [S.COL_ID, S.COL_NAZEV, S.COL_INVESTICE, S.COL_USPORA_KC,
                  S.COL_NPV, S.COL_IRR, S.COL_TSD, S.COL_NAV]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for r in aktivni:
        ek = r.ekonomika
        _data_row(table, [
            r.id,
            r.nazev,
            S.fmt_kc(r.investice),
            S.fmt_kc(r.uspora_kc),
            S.fmt_kc(ek.npv) if ek else "–",
            S.fmt_irr(ek.irr) if ek else "–",
            S.fmt_tsd(ek.tsd, horizont) if ek else "–",
            S.fmt_nav(r.prosta_navratnost),
        ], aligns=[_C, _L, _R, _R, _R, _R, _R, _R])

    # Řádek CELKEM
    celk_inv = sum(r.investice for r in aktivni)
    celk_kc = sum(r.uspora_kc for r in aktivni)
    _data_row(table, [
        "", "CELKEM",
        S.fmt_kc(celk_inv),
        S.fmt_kc(celk_kc),
        S.fmt_kc(projekt_ek.npv) if projekt_ek else "–",
        S.fmt_irr(projekt_ek.irr) if projekt_ek else "–",
        S.fmt_tsd(projekt_ek.tsd, horizont) if projekt_ek else "–",
        "–",
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
       aligns=[_C, _L, _R, _R, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka emisní bilance ────────────────────────────────────────────────────

def tabulka_emise(doc: Document,
                  pred: EmiseBilance | None,
                  po: EmiseBilance | None):
    """
    Emisní bilance před a po opatřeních.
    Sloupce: Látka | Jednotka | Stávající | Po opatřeních | Snížení | Snížení [%]
    """
    if pred is None:
        pred = EmiseBilance()
    if po is None:
        po = EmiseBilance()

    headers = ["Látka", "Jednotka", "Stávající stav", "Po opatřeních", "Snížení", "Snížení [%]"]
    col_widths = [Cm(2.8), Cm(2.0), Cm(3.2), Cm(3.2), Cm(2.4), Cm(2.4)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    def _em_row(latka, jednotka, v_pred, v_po):
        delta = v_pred - v_po
        pct = delta / v_pred * 100 if v_pred > 0 else 0.0
        _data_row(table, [
            latka, jednotka,
            f"{v_pred:,.1f}".replace(",", "\u00a0"),
            f"{v_po:,.1f}".replace(",", "\u00a0"),
            f"{delta:,.1f}".replace(",", "\u00a0"),
            f"{pct:.1f}\u00a0%",
        ], aligns=[_L, _C, _R, _R, _R, _R])

    _em_row("CO\u2082", "kg/rok", pred.co2_kg, po.co2_kg)
    _em_row("NO\u2093", "kg/rok", pred.nox_kg, po.nox_kg)
    _em_row("SO\u2082", "kg/rok", pred.so2_kg, po.so2_kg)
    _em_row("TZL", "kg/rok", pred.tzl_kg, po.tzl_kg)
    _em_row("EPS", "kg/rok", pred.eps_kg, po.eps_kg)

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka klasifikace obálky ────────────────────────────────────────────────

def tabulka_klasifikace(doc: Document, klas: KlasifikaceObaly):
    """
    Parametry klasifikace obálky budovy A–G.
    """
    headers = ["Parametr", "Hodnota"]
    col_widths = [Cm(7.0), Cm(5.0)]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _header_row(table, headers)

    _trida_popisy = {
        "A": "A – Mimořádně úsporná",
        "B": "B – Velmi úsporná",
        "C": "C – Úsporná (referenční)",
        "D": "D – Méně úsporná",
        "E": "E – Nehospodárná",
        "F": "F – Velmi nehospodárná",
        "G": "G – Mimořádně nehospodárná",
    }
    rows = [
        ("Uem – průměrný součinitel prostupu tepla obálky", f"{klas.uem:.3f} W/m²K"),
        ("Uem,N – referenční hodnota (ČSN 73 0540-2)", f"{klas.uem_n:.3f} W/m²K"),
        ("Poměr Uem / Uem,N", f"{klas.pomer:.3f}"),
        ("Klasifikační třída", _trida_popisy.get(klas.trida, klas.trida)),
    ]
    for label, value in rows:
        _data_row(table, [label, value], aligns=[_L, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka budov ─────────────────────────────────────────────────────────────

def tabulka_budovy(doc: Document, budovy: list[Budova]):
    """
    Seznam budov objektu s geometrickými parametry.
    """
    headers = ["Název", "Objem [m³]", "Podl. plocha [m²]", "Ochlaz. plocha [m²]", "A/V [m⁻¹]"]
    col_widths = [Cm(5.5), Cm(2.5), Cm(2.8), Cm(2.8), Cm(2.0)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for b in budovy:
        ft = f"{b.faktor_tvaru:.3f}" if b.objem_m3 > 0 else "–"
        _data_row(table, [
            b.nazev or "–",
            f"{b.objem_m3:,.0f}".replace(",", "\u00a0") if b.objem_m3 else "–",
            f"{b.podlahova_plocha_m2:,.0f}".replace(",", "\u00a0") if b.podlahova_plocha_m2 else "–",
            f"{b.ochlazovana_plocha_m2:,.0f}".replace(",", "\u00a0") if b.ochlazovana_plocha_m2 else "–",
            ft,
        ], aligns=[_L, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka podkladů ──────────────────────────────────────────────────────────

def tabulka_podklady(doc: Document, podklady: list[Podklad]):
    """
    Seznam vstupních podkladů s příznakem dostupnosti.
    """
    headers = ["#", "Podklad", "K dispozici"]
    col_widths = [Cm(0.8), Cm(12.0), Cm(2.0)]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _header_row(table, headers)

    for i, p in enumerate(podklady, 1):
        _data_row(table, [
            str(i),
            p.nazev,
            "✓" if p.k_dispozici else "✗",
        ], aligns=[_C, _L, _C])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka tepelně technických vlastností ────────────────────────────────────

def tabulka_konstrukce(doc: Document, konstrukce: list[Konstrukce]):
    """
    Tepelně technické vlastnosti obálky budovy.
    Pro každou konstrukci zobrazí hlavní řádek + podtabulku vrstev.
    """
    headers = ["Název", "Typ", "Plocha [m²]", "U [W/m²K]", "UN [W/m²K]", "Splnění"]
    col_widths = [Cm(4.5), Cm(3.5), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.0)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for k in konstrukce:
        u = k.u_effective
        un = k.un_value
        splneni = hodnoceni_splneni(u, un) if un > 0 else "–"
        splneni_bg = S.COLOR_TBL_TOTAL_HEX if splneni == "Vyhovuje" else None
        _data_row(table, [
            k.nazev or "–",
            TYP_POPISY.get(k.typ, k.typ),
            f"{k.plocha_m2:,.0f}".replace(",", "\u00a0") if k.plocha_m2 else "–",
            f"{u:.2f}" if u > 0 else "–",
            f"{un:.2f}" if un > 0 else "–",
            splneni,
        ], aligns=[_L, _L, _R, _R, _R, _C], bg_hex=splneni_bg)

        # Podtabulka vrstev (pokud existují a U není zadané přímo)
        if k.vrstvy and k.u_zadane is None:
            vrstvy_row = table.add_row()
            cell = vrstvy_row.cells[0]
            # Span přes všechny sloupce
            for i in range(1, len(headers)):
                cell.merge(vrstvy_row.cells[i])
            # Vnořená tabulka vrstev
            vtbl = cell.add_table(rows=1, cols=5)
            vtbl.style = "Table Grid"
            v_headers = ["Vrstva", "Název materiálu", "d [mm]", "λ [W/(m·K)]", "R [m²K/W]"]
            v_row = vtbl.rows[0]
            for i, txt in enumerate(v_headers):
                S.cell_text(v_row.cells[i], txt, bold=True, font_size=8)
                S.set_cell_bg(v_row.cells[i], "EBF5FB")
            for vi, v in enumerate(k.vrstvy, 1):
                r_v = v.tloustka_m / v.lambda_wm if v.lambda_wm > 0 else 0.0
                vr = vtbl.add_row()
                vals = [
                    str(vi),
                    v.nazev or "–",
                    f"{v.tloustka_m * 1000:.0f}" if v.tloustka_m > 0 else "–",
                    f"{v.lambda_wm:.3f}" if v.lambda_wm > 0 else "–",
                    f"{r_v:.3f}" if r_v > 0 else "–",
                ]
                aligns_v = [_C, _L, _R, _R, _R]
                for i, val in enumerate(vals):
                    S.cell_text(vr.cells[i], val, font_size=8,
                                align=aligns_v[i] if aligns_v else None)

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka bilance dle účelu (příloha č. 4 vyhl. 141/2021) ──────────────────

def tabulka_bilance_pouziti(doc: Document, bilance: BilancePouzitiEnergie):
    """
    Celková energetická bilance dle způsobu užití.
    Vzor: příloha č. 4 k vyhlášce č. 141/2021 Sb.
    """
    headers = ["Způsob užití energie", "MWh/rok", "tis. Kč/rok"]
    col_widths = [Cm(9.0), Cm(3.0), Cm(3.0)]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _header_row(table, headers)

    radky = [
        ("Vytápění", bilance.vytapeni_mwh, bilance.vytapeni_kc),
        ("Chlazení", bilance.chlazeni_mwh, bilance.chlazeni_kc),
        ("Příprava teplé vody", bilance.tuv_mwh, bilance.tuv_kc),
        ("Nucené větrání", bilance.vetrání_mwh, bilance.vetrání_kc),
        ("Osvětlení", bilance.osvetleni_mwh, bilance.osvetleni_kc),
        ("Technologie a spotřebiče", bilance.technologie_mwh, bilance.technologie_kc),
        ("PHM", bilance.phm_mwh, bilance.phm_kc),
    ]
    for nazev, mwh, kc in radky:
        _data_row(table, [
            nazev,
            S.fmt_mwh(mwh) if mwh > 0 else "–",
            f"{kc / 1000:.1f}" if kc > 0 else "–",
        ], aligns=[_L, _R, _R])

    _data_row(table, [
        "CELKEM",
        S.fmt_mwh(bilance.celkem_mwh),
        f"{bilance.celkem_kc / 1000:.1f}",
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX, aligns=[_L, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka PENB dat ──────────────────────────────────────────────────────────

def tabulka_penb(doc: Document, penb: PenbData):
    """
    Ukazatele energetické náročnosti budovy z PENB.
    """
    headers = ["Ukazatel", "Stávající stav", "Navrhovaný stav"]
    col_widths = [Cm(8.0), Cm(3.5), Cm(3.5)]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _header_row(table, headers)

    radky = [
        ("Klasifikační třída PENB", penb.trida_stavajici, penb.trida_navrhovy),
        ("Měrná potřeba tepla na vytápění [kWh/(m²·rok)]",
         f"{penb.merná_potreba_tepla:.0f}" if penb.merná_potreba_tepla else "–", "–"),
        ("Celková dodaná energie [kWh/(m²·rok)]",
         f"{penb.celkova_dodana_energie:.0f}" if penb.celkova_dodana_energie else "–", "–"),
        ("Primární energie z neobn. zdrojů [kWh/(m²·rok)]",
         f"{penb.primarni_neobnovitelna:.0f}" if penb.primarni_neobnovitelna else "–", "–"),
        ("Energeticky vztažná plocha [m²]",
         f"{penb.energeticka_vztazna_plocha:,.0f}".replace(",", "\u00a0")
         if penb.energeticka_vztazna_plocha else "–", "–"),
    ]
    for nazev, stav, navrhovy in radky:
        _data_row(table, [nazev, stav, navrhovy], aligns=[_L, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka měsíční spotřeby (B.2) ───────────────────────────────────────────

_MESICE = ["Leden", "Únor", "Březen", "Duben", "Květen", "Červen",
           "Červenec", "Srpen", "Září", "Říjen", "Listopad", "Prosinec"]


def tabulka_mesicni_spotreba(doc: Document, roky: list[HistorieRok]) -> None:
    """
    Měsíční spotřeba pro jeden energonosič (roky jsou předfiltrovány).
    Řádky = měsíce + CELKEM. Sloupce = Měsíc | rok1 [MWh] | rok2 [MWh] | …
    """
    if not roky:
        return
    rok_labels = [str(r.rok) if r.rok else "–" for r in roky]
    headers = ["Měsíc"] + [f"{rl}\n[MWh]" for rl in rok_labels]
    col_widths = [Cm(3.0)] + [Cm(2.5) for _ in roky]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for mi, mesic in enumerate(_MESICE):
        vals = [mesic]
        for r in roky:
            m = r.mesicni_mwh[mi] if r.mesicni_mwh and mi < len(r.mesicni_mwh) else 0.0
            vals.append(f"{m:.2f}" if m > 0 else "–")
        _data_row(table, vals, aligns=[_L] + [_R] * len(roky))

    # Řádek CELKEM (roční součty)
    celkem_vals = ["CELKEM"]
    for r in roky:
        celkem_vals.append(
            f"{r.spotreba_mwh:.1f}" if r.spotreba_mwh > 0 else "–"
        )
    _data_row(table, celkem_vals, bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
              aligns=[_L] + [_R] * len(roky))

    _set_col_widths(table, col_widths)


# ── Tabulka porovnání U s UN/Urec/Upas (C.2) ─────────────────────────────────

def tabulka_u_porovnani(doc: Document, konstrukce: list[Konstrukce]):
    """
    Porovnání U s normovými hodnotami UN / Urec / Upas dle ČSN 73 0540-2:2011.
    Sloupce: Název | Typ | Plocha | U | UN | Urec | Upas | Hodnocení
    """
    headers = [
        "Název konstrukce", "Typ", "Plocha\n[m²]",
        "U\n[W/m²K]", "UN\n[W/m²K]", "Urec\n[W/m²K]", "Upas\n[W/m²K]",
        "Hodnocení\nvůči UN",
    ]
    col_widths = [Cm(4.2), Cm(2.8), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.6), Cm(2.2)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    from epc_engine.tepelna_technika import un_pozadovana as _un_poz
    for k in konstrukce:
        u = k.u_effective
        un = k.un_value if k.un_value > 0 else _un_poz(k.typ)
        urec = urec_doporucena(k.typ)
        upas = upas_pasivni(k.typ)
        hodnoceni = hodnoceni_splneni(u, un) if un > 0 else "–"
        bg = S.COLOR_TBL_TOTAL_HEX if hodnoceni == "Vyhovuje" else None
        _data_row(table, [
            k.nazev or "–",
            TYP_POPISY.get(k.typ, k.typ),
            f"{k.plocha_m2:,.0f}".replace(",", "\u00a0") if k.plocha_m2 else "–",
            f"{u:.2f}" if u > 0 else "–",
            f"{un:.2f}" if un > 0 else "–",
            f"{urec:.2f}",
            f"{upas:.2f}",
            hodnoceni,
        ], aligns=[_L, _L, _R, _R, _R, _R, _R, _C], bg_hex=bg)

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka hodnocení EnMS (C.3 / B.5.2) ─────────────────────────────────────

_ENMS_HODNOCENI_POPIS = {
    0: "–",
    1: "1 – Nesplnění",
    2: "2 – Částečné splnění",
    3: "3 – Plné splnění",
}
_ENMS_BG = {
    1: "FADBD8",   # světle červená
    2: "FDEBD0",   # světle oranžová
    3: "D5F5E3",   # světle zelená
}


def tabulka_enms(doc: Document, enms: EnMSHodnoceni):
    """
    Hodnocení úrovně EnMS dle ČSN EN ISO 50001, škála 1–3.
    Sloupce: Oblast | Stávající stav | Hodnocení
    """
    headers = ["Oblast EnMS", "Popis stávajícího stavu", "Hodnocení (1–3)"]
    col_widths = [Cm(4.5), Cm(8.5), Cm(3.0)]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _header_row(table, headers)

    for oblast in enms.oblasti:
        bg = _ENMS_BG.get(oblast.hodnoceni)
        _data_row(table, [
            oblast.nazev,
            oblast.stav or "–",
            _ENMS_HODNOCENI_POPIS.get(oblast.hodnoceni, "–"),
        ], aligns=[_L, _L, _C], bg_hex=bg)

    # Řádek souhrnu: certifikace
    cert_text = "Certifikováno dle ČSN EN ISO 50001" if enms.certifikovan \
        else "Bez certifikace dle ČSN EN ISO 50001"
    _data_row(table, ["Certifikace EnMS", cert_text, ""],
              bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
              aligns=[_L, _L, _C])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka bilance energetických vstupů EA (příl. 3 vyhl. 140/2021) ─────────

# Emisní faktory CO₂ [t CO₂/MWh] – z emissions.py (kg/MWh → t/MWh děleno 1000)
# Zdroj: vyhl. 140/2021 příl. 8 + MPO 2022; EE=442 kg/MWh, ZP=200 kg/MWh, CZT=180 kg/MWh

def _co2_faktor(nazev: str) -> float:
    """Vrátí emisní faktor CO₂ [t/MWh] pro energonositel dle názvu."""
    n = nazev.lower()
    if "elektr" in n or " ee" in n:
        return EMISNI_FAKTORY["ee"]["co2"] / 1000
    if "teplo" in n or "czt" in n or "szt" in n or "dálkové" in n:
        return EMISNI_FAKTORY["teplo"]["co2"] / 1000
    # Výchozí: zemní plyn
    return EMISNI_FAKTORY["zp"]["co2"] / 1000


def _energie_pro_nositel(nazev: str, energie: EnergyInputs) -> tuple[float, float]:
    """Vrátí (MWh/rok, Kč/rok) pro daný energonositel z EnergyInputs."""
    n = nazev.lower()
    if "zemní plyn" in n or "zp" in n:
        mwh = energie.zp_total if energie.pouzit_zp else 0.0
        return mwh, mwh * energie.cena_zp
    if "elektr" in n or "ee" in n:
        return energie.ee, energie.ee * energie.cena_ee
    if "teplo" in n or "czt" in n or "szt" in n:
        mwh = energie.teplo_total if energie.pouzit_teplo else 0.0
        return mwh, mwh * energie.cena_teplo
    return 0.0, 0.0


def tabulka_bilance_vstupu_ea(
    doc: Document,
    energonositele: list[EnergonositelEA],
    energie: EnergyInputs,
):
    """
    Bilance energetických vstupů energetického hospodářství.
    Příloha č. 3 k vyhlášce č. 140/2021 Sb.
    Sloupce: Energonositel | Oblast užití | MWh/rok | tis. Kč/rok | t CO₂/rok
    Řádky jsou seskupeny dle kategorie (NOZE / OZE / Druhotné).
    """
    _KAT_PORADI = ["NOZE", "OZE", "Druhotné"]
    _KAT_NAZEV = {
        "NOZE": "Neobnovitelné zdroje energie",
        "OZE": "Obnovitelné zdroje energie",
        "Druhotné": "Druhotné zdroje energie",
    }

    headers = [
        "Energonositel / ucelená část",
        "Oblast užití",
        "MWh/rok",
        "tis. Kč/rok",
        "t CO₂/rok",
    ]
    col_widths = [Cm(5.5), Cm(2.8), Cm(2.2), Cm(2.2), Cm(2.2)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    total_mwh = 0.0
    total_kc = 0.0
    total_co2 = 0.0

    # Seskupit dle kategorie v požadovaném pořadí
    from collections import defaultdict
    skupiny: dict[str, list] = defaultdict(list)
    for en in energonositele:
        skupiny[en.kategorie].append(en)

    for kat in _KAT_PORADI:
        skupina = skupiny.get(kat, [])
        if not skupina:
            continue
        # Skupinový záhlaví řádek
        row = table.add_row()
        cell = row.cells[0]
        # Sloučit celý řádek
        for i in range(1, len(headers)):
            cell.merge(row.cells[i])
        p = cell.paragraphs[0]
        p.text = _KAT_NAZEV.get(kat, kat)
        p.runs[0].bold = True
        S.set_cell_bg(cell, "D9D9D9")

        for en in skupina:
            mwh, kc = _energie_pro_nositel(en.nazev, energie)
            co2 = mwh * _co2_faktor(en.nazev)
            total_mwh += mwh
            total_kc += kc
            total_co2 += co2
            _data_row(table, [
                en.nazev,
                en.oblast,
                S.fmt_mwh(mwh) if mwh > 0 else "–",
                f"{kc / 1000:.1f}" if kc > 0 else "–",
                f"{co2:.3f}" if co2 > 0 else "–",
            ], aligns=[_L, _C, _R, _R, _R])

    _data_row(table, [
        "CELKEM", "",
        S.fmt_mwh(total_mwh),
        f"{total_kc / 1000:.1f}",
        f"{total_co2:.3f}",
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
       aligns=[_L, _C, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka ukazatelů EnPI (příl. 5 vyhl. 140/2021) ──────────────────────────

def tabulka_enpi(doc: Document, enpi: list[EnPIUkazatel]):
    """
    Ukazatele energetické náročnosti (EnPI).
    Příloha č. 5 k vyhlášce č. 140/2021 Sb.
    Sloupce: Ukazatel | Jednotka | Popis stanovení | Typ | Výchozí hodnota | Navrhovaná hodnota | Změna
    """
    headers = [
        "Ukazatel EnPI / Spotřebič",
        "Jedn.",
        "Popis stanovení ukazatele",
        "Typ¹⁾",
        "Stávající",
        "Navrhovaný",
        "Změna",
    ]
    col_widths = [Cm(3.5), Cm(1.5), Cm(4.5), Cm(1.5), Cm(1.8), Cm(1.8), Cm(1.8)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for u in enpi:
        zmena = u.hodnota_navrhova - u.hodnota_stavajici if u.hodnota_stavajici else 0.0
        zmena_str = f"{zmena:+.2f}" if u.hodnota_stavajici > 0 else "–"
        typ = "S" if getattr(u, "je_stavajici", True) else "N"
        _data_row(table, [
            u.nazev or "–",
            u.jednotka or "–",
            getattr(u, "popis_stanoveni", "") or "Měřeno / odečteno z faktur",
            typ,
            f"{u.hodnota_stavajici:.2f}" if u.hodnota_stavajici > 0 else "–",
            f"{u.hodnota_navrhova:.2f}" if u.hodnota_navrhova > 0 else "–",
            zmena_str,
        ], aligns=[_L, _C, _L, _C, _R, _R, _R])

    # Poznámka pod tabulkou
    note = doc.add_paragraph()
    note.add_run("¹⁾ Typ: S = stávající ukazatel (již zaveden), N = navrhovaný nový ukazatel.").italic = True
    note.runs[0].font.size = Pt(8)

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka multikriteriálního hodnocení (příl. 9 vyhl. 140/2021) ─────────────

def _mkh_hodnota(r: MeasureResult, key: str) -> float:
    """Extrahuje hodnotu kritéria z MeasureResult."""
    if key == "npv":
        return r.ekonomika.npv / 1000 if r.ekonomika and r.ekonomika.npv is not None else 0.0
    if key == "td":
        if r.ekonomika and r.ekonomika.tsd is not None:
            return r.ekonomika.tsd
        return r.prosta_navratnost if r.prosta_navratnost is not None else 999.0
    if key == "mwh":
        return r.uspora_teplo + r.uspora_zp + r.uspora_ee
    if key == "kc":
        return r.uspora_kc / 1000
    return 0.0


def vypocitej_mkh(
    aktivni: list[MeasureResult],
    kriteria: list[MKHKriterium],
) -> list[tuple[MeasureResult, float, int]]:
    """
    Vypočítá multikriteriální hodnocení dle příl. 9 vyhl. 140/2021 Sb.
    (metoda váženého součtu s min-max normalizací).

    Vrátí: list of (MeasureResult, celkova_uzitnost, poradi)
    """
    if not aktivni or not kriteria:
        return [(r, 0.0, i + 1) for i, r in enumerate(aktivni)]

    aktivni_krit = [k for k in kriteria if k.vaha > 0]
    if not aktivni_krit:
        return [(r, 0.0, i + 1) for i, r in enumerate(aktivni)]

    # Extrakce hodnot per opatření per kritérium
    hodnoty: dict[str, list[float]] = {
        k.key: [_mkh_hodnota(r, k.key) for r in aktivni]
        for k in aktivni_krit
    }

    # Normalizace min-max
    def _norm(vals: list[float], typ: str) -> list[float]:
        vmin, vmax = min(vals), max(vals)
        if vmax == vmin:
            return [1.0] * len(vals)
        if typ == "max":
            return [(v - vmin) / (vmax - vmin) for v in vals]
        return [(vmax - v) / (vmax - vmin) for v in vals]

    normy: dict[str, list[float]] = {
        k.key: _norm(hodnoty[k.key], k.typ)
        for k in aktivni_krit
    }

    # Vážený součet
    total_vaha = sum(k.vaha for k in aktivni_krit)
    skore = []
    for i in range(len(aktivni)):
        s = sum(
            (k.vaha / total_vaha) * normy[k.key][i]
            for k in aktivni_krit
        )
        skore.append(s)

    # Seřazení a přiřazení pořadí
    poradi_idx = sorted(range(len(aktivni)), key=lambda i: skore[i], reverse=True)
    poradi_map = {idx: rank + 1 for rank, idx in enumerate(poradi_idx)}

    return [(aktivni[i], skore[i], poradi_map[i]) for i in range(len(aktivni))]


def tabulka_mkh(
    doc: Document,
    aktivni: list[MeasureResult],
    kriteria: list[MKHKriterium],
):
    """
    Multikriteriální hodnocení příležitostí ke snížení energetické náročnosti.
    Příloha č. 9 k vyhlášce č. 140/2021 Sb. (metoda váženého součtu).
    """
    aktivni_krit = [k for k in kriteria if k.vaha > 0]
    headers = (
        ["ID", "Název opatření"]
        + [f"{k.nazev}\n[{k.jednotka}]" for k in aktivni_krit]
        + ["Celková\nužitnost", "Pořadí"]
    )
    col_widths = (
        [S.COL_ID, S.COL_NAZEV]
        + [Cm(2.2) for _ in aktivni_krit]
        + [Cm(2.2), Cm(1.6)]
    )

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    # Záhlaví s typem a váhou
    typ_row = table.add_row()
    for i, cell in enumerate(typ_row.cells):
        if i < 2:
            cell.text = ""
        elif i - 2 < len(aktivni_krit):
            k = aktivni_krit[i - 2]
            cell.text = f"[{k.typ}] váha {k.vaha:.0f}"
            para = cell.paragraphs[0]
            para.alignment = _C
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = S.FONT_BODY
                run.italic = True
        S.set_cell_bg(cell, "EBF5FB")

    vysledky = vypocitej_mkh(aktivni, kriteria)

    for r, uzitnost, por in vysledky:
        hodnoty_krit = [
            f"{_mkh_hodnota(r, k.key):.1f}" for k in aktivni_krit
        ]
        bg = S.COLOR_TBL_TOTAL_HEX if por == 1 else None
        _data_row(table, [
            r.id, r.nazev,
            *hodnoty_krit,
            f"{uzitnost:.3f}",
            str(por),
        ], bg_hex=bg, aligns=[_C, _L] + [_R] * len(aktivni_krit) + [_R, _C])

    _set_col_widths(table, col_widths)
    return table


# ── Evidenční list příl. 1B (souhrn příležitostí) ────────────────────────────

def tabulka_evidencni_list_1b(
    doc: Document,
    aktivni: list[MeasureResult],
    kriteria: list[MKHKriterium],
):
    """
    Výstupy hodnocení příležitostí ke snížení energetické náročnosti.
    Část B přílohy č. 1 k vyhlášce č. 140/2021 Sb. (evidenční list).
    """
    headers = [
        "Ozn.", "Název příležitosti",
        "Úspora NOZE\n[MWh/rok]", "Úspora CO₂\n[t/rok]",
        "Náklady\n[tis. Kč]", "Úspora nákl.\n[tis. Kč/rok]",
        "NPV\n[tis. Kč]", "Td\n[roky]",
        "Priorita",
    ]
    col_widths = [
        S.COL_ID, S.COL_NAZEV,
        Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.2),
        Cm(2.0), Cm(1.6), Cm(1.8),
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    vysledky = vypocitej_mkh(aktivni, kriteria)
    # seřadit dle pořadí
    vysledky_sorted = sorted(vysledky, key=lambda x: x[2])

    for r, _, por in vysledky_sorted:
        mwh = r.uspora_teplo + r.uspora_zp + r.uspora_ee
        co2 = (r.uspora_zp * 0.200 + r.uspora_ee * 0.860
               + r.uspora_teplo * 0.100)
        npv_str = (
            S.fmt_kc(r.ekonomika.npv / 1000 * 1000)
            if r.ekonomika and r.ekonomika.npv is not None else "–"
        )
        td_str = (
            S.fmt_tsd(r.ekonomika.tsd if r.ekonomika else None, 20)
            if r.ekonomika else S.fmt_nav(r.prosta_navratnost)
        )
        bg = S.COLOR_TBL_TOTAL_HEX if por == 1 else None
        _data_row(table, [
            r.id, r.nazev,
            f"{mwh:.1f}" if mwh > 0 else "–",
            f"{co2:.1f}" if co2 > 0 else "–",
            f"{r.investice / 1000:.0f}" if r.investice > 0 else "–",
            f"{r.uspora_kc / 1000:.0f}" if r.uspora_kc > 0 else "–",
            npv_str,
            td_str,
            str(por),
        ], bg_hex=bg, aligns=[_C, _L, _R, _R, _R, _R, _R, _R, _C])

    _set_col_widths(table, col_widths)
    return table


# ── Výchozí roční energetická bilance ────────────────────────────────────────

def tabulka_bilance_rocni(
    doc: Document,
    energie: EnergyInputs,
) -> None:
    """
    Výchozí roční energetická bilance dle vzorového EA.
    Řádky: ZP-ÚT / ZP-TUV / EE / Teplo / CELKEM
    Sloupce: Energonositel | GJ/rok | MWh/rok | tis. Kč/rok | Podíl [%]
    """
    headers = ["Energonositel / ukazatel", "GJ/rok", "MWh/rok", "tis. Kč/rok", "Podíl [%]"]
    col_widths = [Cm(6.0), Cm(2.3), Cm(2.3), Cm(2.5), Cm(2.3)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    rows_data = []
    if energie.pouzit_zp:
        zp_celkem = energie.zp_ut + energie.zp_tuv
        if zp_celkem > 0:
            rows_data.append(("Zemní plyn – vytápění", energie.zp_ut,
                               energie.zp_ut * energie.cena_zp))
            rows_data.append(("Zemní plyn – příprava TUV", energie.zp_tuv,
                               energie.zp_tuv * energie.cena_zp))
    if energie.pouzit_teplo:
        teplo = energie.teplo_ut + energie.teplo_tuv
        if teplo > 0:
            rows_data.append(("Tepelná energie (CZT)", teplo,
                               teplo * energie.cena_teplo))
    if energie.ee > 0:
        rows_data.append(("Elektrická energie", energie.ee,
                           energie.ee * energie.cena_ee))

    total_mwh = sum(r[1] for r in rows_data)
    total_kc  = sum(r[2] for r in rows_data)

    for nazev, mwh, kc in rows_data:
        podil = mwh / total_mwh * 100 if total_mwh > 0 else 0.0
        _data_row(table, [
            nazev,
            f"{mwh * 3.6:.1f}",
            S.fmt_mwh(mwh),
            f"{kc / 1000:.1f}",
            f"{podil:.1f} %",
        ], aligns=[_L, _R, _R, _R, _R])

    _data_row(table, [
        "CELKEM",
        f"{total_mwh * 3.6:.1f}",
        S.fmt_mwh(total_mwh),
        f"{total_kc / 1000:.1f}",
        "100,0 %",
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
       aligns=[_L, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Cash flow tabulka per opatření ────────────────────────────────────────────

def tabulka_cash_flow(
    doc: Document,
    r: MeasureResult,
    parametry: "EkonomickeParametry | None" = None,
    horizont: int = 20,
) -> None:
    """
    Rok po roku cash flow tabulka pro jedno opatření.
    Sloupce: Rok | Náklady stáv. | Náklady po | Roční tok | Kumulace
    Formát odpovídá vzorovému EA (str. 83).
    """
    import math

    discount = (parametry.diskontni_sazba if parametry else 0.04)
    inflace  = (parametry.inflace_energie if parametry else 0.03)

    headers = [
        "Rok", "Prov. náklady stáv.\n[tis. Kč]",
        "Prov. náklady po\n[tis. Kč]",
        "Roční tok\n[tis. Kč]",
        "Roční tok\ndiskont. [tis. Kč]",
        "Kumulace\nnediskont. [tis. Kč]",
        "Kumulace\ndiskont. [tis. Kč]",
    ]
    col_widths = [Cm(1.4), Cm(2.5), Cm(2.3), Cm(2.0), Cm(2.2), Cm(2.3), Cm(2.2)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers, font_size=8)

    # Rok 0 – investice
    _data_row(table, [
        "0",
        "–", "–",
        f"{-r.investice / 1000:.0f}",
        f"{-r.investice / 1000:.0f}",
        f"{-r.investice / 1000:.0f}",
        f"{-r.investice / 1000:.0f}",
    ], font_size=8, aligns=[_C, _R, _R, _R, _R, _R, _R])

    import datetime
    rok_start = datetime.date.today().year + 1
    kumulace_nd = -r.investice / 1000
    kumulace_d  = -r.investice / 1000

    for t in range(1, horizont + 1):
        uspora_kc = r.uspora_kc * (1 + inflace) ** t
        tok_nd = uspora_kc / 1000
        tok_d  = tok_nd / (1 + discount) ** t
        kumulace_nd += tok_nd
        kumulace_d  += tok_d

        # Zvýraznit rok návratnosti
        navratnost_dosazena = (
            kumulace_nd >= 0 and (kumulace_nd - tok_nd) < 0
        )
        bg = S.COLOR_TBL_TOTAL_HEX if navratnost_dosazena else None
        _data_row(table, [
            str(rok_start + t - 1),
            "–",
            "–",
            f"{tok_nd:.0f}",
            f"{tok_d:.0f}",
            f"{kumulace_nd:.0f}",
            f"{kumulace_d:.0f}",
        ], font_size=8, bg_hex=bg, aligns=[_C, _R, _R, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Souhrnná tabulka variant opatření ────────────────────────────────────────

def tabulka_soucet_variant(
    doc: Document,
    aktivni: list[MeasureResult],
    parametry: "EkonomickeParametry | None" = None,
) -> None:
    """
    Souhrnná srovnávací tabulka všech navrhovaných příležitostí.
    Vzor: Tabulka 56 ze vzorového EA.
    """
    headers = [
        "Označení", "Název příležitosti",
        "Investice\n[tis. Kč]",
        "Úspora energie\n[MWh/rok]",
        "Úspora nákladů\n[tis. Kč/rok]",
        "NPV\n[tis. Kč]",
        "IRR\n[%]",
        "Prostá Ts\n[let]",
        "Reálná Tsd\n[let]",
    ]
    col_widths = [Cm(1.3), Cm(4.5), Cm(1.8), Cm(2.0), Cm(2.0), Cm(1.8), Cm(1.4), Cm(1.5), Cm(1.5)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    total_inv = 0.0
    total_mwh = 0.0
    total_kc  = 0.0

    for r in aktivni:
        npv_str = (
            S.fmt_kc(r.ekonomika.npv) if r.ekonomika and r.ekonomika.npv is not None
            else "–"
        )
        irr_str = (
            S.fmt_irr(r.ekonomika.irr) if r.ekonomika and r.ekonomika.irr is not None
            else "–"
        )
        ts_str  = S.fmt_nav(r.prosta_navratnost)
        tsd_str = (
            S.fmt_tsd(r.ekonomika.tsd if r.ekonomika else None, 20)
        )
        mwh = r.uspora_teplo + r.uspora_zp + r.uspora_ee
        total_inv += r.investice
        total_mwh += mwh
        total_kc  += r.uspora_kc

        _data_row(table, [
            r.id, r.nazev,
            f"{r.investice / 1000:.0f}",
            f"{mwh:.1f}",
            f"{r.uspora_kc / 1000:.0f}",
            npv_str, irr_str, ts_str, tsd_str,
        ], aligns=[_C, _L, _R, _R, _R, _R, _R, _R, _R])

    # CELKEM řádek
    _data_row(table, [
        "CELKEM", "",
        f"{total_inv / 1000:.0f}",
        f"{total_mwh:.1f}",
        f"{total_kc / 1000:.0f}",
        "–", "–", "–", "–",
    ], bold=True, bg_hex=S.COLOR_TBL_TOTAL_HEX,
       aligns=[_C, _L, _R, _R, _R, _R, _R, _R, _R])

    _set_col_widths(table, col_widths)
    return table


# ── Tabulka účinnosti zdrojů energie (§ 6 vyhl. 406/2000 Sb.) ────────────────

def tabulka_ucinnost_zdroju(
    doc: Document,
    technicke_systemy: "TechnickeSystemy | None",
) -> None:
    """
    Vyhodnocení účinnosti zdrojů energie dle § 6 vyhl. 406/2000 Sb.
    Porovnání s referenční účinností dle vyhl. 441/2012 Sb.
    """
    from epc_engine.models import TechnickeSystemy

    headers = [
        "Zdroj / systém", "Typ", "Jmenovitý výkon\n[kW]",
        "Rok instalace",
        "Účinnost\ndle výrobce [%]",
        "Ref. účinnost\n[%]",
        "Hodnocení",
    ]
    col_widths = [Cm(3.5), Cm(3.0), Cm(2.0), Cm(1.8), Cm(2.0), Cm(2.0), Cm(2.2)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    # Referenční účinnost pro kotle dle vyhl. 441/2012 Sb.
    def _ref_ucinnost(vykon_kw: float, rok: int) -> str:
        if vykon_kw <= 0:
            return "–"
        # Nová zařízení po 2004: kotle do 400 kW → 86 %, nad 400 kW → 88 %
        if rok >= 2004:
            return "86 %" if vykon_kw <= 400 else "88 %"
        return "84 %"

    def _hodnoceni(ucinnost: float, ref_str: str) -> str:
        try:
            ref = float(ref_str.replace(" %", ""))
            return "Vyhovuje" if ucinnost >= ref else "Nevyhovuje"
        except Exception:
            return "–"

    sys = technicke_systemy
    if sys is None:
        sys = TechnickeSystemy()

    for nazev_s, syst in [
        ("Vytápění (zdroj tepla)", sys.vytapeni),
        ("Příprava TUV", sys.tuv),
        ("Větrání a klimatizace", sys.vzt),
        ("Osvětlení", sys.osvetleni),
    ]:
        if not syst or syst.vykon_kw <= 0:
            continue
        ref = _ref_ucinnost(syst.vykon_kw, syst.rok_instalace)
        hod = _hodnoceni(syst.ucinnost_pct, ref) if syst.ucinnost_pct > 0 else "–"
        bg = S.COLOR_TBL_TOTAL_HEX if hod == "Vyhovuje" else None
        _data_row(table, [
            nazev_s,
            syst.typ or "–",
            f"{syst.vykon_kw:.1f}" if syst.vykon_kw > 0 else "–",
            str(syst.rok_instalace) if syst.rok_instalace > 0 else "–",
            f"{syst.ucinnost_pct:.1f} %" if syst.ucinnost_pct > 0 else "–",
            ref,
            hod,
        ], bg_hex=bg, aligns=[_L, _L, _R, _C, _R, _R, _C])

    _set_col_widths(table, col_widths)
    return table


# ── Interní helper pro šířky sloupců ─────────────────────────────────────────

def _set_col_widths(table, widths: list) -> None:
    """Nastaví šířky sloupců v tabulce."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                S.set_col_width(cell, widths[i])
