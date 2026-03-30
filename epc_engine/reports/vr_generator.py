"""
Generátor technické přílohy k zadávací dokumentaci výběrového řízení (VŘ) EPC.

Výstup: Word dokument (.docx) se třemi přílohami:
  Příloha č. 1 – Minimální technické požadavky na opatření
  Příloha č. 2 – Referenční a okrajové podmínky
  Příloha č. 3 – Agregovaný položkový rozpočet
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from epc_engine.models import BuildingInfo, ProjectResult
from epc_engine.op_descriptions import OP_INFO, OP_TECHNICKE_PODMINKY, OP_ROZPOCET_POLOZKY

from . import _styles as S
from ._tables import _header_row, _data_row

_TEMPLATE = __import__("pathlib").Path(__file__).parent / "templates" / "ep_template.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _nadpis(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        run.font.name = S.FONT_BODY
        run.font.color.rgb = S.COLOR_H1 if level == 1 else S.COLOR_H2


def _odstavec(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.style.font.name = S.FONT_BODY
    if bold:
        for run in p.runs:
            run.bold = True


def _tabulka_3sl(
    doc: Document,
    zahlavi: list[str],
    radky: list[tuple[str, ...]],
    sirky: list[float] | None = None,
) -> None:
    """Vytvoří tabulku se záhlavím a datovými řádky."""
    if sirky is None:
        sirky = [5.5, 7.0, 4.0]
    table = doc.add_table(rows=1, cols=len(zahlavi))
    table.style = "Table Grid"
    for i, (col, w) in enumerate(zip(table.columns, sirky)):
        col.width = Cm(w)
    _header_row(table, zahlavi)
    aligns = [WD_ALIGN_PARAGRAPH.LEFT] * len(zahlavi)
    aligns[-1] = WD_ALIGN_PARAGRAPH.LEFT
    for r in radky:
        _data_row(table, list(r), aligns=aligns)
    doc.add_paragraph()


def _op_section_row(table, text: str, n_cols: int) -> None:
    """Přidá řádek záhlaví sekce přes celou šířku tabulky (zvýrazněný)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    row = table.add_row()
    cell = row.cells[0]
    # Merge all cells in the row
    if n_cols > 1:
        cell = cell.merge(row.cells[n_cols - 1])
    cell.text = text
    para = cell.paragraphs[0]
    for run in para.runs:
        run.bold = True
        run.font.name = S.FONT_BODY
        run.font.size = Pt(9)
    S.set_cell_bg(cell, "D5E8F0")


# ── Veřejné API ───────────────────────────────────────────────────────────────

def generuj_prilohu_vr(
    budova: BuildingInfo,
    result: ProjectResult,
    aktivni_op: list[str],
    ss: dict,
) -> BytesIO:
    """
    Sestaví Word dokument – technické přílohy k zadávací dokumentaci VŘ.

    Parametry:
        budova      – BuildingInfo ze session state
        result      – ProjectResult z build_project().vypocitej()
        aktivni_op  – seznam aktivních OP ID (např. ["OP1a", "OP2", "OP10"])
        ss          – st.session_state jako dict (pro přístup k cenám, spotřebám apod.)
    """
    doc = Document(_TEMPLATE)

    # ── Titulní strana ────────────────────────────────────────────────────────
    nazev_obj = budova.objekt_nazev or "Objekt"
    adresa_obj = budova.objekt_adresa or ""
    zadavatel = ss.get("zadavatel_nazev", "")
    datum = ss.get("datum", "")
    zpracovatel = ss.get("zpracovatel_zastupce", "")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TECHNICKÉ PŘÍLOHY K ZADÁVACÍ DOKUMENTACI")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = S.FONT_BODY
    run.font.color.rgb = S.COLOR_H1

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("VÝBĚROVÉHO ŘÍZENÍ – EPC PROJEKT")
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = S.FONT_BODY
    run2.font.color.rgb = S.COLOR_H2

    doc.add_paragraph()
    for label, value in [
        ("Objekt:", f"{nazev_obj}  {adresa_obj}".strip()),
        ("Zadavatel:", zadavatel),
        ("Zpracoval:", zpracovatel),
        ("Datum:", datum),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.name = S.FONT_BODY
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.name = S.FONT_BODY
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ── PŘÍLOHA Č. 1 – Minimální technické požadavky ──────────────────────────
    _nadpis(doc, "Příloha č. 1 – Minimální technické požadavky na navrhovaná opatření")
    _odstavec(
        doc,
        "Zhotovitel je povinen splnit nebo překročit všechny parametry uvedené v této příloze. "
        "Splnění požadavků bude doloženo technickými listy výrobců, atesty, certifikáty nebo "
        "protokoly z měření předloženými při přejímce díla.",
    )
    doc.add_paragraph()

    for op_id in aktivni_op:
        info = OP_INFO.get(op_id, {})
        podminky = OP_TECHNICKE_PODMINKY.get(op_id, [])
        _nadpis(doc, f"{op_id} – {info.get('title', op_id)}", level=2)
        if podminky:
            # 5 sloupců: Parametr | Min. požadavek | Norma | Splňuje (Ano/Ne) | Dodavatel nabízí
            table5 = doc.add_table(rows=1, cols=5)
            table5.style = "Table Grid"
            for col, w in zip(table5.columns, [4.0, 5.5, 2.5, 1.8, 2.7]):
                col.width = Cm(w)
            _header_row(table5, [
                "Parametr", "Minimální požadavek", "Norma / reference",
                "Splňuje\n(Ano / Ne)", "Dodavatel nabízí",
            ])
            for p_row in podminky:
                _data_row(
                    table5,
                    [p_row[0], p_row[1], p_row[2], "", ""],
                    aligns=[
                        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.LEFT,
                    ],
                )
            doc.add_paragraph()
        else:
            _odstavec(doc, "Technické požadavky stanoví projektová dokumentace.")

    doc.add_page_break()

    # ── PŘÍLOHA Č. 2 – Referenční a okrajové podmínky ────────────────────────
    _nadpis(doc, "Příloha č. 2 – Referenční a okrajové podmínky")
    _odstavec(
        doc,
        "Tato příloha definuje výchozí (referenční) stav a okrajové podmínky platné "
        "pro měření a ověřování (M&V) dosažených úspor v průběhu EPC kontraktu. "
        "Zhotovitel ručí za dosažení smluvní výše úspor vztažené k níže uvedenému referenčnímu stavu.",
    )
    doc.add_paragraph()

    _nadpis(doc, "2.1  Referenční spotřeby energie", level=2)
    en = result.energie
    _uspora_zp  = result.celkova_uspora_zp
    _uspora_t   = result.celkova_uspora_teplo
    _uspora_ee  = result.celkova_uspora_ee
    _cena_zp    = ss.get("cena_zp", 1800.0)
    _cena_t     = ss.get("cena_teplo", 2200.0)
    _cena_ee    = ss.get("cena_ee", 4500.0)

    _ref_radky = [
        ("Zemní plyn – vytápění", f"{en.zp_ut:,.1f}".replace(",", "\u00a0"), "MWh/rok",
         f"{_cena_zp:,.0f}".replace(",", "\u00a0"), f"{en.zp_ut * _cena_zp / 1000:,.0f}".replace(",", "\u00a0")),
        ("Zemní plyn – TUV", f"{en.zp_tuv:,.1f}".replace(",", "\u00a0"), "MWh/rok",
         f"{_cena_zp:,.0f}".replace(",", "\u00a0"), f"{en.zp_tuv * _cena_zp / 1000:,.0f}".replace(",", "\u00a0")),
        ("Teplo (CZT)", f"{en.teplo_total:,.1f}".replace(",", "\u00a0"), "MWh/rok",
         f"{_cena_t:,.0f}".replace(",", "\u00a0"), f"{en.teplo_total * _cena_t / 1000:,.0f}".replace(",", "\u00a0")),
        ("Elektrická energie", f"{en.ee:,.1f}".replace(",", "\u00a0"), "MWh/rok",
         f"{_cena_ee:,.0f}".replace(",", "\u00a0"), f"{en.ee * _cena_ee / 1000:,.0f}".replace(",", "\u00a0")),
        ("CELKEM", "", "",
         "", f"{en.celkove_naklady:,.0f}".replace(",", "\u00a0")),
    ]
    table_ref = doc.add_table(rows=1, cols=5)
    table_ref.style = "Table Grid"
    for col, w in zip(table_ref.columns, [4.5, 2.5, 2.0, 2.5, 3.0]):
        col.width = Cm(w)
    _header_row(table_ref, ["Energonosič", "Spotřeba", "Jednotka", "Cena [Kč/MWh]", "Roční náklady [Kč]"])
    for r in _ref_radky:
        _data_row(
            table_ref, list(r),
            bold=(r[0] == "CELKEM"),
            bg_hex=(S.COLOR_TBL_TOTAL_HEX if r[0] == "CELKEM" else None),
            aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        )
    doc.add_paragraph()

    _nadpis(doc, "2.2  Klimatické a provozní podmínky", level=2)
    lokalita = ss.get("lokalita_projekt", "Praha (Karlov)")
    theta_i  = ss.get("theta_i", 21.0)
    theta_e  = ss.get("theta_e", -13.0)
    _klim_radky = [
        ("Referenční lokalita", lokalita, ""),
        ("Vnitřní výpočtová teplota", f"{theta_i:.1f}", "°C"),
        ("Venkovní výpočtová teplota", f"{theta_e:.1f}", "°C"),
        ("Referenční rok", "Dle průměru spotřeb za poslední 3 roky před realizací", ""),
        ("Normalizace klimatu", "Spotřeba na vytápění se normalizuje denostupňovou metodou dle ČSN EN ISO 15316-4-1", ""),
    ]
    _tabulka_3sl(
        doc,
        ["Parametr", "Hodnota", "Jednotka"],
        _klim_radky,
        sirky=[6.5, 7.0, 3.0],
    )

    _nadpis(doc, "2.3  Plánované úspory po realizaci opatření", level=2)
    _uspora_kc   = result.celkova_uspora_kc
    _uspora_pct  = _uspora_kc / en.celkove_naklady * 100 if en.celkove_naklady > 0 else 0.0
    _uspora_mwh  = _uspora_zp + _uspora_t + _uspora_ee
    _usp_radky = [
        ("Roční úspora nákladů na energie",
         f"{_uspora_kc:,.0f}".replace(",", "\u00a0"), "Kč/rok"),
        ("Roční úspora energie – celkem",
         f"{_uspora_mwh:,.1f}".replace(",", "\u00a0"), "MWh/rok"),
        ("Relativní úspora",
         f"{_uspora_pct:.1f}", "%"),
    ]
    _tabulka_3sl(
        doc,
        ["Ukazatel", "Hodnota", "Jednotka"],
        _usp_radky,
        sirky=[7.0, 4.0, 2.5],
    )

    doc.add_page_break()

    # ── PŘÍLOHA Č. 3 – Agregovaný položkový rozpočet ─────────────────────────
    _nadpis(doc, "Příloha č. 3 – Agregovaný položkový rozpočet")
    _odstavec(
        doc,
        "Soupis agregovaných položek vychází z energetického posudku. "
        "Uchazeči ocení každou položku na základě jimi navrhovaného technického řešení "
        "splňujícího nebo překračujícího minimální technické požadavky (Příloha č. 1). "
        "Celková smluvní cena musí pokrýt veškeré práce, materiály, zkoušky a dokumentaci "
        "nezbytné k realizaci opatření.",
    )
    doc.add_paragraph()

    _celkem_inv = result.celkova_investice

    # Tabulka s 6 sloupci: Č. | Popis položky | MJ | Množství | Jedn. cena [Kč] | Celkem bez DPH [Kč]
    # Čísla a ceny vyplňuje uchazeč. Orientační celkový náklad OP je uveden v souhrnném řádku.
    _ROZP_SIRKY = [0.8, 7.2, 1.4, 1.8, 2.4, 2.9]
    _ROZP_ZAHLAVI = ["Č.", "Popis položky", "MJ", "Množství", "Jedn. cena\n[Kč]", "Celkem\nbez DPH [Kč]"]
    _ROZP_ALIGNS = [
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
    ]

    table_rozp = doc.add_table(rows=1, cols=6)
    table_rozp.style = "Table Grid"
    for col, w in zip(table_rozp.columns, _ROZP_SIRKY):
        col.width = Cm(w)
    _header_row(table_rozp, _ROZP_ZAHLAVI)

    for mr in result.vysledky:
        if not mr.aktivni:
            continue
        op_nazev = OP_INFO.get(mr.id, {}).get("title", mr.id)
        # Záhlaví sekce opatření
        _op_section_row(table_rozp, f"{mr.id}  –  {op_nazev}", n_cols=6)
        # Položky
        polozky = OP_ROZPOCET_POLOZKY.get(mr.id, [])
        for i, (popis, mj) in enumerate(polozky, start=1):
            _data_row(
                table_rozp,
                [str(i), popis, mj, "", "", ""],
                aligns=_ROZP_ALIGNS,
            )
        # Souhrnný řádek opatření (orientační hodnota z kalkulace)
        _pct = mr.investice / _celkem_inv * 100 if _celkem_inv > 0 else 0.0
        _data_row(
            table_rozp,
            ["", f"Mezisoučet {mr.id}  (orientační hodnota dle EP/EA)",
             "", "", "",
             f"{mr.investice:,.0f}".replace(",", "\u00a0")],
            bold=True,
            bg_hex=S.COLOR_TBL_TOTAL_HEX,
            aligns=_ROZP_ALIGNS,
        )

    # Celkový součet
    _data_row(
        table_rozp,
        ["", "CELKEM – všechna opatření", "", "", "",
         f"{_celkem_inv:,.0f}".replace(",", "\u00a0")],
        bold=True,
        bg_hex="D5E8F0",
        aligns=_ROZP_ALIGNS,
    )
    doc.add_paragraph()

    _odstavec(
        doc,
        "Poznámka: Položkový rozpočet slouží pro orientaci zadavatele a ověření souladu nabídkové "
        "ceny s rozsahem díla. Zadavatel nepožaduje detailní slepý výkaz výměr; uchazeč může "
        "vnitřní strukturu položek přizpůsobit svému řešení za podmínky splnění minimálních "
        "technických požadavků a dosažení smluvní výše úspor.",
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
