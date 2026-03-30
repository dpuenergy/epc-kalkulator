"""
Generátor Energetického posudku (EP) ve formátu .docx.

Struktura dokumentu dle vyhlášky č. 141/2021 Sb.:
  Titulní strana
  A. Identifikační údaje (A.1–A.4)
  B. Popis stávajícího stavu (B.1–B.5)
  C. Vyhodnocení stávajícího stavu (C.1–C.5)
  D. Navrhovaná opatření a stanovisko specialisty (D.1–D.4)
  E. Okrajové podmínky a závěr
  F. Přílohy (placeholders)
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from epc_engine.models import BuildingInfo, ProjectResult
from epc_engine.op_descriptions import OP_INFO
from epc_engine.tepelna_technika import TYP_POPISY, vypocitej_uem_z_konstrukci

from . import _styles as S
from ._tables import (
    tabulka_energeticka_bilance,
    tabulka_opatreni,
    tabulka_ekonomika,
    tabulka_emise,
    tabulka_klasifikace,
    tabulka_budovy,
    tabulka_podklady,
    tabulka_konstrukce,
    tabulka_bilance_pouziti,
    tabulka_penb,
    tabulka_mesicni_spotreba,
    tabulka_u_porovnani,
    tabulka_enms,
)

_TEMPLATE = Path(__file__).parent / "templates" / "ep_template.docx"

# Legislativa a normy – pevný seznam citovaných předpisů
_LEGISLATIVA = [
    "Zákon č. 406/2000 Sb., o hospodaření energií, ve znění pozdějších předpisů",
    "Vyhláška č. 141/2021 Sb., o energetickém posudku a o údajích vedených v Systému monitoringu spotřeby energie",
    "Vyhláška č. 264/2020 Sb., o energetické náročnosti budov (průkaz energetické náročnosti budovy – PENB)",
    "Vyhláška č. 193/2007 Sb., kterou se stanoví podrobnosti účinnosti užití energie při rozvodu tepelné energie a vnitřním rozvodu tepelné energie a chladu",
    "ČSN EN ISO 6946:2017 – Stavební prvky a stavební konstrukce – Tepelný odpor a součinitel prostupu tepla",
    "ČSN 73 0540-2:2011 – Tepelná ochrana budov – část 2: Požadavky",
    "ČSN EN ISO 13789:2009 – Tepelné chování budov – Součinitel tepelné ztráty prostupem tepla a větráním",
    "ČSN EN ISO 50001:2018 – Systémy managementu hospodaření s energií – Požadavky s návodem k použití",
]


# ── Veřejné API ───────────────────────────────────────────────────────────────

def generuj_ep(budova: BuildingInfo, result: ProjectResult) -> BytesIO:
    """Sestaví EP dokument dle vyhl. 141/2021 Sb. a vrátí BytesIO ke stažení."""
    doc = Document(_TEMPLATE)

    _titulni_strana(doc, budova)
    _obsah(doc)

    _sekce_a(doc, budova, result)
    _sekce_b(doc, budova, result)
    _sekce_c(doc, budova, result)
    _sekce_d(doc, budova, result)
    _sekce_e(doc, budova)
    _sekce_f(doc, budova)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Obsah ─────────────────────────────────────────────────────────────────────

_OBSAH_POLOZKY = [
    # (úroveň odsazení, text)
    (0, "A. IDENTIFIKAČNÍ ÚDAJE"),
    (1, "A.1 Identifikace zadavatele, objektu a zpracovatele"),
    (1, "A.2 Účel zpracování"),
    (1, "A.3 Vstupní podklady, legislativa a technické normy"),
    (1, "A.4 Ceny energií"),
    (0, "B. POPIS STÁVAJÍCÍHO STAVU"),
    (1, "B.1 Popis stavby a provozu"),
    (1, "B.2 Historie spotřeby energie"),
    (2, "B.2.1 Klimatická data a klimatická korekce"),
    (1, "B.3 Tepelně technické vlastnosti obálky budovy"),
    (2, "B.3.1 Vymezení systémové hranice budovy"),
    (2, "B.3.2 Fasády a obvodové stěny"),
    (2, "B.3.3 Podlahy"),
    (2, "B.3.4 Střechy a stropy pod nevytápěným podkrovím"),
    (2, "B.3.5 Okna, dveře a prosklené výplně otvorů"),
    (1, "B.4 Technické systémy budovy"),
    (2, "B.4.1 Vytápění a příprava teplé vody"),
    (2, "B.4.2 Větrání a klimatizace"),
    (2, "B.4.3 Osvětlení"),
    (2, "B.4.4 Měření, regulace a monitoring (MaR)"),
    (1, "B.5 Energetická bilance stávajícího stavu"),
    (2, "B.5.1 Celková energetická bilance dle přílohy č. 4 vyhl. č. 141/2021 Sb."),
    (2, "B.5.2 Systém managementu hospodaření s energií (EnMS)"),
    (0, "C. VYHODNOCENÍ STÁVAJÍCÍHO STAVU"),
    (1, "C.1 Principy výpočtu a výpočtový model"),
    (1, "C.2 Vyhodnocení požadavků na součinitel prostupu tepla"),
    (1, "C.3 Hodnocení úrovně EnMS dle ČSN EN ISO 50001"),
    (1, "C.4 Celková energetická bilance"),
    (1, "C.5 Emisní bilance"),
    (0, "D. NAVRHOVANÁ OPATŘENÍ A STANOVISKO ENERGETICKÉHO SPECIALISTY"),
    (1, "D.1 Navrhovaná úsporná opatření"),
    (1, "D.2 Ekonomická analýza"),
    (1, "D.3 Ekologické hodnocení"),
    (1, "D.4 Primární energie z neobnovitelných zdrojů"),
    (0, "E. OKRAJOVÉ PODMÍNKY A ZÁVĚR"),
    (1, "E.1 Okrajové podmínky výpočtu"),
    (1, "E.2 Závěr a doporučení"),
    (1, "E.3 Prohlášení a podpis energetického specialisty"),
    (0, "F. PŘÍLOHY"),
]


def _obsah(doc: Document) -> None:
    """
    Vloží automatický obsah (TOC field) + statický přehled sekcí.
    TOC se automaticky aktualizuje při otevření v MS Word
    (Ctrl+A → F9, nebo pravé tlačítko → Aktualizovat pole).
    """
    doc.add_heading("OBSAH", level=2)

    # Word TOC field – aktualizuje se při otevření dokumentu
    para = doc.add_paragraph()
    run = para.add_run()
    # Begin field
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fc_begin)
    # Field instruction: TOC levels 1-3, hyperlinks, hide page numbers if 0
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    # Separate (placeholder content shown before first update)
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fc_sep)
    # Static fallback text visible before Word updates the field
    fallback = OxmlElement("w:t")
    fallback.text = "[Obsah se aktualizuje po otevření v MS Word: Ctrl+A → F9]"
    run._r.append(fallback)
    # End field
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run._r.append(fc_end)

    # Statický přehled sekcí – viditelný ihned (záloha a přehled struktury)
    doc.add_paragraph()
    doc.add_paragraph("Přehled sekcí dokumentu:")
    for uroven, text in _OBSAH_POLOZKY:
        indent = "\u00a0\u00a0\u00a0\u00a0" * uroven  # nezlomitelné mezery pro odsazení
        p = doc.add_paragraph(f"{indent}{text}")
        # Tučné pro sekce nejvyšší úrovně
        if uroven == 0:
            for run in p.runs:
                run.bold = True
        from docx.shared import Pt
        for run in p.runs:
            run.font.size = Pt(9 if uroven > 0 else 10)

    doc.add_page_break()


# ── Titulní strana ────────────────────────────────────────────────────────────

def _titulni_strana(doc: Document, budova: BuildingInfo,
                    typ: str = "ENERGETICKÝ POSUDEK") -> None:
    """Titulní strana: nadpis, název objektu, identifikační tabulka."""
    h = doc.add_heading(typ, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if budova.objekt_nazev:
        h2 = doc.add_heading(budova.objekt_nazev, level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    rows = [
        ("Zakázka", budova.nazev_zakazky or "–"),
        ("Zadavatel", budova.zadavatel_nazev or "–"),
        ("Adresa objektu", budova.objekt_adresa or "–"),
        ("Zpracovatel", f"DPU ENERGY s.r.o.\n{budova.zpracovatel_zastupce}" if budova.zpracovatel_zastupce else "DPU ENERGY s.r.o."),
        ("Číslo oprávnění", budova.cislo_opravneni or "–"),
        ("Datum vypracování", budova.datum or "–"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        S.cell_text(c0, label, bold=True, font_size=10)
        S.cell_text(c1, value, font_size=10)
        S.set_col_width(c0, Cm(5.0))
        S.set_col_width(c1, Cm(11.0))

    doc.add_page_break()


# ── A. Identifikační údaje ────────────────────────────────────────────────────

def _sekce_a(doc: Document, budova: BuildingInfo, result: ProjectResult) -> None:
    doc.add_heading("A. IDENTIFIKAČNÍ ÚDAJE", level=1)

    # A.1 – Identifikace zadavatele, objektu a zpracovatele
    doc.add_heading("A.1 Identifikace", level=2)

    doc.add_heading("Zadavatel", level=3)
    for line in _nonempty([
        budova.zadavatel_nazev,
        budova.zadavatel_adresa,
        f"IČ: {budova.zadavatel_ico}" if budova.zadavatel_ico else None,
        f"Kontakt: {budova.zadavatel_kontakt}" if budova.zadavatel_kontakt else None,
        f"Tel.: {budova.zadavatel_telefon}" if budova.zadavatel_telefon else None,
        f"E-mail: {budova.zadavatel_email}" if budova.zadavatel_email else None,
    ]):
        doc.add_paragraph(line)

    doc.add_heading("Identifikace objektu", level=3)
    for line in _nonempty([
        budova.objekt_nazev,
        budova.objekt_adresa,
        f"Katastrální území: {budova.objekt_ku}" if budova.objekt_ku else None,
        f"Parcelní číslo: {budova.objekt_parcelni_cislo}" if budova.objekt_parcelni_cislo else None,
        f"Evidenční číslo (ENEX): {budova.evidencni_cislo}" if budova.evidencni_cislo else None,
    ]):
        doc.add_paragraph(line)

    doc.add_heading("Zpracovatel posudku", level=3)
    for line in _nonempty([
        "DPU ENERGY s.r.o.",
        f"Odpovědný zpracovatel: {budova.zpracovatel_zastupce}" if budova.zpracovatel_zastupce else None,
        f"Číslo oprávnění energetického specialisty: {budova.cislo_opravneni}" if budova.cislo_opravneni else None,
        f"Datum vypracování: {budova.datum}" if budova.datum else None,
    ]):
        doc.add_paragraph(line)

    dph_text = "bez DPH" if budova.ceny_bez_dph else "včetně DPH"
    doc.add_paragraph(f"Veškeré ceny a náklady jsou uváděny {dph_text}.")

    # A.2 – Účel zpracování
    doc.add_heading("A.2 Účel zpracování", level=2)
    if budova.ucel_ep:
        doc.add_paragraph(f"Účel zpracování: {budova.ucel_ep}")
    if budova.predmet_analyzy:
        doc.add_paragraph(budova.predmet_analyzy)
    else:
        doc.add_paragraph(
            "Energetický posudek byl zpracován na základě objednávky zadavatele v souladu "
            "s § 9a zákona č. 406/2000 Sb. o hospodaření energií a vyhláškou č. 141/2021 Sb."
        )

    # A.3 – Vstupní podklady, legislativa a ČSN normy
    doc.add_heading("A.3 Vstupní podklady", level=2)

    if budova.podklady:
        doc.add_heading("Seznam podkladů", level=3)
        tabulka_podklady(doc, budova.podklady)
    if budova.poznamka_podklady:
        doc.add_paragraph(budova.poznamka_podklady)

    doc.add_heading("Použitá legislativa a technické normy", level=3)
    for predpis in _LEGISLATIVA:
        doc.add_paragraph(predpis, style="List Bullet")

    # A.4 – Ceny energií a DPH
    doc.add_heading("A.4 Ceny energií", level=2)
    energie = result.energie
    ceny = []
    if energie.pouzit_zp and energie.cena_zp > 0:
        ceny.append(f"Zemní plyn: {S.fmt_kc(energie.cena_zp)} Kč/MWh")
    if energie.pouzit_teplo and energie.cena_teplo > 0:
        ceny.append(f"Teplo (CZT): {S.fmt_kc(energie.cena_teplo)} Kč/MWh")
    if energie.ee > 0 and energie.cena_ee > 0:
        ceny.append(f"Elektrická energie: {S.fmt_kc(energie.cena_ee)} Kč/MWh")
    if energie.voda > 0 and energie.cena_voda > 0:
        ceny.append(f"Voda a stočné: {S.fmt_kc(energie.cena_voda)} Kč/m³")
    if ceny:
        doc.add_paragraph(f"Použité ceny energií ({dph_text}):")
        for c in ceny:
            doc.add_paragraph(c, style="List Bullet")
    else:
        doc.add_paragraph(f"Ceny energií nebyly zadány (výsledky jsou {dph_text}).")


# ── B. Popis stávajícího stavu ────────────────────────────────────────────────

def _sekce_b(doc: Document, budova: BuildingInfo, result: ProjectResult) -> None:
    doc.add_heading("B. POPIS STÁVAJÍCÍHO STAVU", level=1)

    _b1_popis_stavby(doc, budova)
    _b2_historie_spotreby(doc, budova)
    _b3_tepelna_obalka(doc, budova)
    _b4_technicke_systemy(doc, budova)
    _b5_energeticka_bilance(doc, budova, result)


def _b1_popis_stavby(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("B.1 Popis stavby a provozu", level=2)

    for line in _nonempty([
        f"Druh činnosti: {budova.druh_cinnosti}" if budova.druh_cinnosti else None,
        f"Počet zaměstnanců / uživatelů: {budova.pocet_zamestnancu}" if budova.pocet_zamestnancu else None,
        f"Provozní režim: {budova.provozni_rezim}" if budova.provozni_rezim else None,
    ]):
        doc.add_paragraph(line)

    budovy_s_daty = [b for b in budova.budovy if b.objem_m3 > 0 or b.nazev]
    if budovy_s_daty:
        doc.add_heading("Soupis budov a geometrické parametry", level=3)
        tabulka_budovy(doc, budovy_s_daty)

    if budova.prostory:
        doc.add_heading("Využití prostor", level=3)
        for p in budova.prostory:
            text = p.nazev
            if p.ucel:
                text += f" – {p.ucel}"
            if p.provoz:
                text += f" ({p.provoz})"
            doc.add_paragraph(text, style="List Bullet")


def _b2_historie_spotreby(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("B.2 Historie spotřeby energie", level=2)

    # B.2.1 – Klimatická data
    kd = budova.klimaticka_data
    doc.add_heading("B.2.1 Klimatická data a klimatická korekce", level=3)
    if kd:
        for line in _nonempty([
            f"Lokalita: {kd.lokalita}" if kd.lokalita else None,
            f"Normované denostupně D = {kd.stupnodni_normovane:,.0f}\u00a0°C·dny/rok".replace(",", "\u00a0")
            if kd.stupnodni_normovane else None,
            f"Vnitřní výpočtová teplota ti = {kd.teplota_vnitrni:.0f}\u00a0°C",
            f"Venkovní výpočtová teplota te = {kd.teplota_exterieru:.0f}\u00a0°C",
        ]):
            doc.add_paragraph(line)
        if kd.stupnodni_normovane > 0:
            doc.add_paragraph(
                "Spotřeba tepla na vytápění je klimaticky korigována na normované denostupně. "
                f"Korigovaná spotřeba = skutečná spotřeba × D_norm / D_skut, kde "
                f"D_norm\u00a0=\u00a0{kd.stupnodni_normovane:,.0f}\u00a0°C·dny/rok "
                f"(lokalita: {kd.lokalita or '–'}).".replace(",", "\u00a0")
            )
    else:
        doc.add_paragraph("Klimatická data lokality nebyla zadána.")

    # Tabulky měsíční spotřeby per energonosič
    if budova.historie_spotreby:
        doc.add_heading("Přehled měsíčních spotřeb", level=3)
        # Seskupit dle energonosiče
        nosice: dict[str, list] = {}
        for h in budova.historie_spotreby:
            nosice.setdefault(h.energonosic or "Neurčeno", []).append(h)

        for nosic, roky in nosice.items():
            doc.add_paragraph(f"{nosic}:", style="Heading 4" if "Heading 4" in
                              [s.name for s in doc.styles] else "Normal")
            tabulka_mesicni_spotreba(doc, roky)

        # Souhrnná roční tabulka
        doc.add_heading("Roční souhrn spotřeb a nákladů", level=3)
        d_norm = kd.stupnodni_normovane if kd and kd.stupnodni_normovane else 0.0
        has_korekce = d_norm > 0
        headers = ["Rok", "Energonosič", "Spotřeba [MWh/rok]",
                   "Náklady [tis. Kč/rok]", "Skut. D [°C·d]"]
        if has_korekce:
            headers += ["D_norm/D_skut", "Korig. spotřeba [MWh]"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hr = tbl.rows[0]
        for i, h in enumerate(headers):
            S.cell_text(hr.cells[i], h, bold=True, font_size=9)
            S.set_cell_bg(hr.cells[i], S.COLOR_TBL_HEAD_HEX)
        for h in budova.historie_spotreby:
            row = tbl.add_row()
            d_skut = h.stupnodni or 0.0
            if has_korekce and d_skut > 0:
                pomer = d_norm / d_skut
                korig = h.spotreba_mwh * pomer if h.spotreba_mwh else 0.0
                pomer_str = f"{pomer:.3f}"
                korig_str = f"{korig:,.1f}".replace(",", "\u00a0")
            else:
                pomer_str = "–"
                korig_str = "–"
            vals = [
                str(h.rok) if h.rok else "–",
                h.energonosic or "–",
                f"{h.spotreba_mwh:,.1f}".replace(",", "\u00a0") if h.spotreba_mwh else "–",
                f"{h.naklady_kc / 1000:.1f}" if h.naklady_kc else "–",
                f"{d_skut:,.0f}".replace(",", "\u00a0") if d_skut else "–",
            ]
            if has_korekce:
                vals += [pomer_str, korig_str]
            for i, val in enumerate(vals):
                S.cell_text(row.cells[i], val, font_size=9)
        if has_korekce:
            col_widths = [Cm(1.2), Cm(3.5), Cm(2.8), Cm(2.8), Cm(2.2), Cm(2.2), Cm(2.8)]
        else:
            col_widths = [Cm(1.5), Cm(4.0), Cm(3.5), Cm(3.5), Cm(3.0)]
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    S.set_col_width(cell, col_widths[i])
    else:
        doc.add_paragraph(
            "Data historické spotřeby za poslední 3 roky nebyla zadána. "
            "Spotřeba je hodnocena na základě referenčních hodnot."
        )


def _b3_tepelna_obalka(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("B.3 Tepelně technické vlastnosti obálky budovy", level=2)

    # B.3.1 – Systémová hranice
    doc.add_heading("B.3.1 Vymezení systémové hranice budovy", level=3)
    doc.add_paragraph(
        "Systémová hranice budovy je definována dle ČSN EN ISO 13789:2009 jako obálka ohraničující "
        "vytápěný (klimatizovaný) prostor od venkovního prostředí nebo od nevytápěných prostor. "
        "Do výpočtu jsou zahrnuty veškeré konstrukce, jimiž dochází k tepelným ztrátám z vytápěného prostoru."
    )

    if not budova.konstrukce:
        doc.add_paragraph("Tepelně technické vlastnosti stavebních konstrukcí nebyly zadány.")
        return

    # Seskupení konstrukcí dle typu pro subsekce B.3.2–B.3.5
    SKUPINY = [
        ("stena",   "B.3.2 Fasády a obvodové stěny"),
        ("podlaha", "B.3.3 Podlahy"),
        ("strecha", "B.3.4 Střechy a stropy pod nevytápěným podkrovím"),
        ("okno",    "B.3.5 Okna a prosklené výplně otvorů"),
        ("dvere",   "B.3.5 Vstupní dveře"),
    ]
    # Deduplikovat nadpis B.3.5 pokud jsou obě skupiny
    _visited_b35 = False
    for typ, nadpis in SKUPINY:
        skupina = [k for k in budova.konstrukce if k.typ == typ]
        if not skupina:
            continue
        if typ in ("okno", "dvere"):
            if not _visited_b35:
                doc.add_heading("B.3.5 Okna, dveře a prosklené výplně otvorů", level=3)
                _visited_b35 = True
        else:
            doc.add_heading(nadpis, level=3)
        tabulka_konstrukce(doc, skupina)

    # Fotografie obálky budovy
    foto_budova = [f for f in budova.fotografie if f.sekce == "budova"]
    if foto_budova:
        doc.add_heading("Fotodokumentace obálky budovy", level=3)
        for foto in foto_budova:
            _vloz_obrazek(doc, foto.data, foto.popisek, sirka_cm=foto.sirka_cm)


def _b4_technicke_systemy(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("B.4 Technické systémy budovy", level=2)

    ts = budova.technicke_systemy

    # B.4.1 – Vytápění a TUV
    doc.add_heading("B.4.1 Vytápění a příprava teplé vody", level=3)
    for nazev, sys in [("Vytápění", ts.vytapeni), ("Příprava teplé vody (TUV)", ts.tuv)]:
        doc.add_heading(nazev, level=4)
        _popis_systemu(doc, sys)

    # B.4.2 – VZT
    doc.add_heading("B.4.2 Větrání a klimatizace", level=3)
    _popis_systemu(doc, ts.vzt)

    # B.4.3 – Osvětlení
    doc.add_heading("B.4.3 Osvětlení", level=3)
    _popis_systemu(doc, ts.osvetleni)

    # B.4.4 – MaR
    if ts.mereni_ridici:
        doc.add_heading("B.4.4 Měření, regulace a monitoring (MaR)", level=3)
        doc.add_paragraph(ts.mereni_ridici)

    # Fotografie technických systémů
    foto_tech = [f for f in budova.fotografie if f.sekce == "technika"]
    if foto_tech:
        doc.add_heading("Fotodokumentace technických systémů", level=3)
        for foto in foto_tech:
            _vloz_obrazek(doc, foto.data, foto.popisek, sirka_cm=foto.sirka_cm)


def _popis_systemu(doc: Document, sys) -> None:
    """Pomocná funkce: vypíše parametry jednoho technického systému."""
    if sys.typ or sys.vykon_kw or sys.ucinnost_pct or sys.rok_instalace or sys.popis:
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        if sys.typ:
            _add_kv_row(tbl, "Typ / zdroj", sys.typ)
        if sys.vykon_kw:
            _add_kv_row(tbl, "Instalovaný výkon",
                        f"{sys.vykon_kw:,.1f}\u00a0kW".replace(",", "\u00a0"))
        if sys.ucinnost_pct:
            _add_kv_row(tbl, "Účinnost", f"{sys.ucinnost_pct:.1f}\u00a0%")
        if sys.rok_instalace:
            _add_kv_row(tbl, "Rok instalace", str(sys.rok_instalace))
        for row in tbl.rows:
            S.set_col_width(row.cells[0], Cm(5.5))
            S.set_col_width(row.cells[1], Cm(10.5))
        if sys.popis:
            doc.add_paragraph(sys.popis)
    else:
        doc.add_paragraph("Data k systému nebyla zadána.")


def _b5_energeticka_bilance(doc: Document, budova: BuildingInfo,
                             result: ProjectResult) -> None:
    doc.add_heading("B.5 Energetická bilance stávajícího stavu", level=2)

    # B.5.1 – Celková energetická bilance (Příloha č. 4)
    doc.add_heading(
        "B.5.1 Celková energetická bilance dle přílohy č. 4 vyhl. č. 141/2021 Sb.", level=3)
    tabulka_energeticka_bilance(doc, result.energie)

    if budova.bilance_pouziti:
        doc.add_heading("Rozpad spotřeby dle způsobu užití energie", level=4)
        tabulka_bilance_pouziti(doc, budova.bilance_pouziti)

    # B.5.2 – EnMS
    doc.add_heading(
        "B.5.2 Systém managementu hospodaření s energií (EnMS)", level=3)
    if budova.enms:
        doc.add_paragraph(
            "Hodnocení úrovně systému managementu hospodaření s energií dle požadavků "
            "ČSN EN ISO 50001:2018, škála: 1 – nesplnění, 2 – částečné splnění, 3 – plné splnění."
        )
        tabulka_enms(doc, budova.enms)
        if budova.enms.komentar:
            doc.add_paragraph(f"Komentář: {budova.enms.komentar}")
    else:
        doc.add_paragraph(
            "Hodnocení EnMS nebylo provedeno. Organizace nemá zaveden formální systém "
            "managementu hospodaření s energií dle ČSN EN ISO 50001:2018."
        )


# ── C. Vyhodnocení stávajícího stavu ─────────────────────────────────────────

def _sekce_c(doc: Document, budova: BuildingInfo, result: ProjectResult) -> None:
    doc.add_heading("C. VYHODNOCENÍ STÁVAJÍCÍHO STAVU", level=1)

    _c1_principy(doc)
    _c2_tepelne_technicke(doc, budova, result)
    _c3_enms(doc, budova)
    _c4_bilance(doc, result)
    _c5_emise(doc, result)


def _c1_principy(doc: Document) -> None:
    doc.add_heading("C.1 Principy výpočtu a výpočtový model", level=2)
    doc.add_paragraph(
        "Výpočet tepelných ztrát prostupem tepla a energetické bilance je proveden na základě "
        "naměřených a zadaných dat o spotřebě energií, parametrech stavebních konstrukcí a "
        "technických systémech budovy."
    )
    doc.add_paragraph(
        "Součinitele prostupu tepla U jednotlivých stavebních konstrukcí jsou vypočteny dle "
        "ČSN EN ISO 6946:2017 z tepelných odporů vrstev konstrukce (R = d/λ) a povrchových odporů "
        "(Rsi, Rse). Pro okna, dveře a jiné výplně otvorů jsou použity hodnoty U udávané výrobcem."
    )
    doc.add_paragraph(
        "Požadované (UN), doporučené (Urec) a doporučené pro pasivní budovy (Upas) hodnoty "
        "součinitele prostupu tepla jsou převzaty z ČSN 73 0540-2:2011, tabulka 3, "
        "pro budovy s převažující vnitřní teplotou 18–22 °C."
    )


def _c2_tepelne_technicke(doc: Document, budova: BuildingInfo,
                           result: ProjectResult) -> None:
    doc.add_heading(
        "C.2 Vyhodnocení požadavků na součinitel prostupu tepla", level=2)
    doc.add_paragraph(
        "Hodnocení dle ČSN 73 0540-2:2011. "
        "UN = požadovaná hodnota, Urec = doporučená hodnota, Upas = doporučená pro pasivní budovy."
    )

    if budova.konstrukce:
        tabulka_u_porovnani(doc, budova.konstrukce)

        # Uem – průměrný součinitel prostupu tepla obálky
        uem = vypocitej_uem_z_konstrukci(budova.konstrukce)
        if uem is not None:
            doc.add_paragraph(
                f"Průměrný součinitel prostupu tepla obálky budovy: "
                f"Uem\u00a0=\u00a0{uem:.3f}\u00a0W/(m²K)."
            )
    else:
        doc.add_paragraph("Tepelně technické vlastnosti stavebních konstrukcí nebyly zadány.")

    # Klasifikace obálky
    klas = result.klasifikace_pred
    if klas:
        doc.add_heading("Klasifikace obálky budovy", level=3)
        tabulka_klasifikace(doc, klas)
        doc.add_paragraph(
            f"Budova je zařazena do energetické třídy {klas.trida} "
            f"(Uem\u00a0=\u00a0{klas.uem:.3f}\u00a0W/m²K, "
            f"Uem,N\u00a0=\u00a0{klas.uem_n:.3f}\u00a0W/m²K)."
        )


def _c3_enms(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading(
        "C.3 Hodnocení úrovně EnMS dle ČSN EN ISO 50001", level=2)
    if budova.enms:
        doc.add_paragraph(
            "Celkové vyhodnocení stavu systému managementu hospodaření s energií (EnMS) "
            "dle požadavků ČSN EN ISO 50001:2018."
        )
        tabulka_enms(doc, budova.enms)
        # Průměrné hodnocení
        hodnoty = [o.hodnoceni for o in budova.enms.oblasti if o.hodnoceni > 0]
        if hodnoty:
            prumer = sum(hodnoty) / len(hodnoty)
            doc.add_paragraph(
                f"Průměrné hodnocení EnMS: {prumer:.1f} / 3,0. "
                + ("Organizace má zaveden a certifikován EnMS dle ISO 50001."
                   if budova.enms.certifikovan
                   else "Organizace nemá certifikovaný EnMS dle ISO 50001.")
            )
    else:
        doc.add_paragraph(
            "Hodnocení EnMS nebylo provedeno. Doporučujeme zhodnotit zavedení systému "
            "managementu hospodaření s energií dle ČSN EN ISO 50001:2018 jako nástroje "
            "systematického řízení spotřeby energie."
        )


def _c4_bilance(doc: Document, result: ProjectResult) -> None:
    doc.add_heading("C.4 Celková energetická bilance", level=2)
    doc.add_paragraph(
        "Níže uvedená tabulka shrnuje celkovou spotřebu a náklady na energie stávajícího stavu."
    )
    tabulka_energeticka_bilance(doc, result.energie)


def _c5_emise(doc: Document, result: ProjectResult) -> None:
    doc.add_heading("C.5 Emisní bilance", level=2)
    doc.add_paragraph(
        "Emisní bilance je zpracována pro stávající stav a stav po realizaci navrhovaných opatření. "
        "Emise jsou vypočteny z měrných emisních faktorů pro jednotlivé energonosiče. "
        "Ekvivalentní prašná substance (EPS) je stanovena dle vztahu: "
        "EPS\u00a0=\u00a01,0·TZL\u00a0+\u00a00,88·NOₓ\u00a0+\u00a00,54·SO₂\u00a0+\u00a00,64·NH₃."
    )
    tabulka_emise(doc, result.emise_pred, result.emise_po)

    pred = result.emise_pred
    po = result.emise_po
    if pred and po:
        delta_co2 = pred.co2_kg - po.co2_kg
        if delta_co2 > 0:
            doc.add_paragraph(
                f"Realizací navržených opatření dojde ke snížení emisí CO\u2082 "
                f"o {delta_co2:,.0f}\u00a0kg/rok.".replace(",", "\u00a0")
            )


# ── D. Navrhovaná opatření a stanovisko specialisty ──────────────────────────

def _sekce_d(doc: Document, budova: BuildingInfo, result: ProjectResult) -> None:
    doc.add_heading(
        "D. NAVRHOVANÁ OPATŘENÍ A STANOVISKO ENERGETICKÉHO SPECIALISTY", level=1)

    _d1_opatreni(doc, result)
    _d2_ekonomika(doc, result)
    _d3_ekologie(doc, result)
    _d4_primarni_energie(doc, budova)


def _d1_opatreni(doc: Document, result: ProjectResult) -> None:
    doc.add_heading("D.1 Navrhovaná úsporná opatření", level=2)

    aktivni = result.aktivni
    if not aktivni:
        doc.add_paragraph("Nebyla identifikována žádná vhodná úsporná opatření.")
        return

    tabulka_opatreni(doc, aktivni)
    doc.add_paragraph()

    for r in aktivni:
        doc.add_heading(f"{r.id} – {r.nazev}", level=3)
        info = OP_INFO.get(r.id, {})
        popis = info.get("popis", "")
        if popis:
            doc.add_paragraph(popis)
        doc.add_paragraph(
            f"Odhadovaná investice: {S.fmt_kc(r.investice)}.  "
            f"Roční úspora: {S.fmt_kc(r.uspora_kc)}.  "
            f"Prostá návratnost: {S.fmt_nav(r.prosta_navratnost)}."
        )
        dotace = info.get("dotace", "")
        if dotace:
            doc.add_paragraph(f"Dotační podmínky: {dotace}")


def _d2_ekonomika(doc: Document, result: ProjectResult) -> None:
    doc.add_heading("D.2 Ekonomická analýza", level=2)

    par = result.ekonomika_parametry
    if par:
        doc.add_paragraph(
            f"Ekonomická analýza je provedena na hodnotící horizont {par.horizont}\u00a0let, "
            f"diskontní sazba {par.diskontni_sazba * 100:.1f}\u00a0%, "
            f"předpokládaná meziroční inflace cen energií {par.inflace_energie * 100:.1f}\u00a0%/rok."
        )

    aktivni = result.aktivni
    if aktivni:
        tabulka_ekonomika(doc, aktivni, result.ekonomika_projekt, par)

        celk_inv = result.celkova_investice
        celk_kc = result.celkova_uspora_kc
        nav = result.prosta_navratnost_celkem
        ek = result.ekonomika_projekt
        doc.add_paragraph(
            f"Celková investice do navrhovaných opatření: {S.fmt_kc(celk_inv)}. "
            f"Celková roční úspora: {S.fmt_kc(celk_kc)}. "
            f"Prostá návratnost: {S.fmt_nav(nav)}."
            + (f" NPV projektu: {S.fmt_kc(ek.npv)}." if ek else "")
        )
    else:
        doc.add_paragraph("Žádná aktivní opatření nebyla definována.")


def _d3_ekologie(doc: Document, result: ProjectResult) -> None:
    doc.add_heading("D.3 Ekologické hodnocení", level=2)
    doc.add_paragraph(
        "Ekologické hodnocení vychází z emisní bilance zpracované v části C.5. "
        "Níže jsou uvedeny dopady realizace navrhovaných opatření na produkci emisí."
    )
    pred = result.emise_pred
    po = result.emise_po
    tabulka_emise(doc, pred, po)


def _d4_primarni_energie(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading(
        "D.4 Primární energie z neobnovitelných zdrojů", level=2)

    penb = budova.penb
    if penb is None:
        doc.add_paragraph(
            "Data průkazu energetické náročnosti budovy (PENB) nebyla zadána. "
            "PENB je zpracováván samostatně v certifikovaném softwaru (Energie+, TechCon apod.) "
            "a je přikládán jako příloha tohoto posudku."
        )
    else:
        tabulka_penb(doc, penb)
        if penb.poznamka:
            doc.add_paragraph(f"Poznámka k PENB: {penb.poznamka}")

    doc.add_paragraph(
        "Průkaz energetické náročnosti budovy (PENB) dle vyhlášky č. 264/2020 Sb. "
        "je vypracován certifikovaným energetickým specialistou a je přílohou tohoto energetického posudku."
    )


# ── E. Okrajové podmínky a závěr ─────────────────────────────────────────────

def _sekce_e(doc: Document, budova: BuildingInfo) -> None:
    doc.add_heading("E. OKRAJOVÉ PODMÍNKY A ZÁVĚR", level=1)

    doc.add_heading("E.1 Okrajové podmínky výpočtu", level=2)
    if budova.okrajove_podminky:
        doc.add_paragraph(budova.okrajove_podminky)
    doc.add_paragraph(
        "Výpočty jsou zpracovány v souladu s ČSN EN ISO 6946:2017 (tepelný odpor a součinitel "
        "prostupu tepla stavebních prvků), ČSN 73 0540-2:2011 (tepelná ochrana budov – Požadavky), "
        "ČSN EN ISO 13789:2009 (tepelné chování budov, systémová hranice) a metodikou "
        "Státní energetické inspekce."
    )

    doc.add_heading("E.2 Závěr a doporučení", level=2)
    doc.add_paragraph(
        "Na základě provedené analýzy energetické situace objektu jsou v části D navrhována "
        "úsporná opatření vedoucí ke snížení spotřeby energie a provozních nákladů."
    )
    doc.add_paragraph(
        "Doporučujeme zahájit přípravné práce pro realizaci navrhovaných opatření a ověřit "
        "aktuální podmínky dotačních programů (Nová zelená úsporám Light, OPŽP, MPO EFEKT, OP TAK)."
    )

    doc.add_heading("E.3 Prohlášení a podpis energetického specialisty", level=2)
    doc.add_paragraph(
        "Energetický posudek byl zpracován v souladu s požadavky zákona č. 406/2000 Sb. "
        "a vyhlášky č. 141/2021 Sb. osobou s oprávněním energetického specialisty. "
        "Energetický specialista prohlašuje, že údaje v posudku jsou pravdivé a úplné."
    )
    doc.add_paragraph()
    sig_rows = [
        ("Energetický specialista:", budova.zpracovatel_zastupce or ""),
        ("Číslo oprávnění:", budova.cislo_opravneni or ""),
        ("Datum:", budova.datum or ""),
        ("Podpis:", ""),
    ]
    tbl = doc.add_table(rows=len(sig_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(sig_rows):
        S.cell_text(tbl.rows[i].cells[0], label, bold=True, font_size=10)
        S.cell_text(tbl.rows[i].cells[1], value, font_size=10)
        S.set_col_width(tbl.rows[i].cells[0], Cm(5.0))
        S.set_col_width(tbl.rows[i].cells[1], Cm(11.0))


# ── F. Přílohy ────────────────────────────────────────────────────────────────

def _sekce_f(doc: Document, budova: BuildingInfo) -> None:
    doc.add_page_break()
    doc.add_heading("F. PŘÍLOHY", level=1)

    prilohy = [
        "F.1 Kopie oprávnění energetického specialisty",
        "F.2 Průkaz energetické náročnosti budovy (PENB) dle vyhlášky č. 264/2020 Sb.",
        "F.3 Fotodokumentace objektu",
    ]
    doc.add_paragraph(
        "K tomuto posudku jsou přiloženy následující přílohy:"
    )
    for p in prilohy:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_paragraph(
        "[Přílohy jsou součástí tištěné / elektronicky podepsané verze dokumentu.]"
    )

    # F.3 – fotodokumentace označená jako příloha
    foto_priloha = [f for f in budova.fotografie if f.sekce == "priloha"]
    if foto_priloha:
        doc.add_heading("F.3 Fotodokumentace", level=2)
        for foto in foto_priloha:
            _vloz_obrazek(doc, foto.data, foto.popisek, sirka_cm=foto.sirka_cm)


# ── Pomocné funkce ────────────────────────────────────────────────────────────

def _vloz_obrazek(doc: Document, buf, popisek: str = "",
                  sirka_cm: float = 14.0) -> None:
    """Vloží obrázek z bytes/BytesIO do dokumentu s popiskem. buf=None → nic."""
    if buf is None:
        return
    from io import BytesIO as _BytesIO
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    data = buf if isinstance(buf, _BytesIO) else _BytesIO(buf)
    data.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(4)
    p.add_run().add_picture(data, width=Cm(sirka_cm))
    if popisek:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(10)
        run_cap = cap.add_run(popisek)
        run_cap.italic = True
        run_cap.font.size = Pt(8)


def _nonempty(items: list) -> list[str]:
    """Filtruje None a prázdné řetězce ze seznamu."""
    return [x for x in items if x]


def _add_kv_row(table, klic: str, hodnota: str) -> None:
    """Přidá řádek klíč–hodnota do 2-sloupcové tabulky."""
    row = table.add_row()
    S.cell_text(row.cells[0], klic, bold=True, font_size=9)
    S.cell_text(row.cells[1], hodnota, font_size=9)
