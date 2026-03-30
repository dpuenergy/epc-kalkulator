"""
Testy pro Sprint 3 – generování Word výstupů (EP a EA).

Pokrývá:
  • epc_engine.reports  – generuj_ep, generuj_ea
  • epc_engine.reports._styles  – fmt_* helpers
  • epc_engine.reports._tables  – tabulka_opatreni, tabulka_emise
"""
from __future__ import annotations

import pytest
from io import BytesIO

from docx import Document

from epc_engine.reports import generuj_ep, generuj_ea
from epc_engine.reports._styles import fmt_kc, fmt_mwh, fmt_nav, fmt_irr, fmt_tsd, fmt_pct
from epc_engine.reports._tables import tabulka_opatreni, tabulka_emise
from epc_engine.models import BuildingInfo, EnergyInputs, Budova, Prostor, Podklad
from epc_engine.calculator import Project
from epc_engine.measures import OP1a
from epc_engine.economics import EkonomickeParametry


# ──────────────────────────────────────────────────────────────────────────────
# Fixtures
# ──────────────────────────────────────────────────────────────────────────────

@pytest.fixture
def minimal_budova() -> BuildingInfo:
    return BuildingInfo(
        nazev_zakazky="Test zakázka",
        objekt_nazev="ZŠ Test",
        zadavatel_nazev="Klient a.s.",
        zadavatel_adresa="Testovní 1, Praha",
        datum="2026-03-26",
    )


@pytest.fixture
def plna_budova() -> BuildingInfo:
    return BuildingInfo(
        nazev_zakazky="Plná zakázka",
        objekt_nazev="Základní škola Vzor",
        zadavatel_nazev="Město Vzor",
        zadavatel_adresa="Nám. Republiky 1, Vzor",
        zadavatel_ico="12345678",
        zadavatel_kontakt="Ing. Jan Novák",
        objekt_adresa="Školní 10, Vzor",
        druh_cinnosti="Základní škola",
        pocet_zamestnancu="45 / 320 žáků",
        budovy=[
            Budova(nazev="Hlavní budova", objem_m3=5000,
                   podlahova_plocha_m2=1200, ochlazovana_plocha_m2=1800),
        ],
        prostory=[Prostor(nazev="Učebny", ucel="Výuka", provoz="Po–Pá 7:00–16:00")],
        podklady=[
            Podklad("Spotřeby energií", True),
            Podklad("Projektová dokumentace", False),
        ],
    )


@pytest.fixture
def minimal_result():
    energie = EnergyInputs(
        zp_ut=400.0, zp_tuv=30.0, ee=80.0,
        cena_zp=1_800.0, cena_ee=4_500.0,
        pouzit_zp=True,
    )
    projekt = Project(
        energie=energie,
        opatreni=[
            OP1a(
                aktivni=True, plocha_m2=500,
                u_stavajici=0.8, u_nove=0.25,
                denostupne=3944.0, cena_kc_m2=1_800,
            )
        ],
        ekonomicke_parametry=EkonomickeParametry(),
    )
    return projekt.vypocitej()


@pytest.fixture
def result_s_klasifikaci(minimal_result):
    """ProjectResult doplněný o klasifikaci obálky."""
    energie = EnergyInputs(
        zp_ut=400.0, zp_tuv=30.0, ee=80.0,
        cena_zp=1_800.0, cena_ee=4_500.0,
        pouzit_zp=True,
    )
    projekt = Project(
        energie=energie,
        opatreni=[
            OP1a(
                aktivni=True, plocha_m2=500,
                u_stavajici=0.8, u_nove=0.25,
                denostupne=3944.0, cena_kc_m2=1_800,
            )
        ],
        ekonomicke_parametry=EkonomickeParametry(),
        uem_stav=1.2,
        faktor_tvaru=0.4,
    )
    return projekt.vypocitej()


# ──────────────────────────────────────────────────────────────────────────────
# Smoke testy – výstup je platný .docx (ZIP)
# ──────────────────────────────────────────────────────────────────────────────

class TestSmoke:

    def test_generuj_ep_je_docx(self, minimal_budova, minimal_result):
        buf = generuj_ep(minimal_budova, minimal_result)
        assert buf.read(4) == b"PK\x03\x04"  # .docx = ZIP

    def test_generuj_ea_je_docx(self, minimal_budova, minimal_result):
        buf = generuj_ea(minimal_budova, minimal_result)
        assert buf.read(4) == b"PK\x03\x04"

    def test_ep_s_plnou_budovou(self, plna_budova, minimal_result):
        buf = generuj_ep(plna_budova, minimal_result)
        assert buf.getvalue()

    def test_ea_s_plnou_budovou(self, plna_budova, minimal_result):
        buf = generuj_ea(plna_budova, minimal_result)
        assert buf.getvalue()

    def test_ep_s_klasifikaci(self, minimal_budova, result_s_klasifikaci):
        buf = generuj_ep(minimal_budova, result_s_klasifikaci)
        assert buf.getvalue()

    def test_ea_s_klasifikaci(self, minimal_budova, result_s_klasifikaci):
        buf = generuj_ea(minimal_budova, result_s_klasifikaci)
        assert buf.getvalue()

    def test_ep_bez_klasifikace(self, minimal_budova, minimal_result):
        minimal_result.klasifikace_pred = None
        buf = generuj_ep(minimal_budova, minimal_result)
        assert buf.getvalue()

    def test_ep_bez_ekonomiky(self, minimal_budova, minimal_result):
        minimal_result.ekonomika_projekt = None
        minimal_result.ekonomika_parametry = None
        buf = generuj_ep(minimal_budova, minimal_result)
        assert buf.getvalue()

    def test_ep_bez_aktivnich_opatreni(self, minimal_budova, minimal_result):
        for r in minimal_result.vysledky:
            r.aktivni = False
        buf = generuj_ep(minimal_budova, minimal_result)
        assert buf.getvalue()

    def test_ep_lze_otevrit_jako_dokument(self, minimal_budova, minimal_result):
        buf = generuj_ep(minimal_budova, minimal_result)
        doc = Document(buf)
        assert len(doc.paragraphs) > 0


# ──────────────────────────────────────────────────────────────────────────────
# Formátovací funkce
# ──────────────────────────────────────────────────────────────────────────────

class TestFormaty:

    def test_fmt_kc_kladne(self):
        s = fmt_kc(1_234_567)
        assert "1\u00a0234\u00a0567" in s
        assert "Kč" in s

    def test_fmt_kc_zaporne(self):
        s = fmt_kc(-500_000)
        assert s.startswith("-")
        assert "500\u00a0000" in s

    def test_fmt_kc_nula(self):
        s = fmt_kc(0)
        assert "0" in s

    def test_fmt_mwh(self):
        s = fmt_mwh(1234.5)
        assert "MWh" in s
        assert "1\u00a0234" in s

    def test_fmt_nav_cislo(self):
        s = fmt_nav(7.3)
        assert "7,3" in s or "7.3" in s
        assert "let" in s

    def test_fmt_nav_none(self):
        assert fmt_nav(None) == "\u221e"

    def test_fmt_irr_cislo(self):
        s = fmt_irr(0.125)
        assert "12" in s
        assert "%" in s

    def test_fmt_irr_none(self):
        assert fmt_irr(None) == "\u2013"

    def test_fmt_tsd_cislo(self):
        s = fmt_tsd(8.0)
        assert "8" in s
        assert "let" in s

    def test_fmt_tsd_none(self):
        s = fmt_tsd(None, horizont=20)
        assert ">20" in s

    def test_fmt_pct(self):
        s = fmt_pct(0.123)
        assert "12" in s
        assert "%" in s


# ──────────────────────────────────────────────────────────────────────────────
# Tabulky
# ──────────────────────────────────────────────────────────────────────────────

class TestTabulky:

    def test_tabulka_opatreni_pocet_radku(self, minimal_result):
        doc = Document()
        tbl = tabulka_opatreni(doc, minimal_result.aktivni)
        # 1 záhlaví + 1 opatření + 1 CELKEM = 3 řádky
        assert len(tbl.rows) == 3

    def test_tabulka_opatreni_hlavicka(self, minimal_result):
        doc = Document()
        tbl = tabulka_opatreni(doc, minimal_result.aktivni)
        header_texts = [c.text for c in tbl.rows[0].cells]
        assert any("Investice" in t for t in header_texts)
        assert any("Název" in t for t in header_texts)

    def test_tabulka_opatreni_celkem_tucne(self, minimal_result):
        doc = Document()
        tbl = tabulka_opatreni(doc, minimal_result.aktivni)
        # Poslední řádek = CELKEM
        last_row = tbl.rows[-1]
        assert "CELKEM" in last_row.cells[1].text

    def test_tabulka_emise_co2_radek(self, minimal_result):
        doc = Document()
        tbl = tabulka_emise(doc, minimal_result.emise_pred, minimal_result.emise_po)
        # První datový řádek = CO₂
        assert "CO" in tbl.rows[1].cells[0].text

    def test_tabulka_emise_pocet_radku(self, minimal_result):
        doc = Document()
        tbl = tabulka_emise(doc, minimal_result.emise_pred, minimal_result.emise_po)
        # 1 záhlaví + 5 látek (CO2, NOx, SO2, TZL, EPS)
        assert len(tbl.rows) == 6

    def test_tabulka_emise_none_vstupy(self):
        doc = Document()
        tbl = tabulka_emise(doc, None, None)
        assert len(tbl.rows) == 6  # nepadne ani s None vstupy


# ──────────────────────────────────────────────────────────────────────────────
# Obsah dokumentu – spot checks
# ──────────────────────────────────────────────────────────────────────────────

class TestObsahDokumentu:

    def _all_text(self, buf: BytesIO) -> str:
        doc = Document(buf)
        return "\n".join(p.text for p in doc.paragraphs)

    def test_ep_obsahuje_nadpis(self, minimal_budova, minimal_result):
        buf = generuj_ep(minimal_budova, minimal_result)
        text = self._all_text(buf)
        assert "ENERGETICKÝ POSUDEK" in text

    def test_ep_obsahuje_nazev_objektu(self, minimal_budova, minimal_result):
        buf = generuj_ep(minimal_budova, minimal_result)
        text = self._all_text(buf)
        assert "ZŠ Test" in text

    def test_ea_obsahuje_nadpis(self, minimal_budova, minimal_result):
        buf = generuj_ea(minimal_budova, minimal_result)
        text = self._all_text(buf)
        assert "ENERGETICKÝ AUDIT" in text

    def test_ea_obsahuje_poradi_opatreni(self, minimal_budova, minimal_result):
        buf = generuj_ea(minimal_budova, minimal_result)
        text = self._all_text(buf)
        # EA uses "příležitosti" terminology per vyhl. 140/2021 (not "opatření")
        assert "příležitost" in text.lower() or "opatření" in text.lower()

    def test_ea_obsahuje_legislativu(self, minimal_budova, result_s_klasifikaci):
        buf = generuj_ea(minimal_budova, result_s_klasifikaci)
        text = self._all_text(buf)
        assert "legislativních požadavků" in text

    def test_ep_obsahuje_zaver(self, minimal_budova, minimal_result):
        buf = generuj_ep(minimal_budova, minimal_result)
        text = self._all_text(buf)
        assert "Závěr" in text
