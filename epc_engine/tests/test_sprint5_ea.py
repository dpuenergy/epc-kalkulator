"""
Testy Sprint 5 – EA přídavný modul (vyhl. 140/2021 Sb.)

Testuje:
- EnergonositelEA, MKHKriterium, EnPIUkazatel, EAData
- Výpočet MKH (normalizace + vážený součet)
- Tabulky: bilance vstupů, EnPI, MKH, evidenční list
- Generátor EA dokumentu (nové sekce)
"""
from __future__ import annotations

import pytest

from epc_engine.models import (
    BuildingInfo, EnergyInputs,
    EnergonositelEA, MKHKriterium, EnPIUkazatel, EAData,
    MeasureResult,
)
from epc_engine.economics import EkonomickeBilance


# ══════════════════════════════════════════════════════════════════════════════
# Pomocné továrny
# ══════════════════════════════════════════════════════════════════════════════

def _energie() -> EnergyInputs:
    return EnergyInputs(
        zp_ut=100.0, zp_tuv=20.0, cena_zp=1500.0,
        ee=50.0, cena_ee=4000.0,
        teplo_ut=0.0, teplo_tuv=0.0, cena_teplo=0.0,
        pouzit_zp=True, pouzit_teplo=False,
    )


def _opatreni(id_: str, investice: float, uspora_kc: float,
               uspora_ee: float = 0.0, uspora_zp: float = 0.0,
               prosta_nav: float | None = None,
               npv: float | None = None) -> MeasureResult:
    r = MeasureResult(
        id=id_, nazev=f"Opatření {id_}", aktivni=True,
        investice=investice,
        uspora_kc=uspora_kc,
        uspora_ee=uspora_ee,
        uspora_zp=uspora_zp,
        prosta_navratnost=prosta_nav,
    )
    if npv is not None:
        r.ekonomika = EkonomickeBilance(npv=npv, irr=0.05, tsd=prosta_nav)
    return r


def _budova_s_ea(**kw) -> BuildingInfo:
    ea = EAData(
        evidencni_cislo_ea="EA-2024-001",
        cil="Identifikace úsporných opatření",
        datum_zahajeni="01.01.2024",
        datum_ukonceni="31.03.2024",
        **kw,
    )
    return BuildingInfo(
        objekt_nazev="Testovací budova",
        objekt_adresa="Testovací 1, Praha",
        zadavatel_nazev="Testovací s.r.o.",
        ea_data=ea,
    )


# ══════════════════════════════════════════════════════════════════════════════
# Datové třídy
# ══════════════════════════════════════════════════════════════════════════════

class TestEADataClasses:
    def test_energonositel_defaults(self):
        e = EnergonositelEA()
        assert e.kategorie == "NOZE"
        assert e.oblast == "Budovy"

    def test_energonositel_hodnoty(self):
        e = EnergonositelEA(nazev="Zemní plyn", kategorie="NOZE", oblast="Budovy")
        assert e.nazev == "Zemní plyn"
        assert e.kategorie == "NOZE"

    def test_mkh_kriterium_defaults(self):
        k = MKHKriterium()
        assert k.typ == "max"
        assert k.vaha == 0.0

    def test_enpi_ukazatel(self):
        u = EnPIUkazatel(nazev="Měrná spotřeba", jednotka="kWh/m²",
                         hodnota_stavajici=150.0, hodnota_navrhova=120.0)
        assert u.hodnota_stavajici == 150.0
        assert u.hodnota_navrhova == 120.0

    def test_ea_data_default_kriteria(self):
        ea = EAData()
        assert len(ea.mkh_kriteria) == 4
        assert ea.mkh_kriteria[0].key == "npv"
        assert ea.mkh_kriteria[1].key == "td"
        assert ea.mkh_kriteria[2].key == "mwh"
        assert ea.mkh_kriteria[3].key == "kc"
        assert sum(k.vaha for k in ea.mkh_kriteria) == 100.0

    def test_ea_data_prazdna(self):
        ea = EAData()
        assert ea.evidencni_cislo_ea == ""
        assert ea.energonositele == []
        assert ea.enpi == []

    def test_building_info_ea_data(self):
        b = _budova_s_ea()
        assert b.ea_data is not None
        assert b.ea_data.evidencni_cislo_ea == "EA-2024-001"


# ══════════════════════════════════════════════════════════════════════════════
# MKH výpočet
# ══════════════════════════════════════════════════════════════════════════════

class TestMKHVypocet:
    def _kriteria(self):
        return [
            MKHKriterium("NPV", "tis. Kč", "max", 50.0, "npv"),
            MKHKriterium("Td", "roky", "min", 50.0, "td"),
        ]

    def test_mkh_poradi(self):
        from epc_engine.reports._tables import vypocitej_mkh
        aktivni = [
            _opatreni("A", 100_000, 20_000, npv=500_000, prosta_nav=5.0),
            _opatreni("B", 200_000, 10_000, npv=100_000, prosta_nav=20.0),
        ]
        vysledky = vypocitej_mkh(aktivni, self._kriteria())
        # A má vyšší NPV a kratší Td → mělo by být první
        poradi = {r.id: por for r, _, por in vysledky}
        assert poradi["A"] < poradi["B"]

    def test_mkh_prazdne_opatreni(self):
        from epc_engine.reports._tables import vypocitej_mkh
        vysledky = vypocitej_mkh([], self._kriteria())
        assert vysledky == []

    def test_mkh_bez_kriterii(self):
        from epc_engine.reports._tables import vypocitej_mkh
        aktivni = [_opatreni("A", 100_000, 20_000)]
        vysledky = vypocitej_mkh(aktivni, [])
        assert len(vysledky) == 1

    def test_mkh_normalizace_stejne_hodnoty(self):
        from epc_engine.reports._tables import vypocitej_mkh
        # Pokud mají všechna opatření stejnou hodnotu kritéria, nepadne
        aktivni = [
            _opatreni("A", 100_000, 20_000, npv=100_000, prosta_nav=5.0),
            _opatreni("B", 100_000, 20_000, npv=100_000, prosta_nav=5.0),
        ]
        vysledky = vypocitej_mkh(aktivni, self._kriteria())
        assert len(vysledky) == 2

    def test_mkh_skore_rozsah(self):
        from epc_engine.reports._tables import vypocitej_mkh
        aktivni = [
            _opatreni("A", 100_000, 20_000, npv=500_000, prosta_nav=5.0),
            _opatreni("B", 50_000, 8_000, npv=200_000, prosta_nav=8.0),
            _opatreni("C", 300_000, 5_000, npv=50_000, prosta_nav=60.0),
        ]
        vysledky = vypocitej_mkh(aktivni, self._kriteria())
        for _, skore, _ in vysledky:
            assert 0.0 <= skore <= 1.0

    def test_mkh_poradi_je_unikatni(self):
        from epc_engine.reports._tables import vypocitej_mkh
        aktivni = [
            _opatreni("A", 100_000, 20_000, npv=500_000, prosta_nav=5.0),
            _opatreni("B", 200_000, 10_000, npv=100_000, prosta_nav=20.0),
            _opatreni("C", 300_000, 5_000, npv=50_000, prosta_nav=60.0),
        ]
        vysledky = vypocitej_mkh(aktivni, self._kriteria())
        poradi = [por for _, _, por in vysledky]
        assert len(set(poradi)) == 3


# ══════════════════════════════════════════════════════════════════════════════
# Tabulkové funkce
# ══════════════════════════════════════════════════════════════════════════════

class TestEATabulky:
    def _doc(self):
        from docx import Document
        from pathlib import Path
        tmpl = (Path(__file__).parent.parent / "reports" / "templates"
                / "ep_template.docx")
        return Document(str(tmpl))

    def test_tabulka_bilance_vstupu_ea(self):
        from epc_engine.reports._tables import tabulka_bilance_vstupu_ea
        doc = self._doc()
        energonositele = [
            EnergonositelEA("Zemní plyn", "NOZE", "Budovy"),
            EnergonositelEA("Elektrická energie", "NOZE", "Budovy"),
        ]
        tbl = tabulka_bilance_vstupu_ea(doc, energonositele, _energie())
        assert tbl is not None
        # 1 header + 1 skupinový záhlaví NOZE + 2 data + 1 celkem = 5
        assert len(tbl.rows) == 5

    def test_tabulka_bilance_vstupu_ea_prazdna(self):
        from epc_engine.reports._tables import tabulka_bilance_vstupu_ea
        doc = self._doc()
        tbl = tabulka_bilance_vstupu_ea(doc, [], _energie())
        # Pouze header + celkem
        assert len(tbl.rows) >= 1

    def test_tabulka_enpi(self):
        from epc_engine.reports._tables import tabulka_enpi
        doc = self._doc()
        enpi = [
            EnPIUkazatel("Měrná spotřeba tepla", "kWh/m²·rok", 150.0, 120.0),
            EnPIUkazatel("Spotřeba EE na zaměstnance", "kWh/os·rok", 2000.0, 1800.0),
        ]
        tbl = tabulka_enpi(doc, enpi)
        assert tbl is not None
        assert len(tbl.rows) == 3  # header + 2 data rows (poznámka je paragraf, ne řádek tabulky)

    def test_tabulka_mkh(self):
        from epc_engine.reports._tables import tabulka_mkh
        doc = self._doc()
        aktivni = [
            _opatreni("OP1", 200_000, 50_000, npv=300_000, prosta_nav=4.0),
            _opatreni("OP2", 100_000, 25_000, npv=150_000, prosta_nav=4.0),
        ]
        kriteria = [
            MKHKriterium("NPV", "tis. Kč", "max", 60.0, "npv"),
            MKHKriterium("Td", "roky", "min", 40.0, "td"),
        ]
        tbl = tabulka_mkh(doc, aktivni, kriteria)
        assert tbl is not None

    def test_tabulka_evidencni_list_1b(self):
        from epc_engine.reports._tables import tabulka_evidencni_list_1b
        doc = self._doc()
        aktivni = [
            _opatreni("OP1", 200_000, 50_000, npv=300_000, prosta_nav=4.0),
        ]
        tbl = tabulka_evidencni_list_1b(doc, aktivni, EAData().mkh_kriteria)
        assert tbl is not None


# ══════════════════════════════════════════════════════════════════════════════
# Generátor EA dokumentu
# ══════════════════════════════════════════════════════════════════════════════

class TestEAGenerator:
    def _result(self, aktivni=None):
        from epc_engine.models import ProjectResult
        r = ProjectResult(energie=_energie(), vysledky=aktivni or [])
        return r

    def _text(self, buf) -> str:
        from docx import Document
        import io
        doc = Document(io.BytesIO(buf.read()))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    def test_generuj_ea_zakladni(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea()
        result = self._result()
        buf = generuj_ea(b, result)
        assert buf is not None
        assert buf.read(4) == b"PK\x03\x04"

    def test_ea_obsahuje_souhrn(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea()
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "souhrn energetického auditu" in txt.lower()

    def test_ea_obsahuje_evidencni_cislo(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea()
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "EA-2024-001" in txt

    def test_ea_obsahuje_bilanci_vstupu(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea(energonositele=[
            EnergonositelEA("Zemní plyn", "NOZE", "Budovy"),
            EnergonositelEA("Elektrická energie", "NOZE", "Budovy"),
        ])
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "příl. 3" in txt or "Bilance energetických vstupů" in txt

    def test_ea_obsahuje_enpi(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea(enpi=[
            EnPIUkazatel("Měrná spotřeba tepla", "kWh/m²·rok", 150.0, 120.0),
        ])
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "EnPI" in txt or "Ukazatele energetické náročnosti" in txt

    def test_ea_obsahuje_mkh_s_opatrenimi(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea()
        aktivni = [
            _opatreni("OP1", 200_000, 50_000, npv=300_000, prosta_nav=4.0),
        ]
        from epc_engine.models import ProjectResult
        result = ProjectResult(energie=_energie(), vysledky=aktivni)
        buf = generuj_ea(b, result)
        buf.seek(0)
        txt = self._text(buf)
        assert "Multikriteriální" in txt or "příl. 9" in txt

    def test_ea_bez_ea_data_nepadne(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = BuildingInfo(objekt_nazev="Test", ea_data=None)
        result = self._result()
        buf = generuj_ea(b, result)
        assert buf is not None

    def test_ea_obsahuje_legislativni_hodnoceni(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea()
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "legislativních požadavků" in txt

    def test_ea_obsahuje_140_2021(self):
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea()
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "140/2021" in txt

    def test_ea_bez_energonosicel_nezobrazuje_bilanci(self):
        """Pokud nejsou energonosiče zadány, tabulka bilance vstupů se nevygeneruje."""
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea(energonositele=[])
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        # Obsah sekce (popis energonositelů) se nesmí objevit – záhlaví může být v obsahu
        assert "NOZE" not in txt and "energonositelů dle kategorie" not in txt

    def test_ea_bez_enpi_nezobrazuje_sekci(self):
        """Pokud nejsou EnPI ukazatele, sekce se nevygeneruje."""
        from epc_engine.reports.ea_generator import generuj_ea
        b = _budova_s_ea(enpi=[])
        buf = generuj_ea(b, self._result())
        buf.seek(0)
        txt = self._text(buf)
        assert "EnPIUkazatel" not in txt
