"""
Testy Sprint 4 – tepelně technické vlastnosti, nové sekce EP/EA, modely.

Pokrývá:
  - epc_engine.tepelna_technika
  - epc_engine.models (Vrstva, Konstrukce, TechnickeSystemy, BilancePouzitiEnergie,
                        PenbData, HistorieRok, KlimatickaData)
  - ep_generator + ea_generator s novými sekcemi
"""
from __future__ import annotations

import pytest
from io import BytesIO
from docx import Document

from epc_engine.tepelna_technika import (
    vypocitej_u_z_vrstev,
    un_pozadovana,
    hodnoceni_splneni,
    vypocitej_uem_z_konstrukci,
)
from epc_engine.models import (
    Vrstva, Konstrukce, TechnickySytem, TechnickeSystemy,
    BilancePouzitiEnergie, PenbData, HistorieRok, KlimatickaData,
    BuildingInfo, EnergyInputs,
)
from epc_engine.reports.ep_generator import generuj_ep
from epc_engine.reports.ea_generator import generuj_ea
from epc_engine.models import ProjectResult


# ── Helpers ───────────────────────────────────────────────────────────────────

def _minimal_result() -> ProjectResult:
    """Minimální ProjectResult bez žádných opatření."""
    return ProjectResult(energie=EnergyInputs(pouzit_teplo=True, teplo_ut=100.0, cena_teplo=800.0))


def _doc_text(buf: BytesIO) -> str:
    """Vrátí veškerý text z Word dokumentu včetně obsahu tabulek."""
    buf.seek(0)
    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
                # rekurzivně vnořené tabulky
                for nested in cell.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            parts.append(nc.text)
    return "\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# Tepelná technika – výpočet U z vrstev
# ══════════════════════════════════════════════════════════════════════════════

class TestVypocitejUZVrstev:
    def test_jednoducha_betonova_stena(self):
        """Beton 200 mm, λ=1.58: U = 1/(0.13 + 0.200/1.58 + 0.04) ≈ 3.37 W/m²K."""
        vrstvy = [Vrstva("Beton", tloustka_m=0.200, lambda_wm=1.58)]
        u = vypocitej_u_z_vrstev(vrstvy, "stena")
        assert 3.2 < u < 3.6

    def test_zateplena_stena_eps(self):
        """Beton 200mm + EPS 120mm → U < 0.35 W/m²K."""
        vrstvy = [
            Vrstva("Beton", 0.200, 1.58),
            Vrstva("EPS", 0.120, 0.037),
        ]
        u = vypocitej_u_z_vrstev(vrstvy, "stena")
        assert u < 0.35

    def test_strecha_odlisne_odpory(self):
        """Střecha má Rsi=0.10, Rse=0.04 – jiné než stěna."""
        vrstvy = [Vrstva("Min. vata", 0.200, 0.040)]
        u_stena = vypocitej_u_z_vrstev(vrstvy, "stena")
        u_strecha = vypocitej_u_z_vrstev(vrstvy, "strecha")
        # Rsi střechy (0.10) < Rsi stěny (0.13) → U střechy > U stěny
        assert u_strecha > u_stena

    def test_nulova_lambda_ignorovana(self):
        """Vrstva s λ=0 nesmí způsobit ZeroDivisionError."""
        vrstvy = [
            Vrstva("Vzduch", 0.050, 0.0),    # λ=0 → ignorovat
            Vrstva("EPS", 0.100, 0.037),
        ]
        u = vypocitej_u_z_vrstev(vrstvy, "stena")
        assert u > 0

    def test_prazdne_vrstvy_vraci_pouze_odpory(self):
        """Prázdný seznam vrstev → U se vypočítá jen z povrchových odporů (Rsi+Rse)."""
        u = vypocitej_u_z_vrstev([], "stena")
        # U = 1/(Rsi+Rse) = 1/(0.13+0.04) = 5.88 – žádná izolace, ale není 0
        assert u > 0
        assert u == pytest.approx(1.0 / (0.13 + 0.04), rel=1e-3)

    def test_podlaha_k_zemini_rsi_vetsi(self):
        """Podlaha má Rsi=0.17 (větší než stěna Rsi=0.13), celkový Rtotal je stejný (0.17+0=0.13+0.04).
        Ověříme jen, že funkce nepadá a vrátí kladné U."""
        vrstvy = [Vrstva("Beton", 0.150, 1.58)]
        u_podlaha = vypocitej_u_z_vrstev(vrstvy, "podlaha")
        assert u_podlaha > 0
        # Podlaha Rtotal = 0.17 + 0.0950 + 0.00 = 0.2650; U ≈ 3.77 W/m²K
        assert u_podlaha == pytest.approx(1.0 / (0.17 + 0.150/1.58 + 0.00), rel=1e-3)


class TestUnPozadovana:
    def test_normove_hodnoty(self):
        assert un_pozadovana("stena") == pytest.approx(0.30)
        assert un_pozadovana("strecha") == pytest.approx(0.24)
        assert un_pozadovana("podlaha") == pytest.approx(0.45)
        assert un_pozadovana("okno") == pytest.approx(1.50)
        assert un_pozadovana("dvere") == pytest.approx(1.70)

    def test_neznamy_typ_vrati_default(self):
        assert un_pozadovana("neznamy_typ") == pytest.approx(0.30)


class TestHodnoceniSplneni:
    def test_vyhovuje(self):
        assert hodnoceni_splneni(0.25, 0.30) == "Vyhovuje"

    def test_nevyhovuje(self):
        assert hodnoceni_splneni(0.35, 0.30) == "Nevyhovuje"

    def test_presne_rovno_vyhovuje(self):
        assert hodnoceni_splneni(0.30, 0.30) == "Vyhovuje"

    def test_nulovy_un(self):
        assert hodnoceni_splneni(0.25, 0.0) == "–"


class TestVypocitejUemZKonstrukci:
    def test_jednoducha_stena(self):
        """Jedna stěna, A=100m², U=0.25 → Uem ≈ 0.25 + 0.02 = 0.27 W/m²K."""
        k = Konstrukce("Stěna", "stena", plocha_m2=100.0, u_zadane=0.25, un_value=0.30)
        uem = vypocitej_uem_z_konstrukci([k])
        assert uem == pytest.approx(0.27, abs=0.001)

    def test_prazdny_seznam(self):
        assert vypocitej_uem_z_konstrukci([]) is None

    def test_nulova_plocha(self):
        k = Konstrukce(plocha_m2=0.0, u_zadane=0.3)
        assert vypocitej_uem_z_konstrukci([k]) is None


# ══════════════════════════════════════════════════════════════════════════════
# Datové třídy Sprint 4
# ══════════════════════════════════════════════════════════════════════════════

class TestKonstrukce:
    def test_u_effective_zadane(self):
        """Pokud je u_zadane, přebíjí výpočet z vrstev."""
        k = Konstrukce(u_zadane=1.1, vrstvy=[Vrstva("Beton", 0.2, 1.58)])
        assert k.u_effective == pytest.approx(1.1)

    def test_u_effective_z_vrstev(self):
        """Bez u_zadane se U vypočítá z vrstev."""
        k = Konstrukce(typ="stena", vrstvy=[Vrstva("EPS", 0.120, 0.037)])
        assert k.u_effective < 0.40

    def test_u_effective_bez_vrstev_a_bez_zadane(self):
        """Prázdné vrstvy a žádné u_zadane → U z čistých povrchových odporů (> 0)."""
        k = Konstrukce(typ="stena")
        assert k.u_effective > 0

    def test_validni_typy(self):
        """Všechny typy z TYP_POPISY jsou platné."""
        from epc_engine.tepelna_technika import TYP_POPISY
        for typ in TYP_POPISY:
            k = Konstrukce(typ=typ, vrstvy=[Vrstva("X", 0.1, 0.5)])
            assert k.u_effective > 0


class TestBilancePouzitiEnergie:
    def test_celkem_mwh(self):
        b = BilancePouzitiEnergie(vytapeni_mwh=100, osvetleni_mwh=20, technologie_mwh=50)
        assert b.celkem_mwh == pytest.approx(170.0)

    def test_celkem_kc(self):
        b = BilancePouzitiEnergie(vytapeni_kc=500_000, tuv_kc=100_000)
        assert b.celkem_kc == pytest.approx(600_000.0)


class TestHistorieRokKlimatickaData:
    def test_historierok_defaulty(self):
        h = HistorieRok()
        assert h.rok == 0
        assert h.spotreba_mwh == 0.0

    def test_klimaticka_data_defaulty(self):
        kd = KlimatickaData()
        assert kd.stupnodni_normovane == pytest.approx(3600.0)
        assert kd.teplota_vnitrni == pytest.approx(19.0)


# ══════════════════════════════════════════════════════════════════════════════
# Generátor EP – nové sekce
# ══════════════════════════════════════════════════════════════════════════════

class TestEpGeneratorNoveSekceSprint4:
    def _budova_plna(self) -> BuildingInfo:
        """BuildingInfo se všemi Sprint 4 poli vyplněnými."""
        return BuildingInfo(
            objekt_nazev="Testovací objekt",
            objekt_adresa="Testovací 1, Praha",
            objekt_ku="Praha-Dejvice",
            objekt_parcelni_cislo="123/4",
            evidencni_cislo="ENEX-001",
            cislo_opravneni="OP-12345",
            ucel_ep="EPC studie",
            ceny_bez_dph=True,
            klimaticka_data=KlimatickaData(
                lokalita="Praha",
                stupnodni_normovane=3600.0,
                teplota_vnitrni=20.0,
                teplota_exterieru=-12.0,
            ),
            historie_spotreby=[
                HistorieRok(rok=2022, energonosic="Teplo (CZT)", spotreba_mwh=500.0,
                            naklady_kc=400_000, stupnodni=3500),
                HistorieRok(rok=2023, energonosic="Teplo (CZT)", spotreba_mwh=480.0,
                            naklady_kc=420_000, stupnodni=3400),
            ],
            konstrukce=[
                Konstrukce("Obvodová stěna", "stena", plocha_m2=800.0,
                           vrstvy=[Vrstva("Beton", 0.200, 1.58), Vrstva("EPS", 0.120, 0.037)],
                           un_value=0.30),
                Konstrukce("Střecha", "strecha", plocha_m2=300.0,
                           vrstvy=[Vrstva("Min. vata", 0.200, 0.040)],
                           un_value=0.24),
                Konstrukce("Okna", "okno", plocha_m2=120.0, u_zadane=1.1, un_value=1.50),
            ],
            technicke_systemy=TechnickeSystemy(
                vytapeni=TechnickySytem(
                    typ="Horkovod CZT", vykon_kw=500.0, ucinnost_pct=95.0,
                    rok_instalace=2015,
                    popis="Napojení na centrální zásobování teplem, předávací stanice v suterénu.",
                ),
                osvetleni=TechnickySytem(
                    typ="LED osvětlení", popis="Osvětlení chodeb a kancelářských prostor.",
                ),
            ),
            bilance_pouziti=BilancePouzitiEnergie(
                vytapeni_mwh=400.0, tuv_mwh=80.0, osvetleni_mwh=50.0,
                vytapeni_kc=320_000, tuv_kc=64_000, osvetleni_kc=40_000,
            ),
            penb=PenbData(
                trida_stavajici="D", trida_navrhovy="C",
                merná_potreba_tepla=120.0, celkova_dodana_energie=180.0,
                primarni_neobnovitelna=200.0, energeticka_vztazna_plocha=2500.0,
                poznamka="PENB zpracován Ing. Novák, 03/2026",
            ),
            okrajove_podminky=(
                "Výpočty vychází z naměřených spotřeb a prohlídky objektu dne 10. 3. 2026."
            ),
        )

    def test_ep_obsahuje_identifikacni_pole(self):
        """Dokument musí obsahovat KÚ, parcelní číslo, číslo oprávnění."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Praha-Dejvice" in text
        assert "123/4" in text
        assert "OP-12345" in text

    def test_ep_obsahuje_historii_spotreby(self):
        """Sekce B.2 – tabulka historické spotřeby."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "2022" in text
        assert "2023" in text
        assert "Teplo (CZT)" in text

    def test_ep_obsahuje_klimaticka_data(self):
        """Sekce B.2 – klimatická lokalita."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Praha" in text

    def test_ep_obsahuje_tepelne_technicke_vlastnosti(self):
        """Sekce B.3 – tabulka konstrukcí s U-hodnotami."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Obvodová stěna" in text
        assert "Střecha" in text
        assert "Okna" in text
        assert "Tepelně technické" in text

    def test_ep_obsahuje_technicke_systemy(self):
        """Sekce B.4 – popis technických systémů."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Horkovod CZT" in text
        assert "LED osvětlení" in text

    def test_ep_obsahuje_bilanci_pouziti(self):
        """Sekce C.4 – tabulka dle přílohy č. 4 vyhl. 141/2021."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Způsob užití" in text or "způsobu užití" in text
        assert "Vytápění" in text

    def test_ep_obsahuje_penb(self):
        """Sekce D.4 – PENB data."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Primární energie" in text
        assert "Ing. Novák" in text

    def test_ep_obsahuje_okrajove_podminky(self):
        """Sekce E – okrajové podmínky."""
        budova = self._budova_plna()
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "prohlídky objektu" in text

    def test_ep_bez_penb_nepadne(self):
        """penb=None nesmí způsobit chybu – je to volitelné pole."""
        budova = BuildingInfo(penb=None)
        buf = generuj_ep(budova, _minimal_result())
        assert buf is not None
        text = _doc_text(buf)
        assert "PENB je zpracováván samostatně" in text

    def test_ep_bez_konstrukci_nepadne(self):
        """Prázdný seznam konstrukcí nesmí způsobit chybu."""
        budova = BuildingInfo(konstrukce=[])
        buf = generuj_ep(budova, _minimal_result())
        assert buf is not None

    def test_ep_bez_historie_nepadne(self):
        """Prázdná historická data nesmí způsobit chybu."""
        budova = BuildingInfo(historie_spotreby=[], klimaticka_data=None)
        buf = generuj_ep(budova, _minimal_result())
        assert buf is not None


# ══════════════════════════════════════════════════════════════════════════════
# Generátor EA – nové sekce
# ══════════════════════════════════════════════════════════════════════════════

class TestEaGeneratorNoveSekceSprint4:
    def test_ea_obsahuje_technicke_systemy(self):
        """EA musí také obsahovat technické systémy."""
        budova = BuildingInfo(
            technicke_systemy=TechnickeSystemy(
                vytapeni=TechnickySytem(typ="Plynový kotel Buderus", vykon_kw=120.0),
            )
        )
        buf = generuj_ea(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Buderus" in text

    def test_ea_obsahuje_historii(self):
        """EA musí zahrnout historii spotřeby."""
        budova = BuildingInfo(
            historie_spotreby=[
                HistorieRok(rok=2021, energonosic="Zemní plyn", spotreba_mwh=300.0),
            ]
        )
        buf = generuj_ea(budova, _minimal_result())
        text = _doc_text(buf)
        assert "2021" in text
        assert "Zemní plyn" in text

    def test_ea_obsahuje_ucel_ep(self):
        """Účel EP/EA musí být v dokumentu."""
        budova = BuildingInfo(ucel_ep="Dotace OPŽP")
        buf = generuj_ea(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Dotace OPŽP" in text


# ══════════════════════════════════════════════════════════════════════════════
# Sprint 5 – Nová struktura A/B/C/D/E + EnMS + měsíční tabulky + U-porovnání
# ══════════════════════════════════════════════════════════════════════════════

from epc_engine.models import EnMSOblast, EnMSHodnoceni
from epc_engine.tepelna_technika import urec_doporucena, upas_pasivni


class TestEnMSModel:
    def test_enms_hodnoceni_default_oblasti(self):
        """EnMSHodnoceni má 7 defaultních oblastí."""
        enms = EnMSHodnoceni()
        assert len(enms.oblasti) == 7

    def test_enms_oblast_hodnoceni_range(self):
        """Hodnocení oblasti je 1–3."""
        o = EnMSOblast("Test", "Stav", 2)
        assert o.hodnoceni == 2

    def test_enms_certifikace_default_false(self):
        enms = EnMSHodnoceni()
        assert enms.certifikovan is False


class TestUrec_Upas:
    def test_urec_stena(self):
        assert urec_doporucena("stena") == pytest.approx(0.25)

    def test_upas_stena(self):
        assert upas_pasivni("stena") == pytest.approx(0.18)

    def test_urec_strecha(self):
        assert urec_doporucena("strecha") == pytest.approx(0.16)

    def test_upas_okno(self):
        assert upas_pasivni("okno") == pytest.approx(0.80)


class TestMesicniTabulka:
    def test_mesicni_mwh_default_length(self):
        """HistorieRok.mesicni_mwh má 12 hodnot."""
        h = HistorieRok(rok=2023)
        assert len(h.mesicni_mwh) == 12

    def test_mesicni_mwh_soucet(self):
        """Součet měsíčních hodnot odpovídá roční spotřebě."""
        mesice = [50.0, 45.0, 40.0, 20.0, 10.0, 5.0, 5.0, 5.0, 10.0, 20.0, 40.0, 50.0]
        h = HistorieRok(rok=2023, energonosic="Teplo (CZT)",
                        spotreba_mwh=sum(mesice), mesicni_mwh=mesice)
        assert h.spotreba_mwh == pytest.approx(300.0)

    def test_ep_s_mesicnimi_daty(self):
        """EP s měsíčními daty nesmí padat a musí obsahovat název energonosiče."""
        mesice = [40.0] * 12
        budova = BuildingInfo(
            historie_spotreby=[
                HistorieRok(rok=2023, energonosic="Zemní plyn",
                            spotreba_mwh=480.0, mesicni_mwh=mesice),
                HistorieRok(rok=2024, energonosic="Zemní plyn",
                            spotreba_mwh=460.0, mesicni_mwh=[38.0] * 12),
            ]
        )
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Zemní plyn" in text
        assert "2023" in text
        assert "2024" in text


class TestStrukturaDokumentu:
    """Ověří, že dokument obsahuje nové nadpisy A–E podle vyhlášky."""

    def test_ep_obsahuje_nadpis_a(self):
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "IDENTIFIKAČNÍ ÚDAJE" in text.upper() or "Identifikační" in text

    def test_ep_obsahuje_nadpis_b(self):
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "POPIS STÁVAJÍCÍHO STAVU" in text.upper() or "stávající" in text.lower()

    def test_ep_obsahuje_nadpis_c(self):
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "VYHODNOCENÍ" in text.upper()

    def test_ep_obsahuje_nadpis_d(self):
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "NAVRHOVANÁ OPATŘENÍ" in text.upper() or "Navrhovaná" in text

    def test_ep_obsahuje_nadpis_e(self):
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "OKRAJOVÉ PODMÍNKY" in text.upper() or "Okrajové" in text

    def test_ep_obsahuje_prilohy(self):
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "PŘÍLOHY" in text.upper() or "Přílohy" in text

    def test_ep_s_enms(self):
        """Dokument s EnMS musí obsahovat hodnocení EnMS."""
        enms = EnMSHodnoceni(
            certifikovan=False,
            oblasti=[
                EnMSOblast("Energetická politika", "Bez formální politiky.", 1),
                EnMSOblast("Energetické plánování", "Základní plánování probíhá.", 2),
            ],
            komentar="EnMS v počáteční fázi zavádění.",
        )
        budova = BuildingInfo(enms=enms)
        buf = generuj_ep(budova, _minimal_result())
        text = _doc_text(buf)
        assert "Energetická politika" in text
        assert "EnMS v počáteční fázi" in text

    def test_ep_obsahuje_legislativu(self):
        """Sekce A.3 musí obsahovat odkaz na vyhlášku č. 141/2021 Sb."""
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "141/2021" in text

    def test_ep_obsahuje_csn_6946(self):
        """Sekce E musí zmiňovat ČSN EN ISO 6946."""
        buf = generuj_ep(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "6946" in text

    def test_ea_obsahuje_legislativni_hodnoceni(self):
        """EA musí obsahovat sekci hodnocení legislativních požadavků."""
        buf = generuj_ea(BuildingInfo(), _minimal_result())
        text = _doc_text(buf)
        assert "legislativních" in text.lower() or "Legislativní" in text

    def test_ep_bez_enms_nepadne(self):
        """EP bez EnMS (enms=None) nesmí padat."""
        budova = BuildingInfo(enms=None)
        buf = generuj_ep(budova, _minimal_result())
        assert buf is not None
